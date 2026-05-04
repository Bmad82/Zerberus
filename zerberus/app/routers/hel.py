"""
Hel Router – Admin-Dashboard und Konfiguration.
"""
import logging
import json
import re
import secrets
import os
import tempfile
import asyncio
import io
from pathlib import Path
from typing import Optional
from fastapi import APIRouter, Request, HTTPException, Depends, status, UploadFile, File, Form
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from fastapi.responses import HTMLResponse, JSONResponse
import httpx
import yaml

from zerberus.core.config import get_settings
from zerberus.core.database import get_all_sessions, get_session_messages, get_latest_metrics, get_metrics_summary, get_message_costs, get_last_cost
from zerberus.core.config import get_settings, reload_settings, invalidates_settings, Settings
from zerberus.core.cleaner import clean_transcript
from zerberus.modules.telegram.bot import DEFAULT_HUGINN_PROMPT  # Patch 158
from zerberus.core.dialect import detect_dialect_marker, apply_dialect
from zerberus.core.llm import LLMService
import hashlib

try:
    from docx import Document as DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    import pdfplumber
    _PDF_OK = True
except ImportError:
    _PDF_OK = False

logger = logging.getLogger(__name__)


def _sanitize_unicode(text: str) -> str:
    """Entfernt Surrogate-Pairs und andere ungültige Unicode-Zeichen."""
    if not isinstance(text, str):
        return str(text)
    return text.encode('utf-8', errors='replace').decode('utf-8', errors='replace')


security = HTTPBasic()

def verify_admin(credentials: HTTPBasicCredentials = Depends(security)):
    """Prüft Admin-Credentials über HTTP Basic Auth."""
    correct_username = os.getenv("ADMIN_USER", "admin")
    correct_password = os.getenv("ADMIN_PASSWORD", "admin")
    if correct_username == "admin" and correct_password == "admin":
        print("⚠️  WARNUNG: Admin-Zugang mit Standard-Credentials (admin/admin) aktiv! Bitte in .env ändern.")
    is_correct_username = secrets.compare_digest(credentials.username, correct_username)
    is_correct_password = secrets.compare_digest(credentials.password, correct_password)
    if not (is_correct_username and is_correct_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Basic"},
        )
    return credentials.username





router = APIRouter(prefix="/hel", tags=["Hel"], dependencies=[Depends(verify_admin)])

# ========== DEBUG‑ENDKNOTEN ==========
@router.get("/debug/trace/{session_id}")
async def debug_trace(session_id: str, request: Request):
    """Zeigt detaillierte Trace‑Info für eine Session."""
    from zerberus.core.database import get_session_messages
    from zerberus.core.config import get_settings
    import json as jsonlib

    messages = await get_session_messages(session_id)
    if not messages:
        raise HTTPException(404, "Session nicht gefunden")

    # Zusätzlich: letzte LLM‑Prompt aus dem Cost‑Table? (nicht direkt)
    # Wir holen die letzte Assistant‑Nachricht und deren Metadaten
    last_assistant = None
    for m in reversed(messages):
        if m["role"] == "assistant":
            last_assistant = m
            break

    # Patch 112: config.yaml ist Single Source of Truth. config.json nur noch
    # als Legacy-Hinweis anzeigen, falls Datei noch existiert.
    settings = get_settings()
    config_json_path = Path("config.json")
    legacy_json_present = config_json_path.exists()

    return {
        "session_id": session_id,
        "message_count": len(messages),
        "last_assistant": last_assistant,
        "config": {
            "from_yaml": {
                "cloud_model": settings.legacy.models.cloud_model,
                "temperature": settings.legacy.settings.ai_temperature,
                "threshold": settings.legacy.settings.threshold_length,
            },
            "legacy_config_json_present": legacy_json_present,
        },
        "active_modules": ["emotional", "nudge", "preparer", "rag"],  # später dynamisch
    }

@router.get("/debug/state")
async def debug_state():
    """Zeigt System‑Zustand: FAISS‑Status, Sessions, Config‑Hash, Modul‑Health."""
    from zerberus.core.database import _async_session_maker
    from sqlalchemy import text
    import hashlib

    # FAISS‑Status prüfen
    faiss_ok = False
    try:
        import faiss
        # Prüfe, ob Index geladen (später)
        faiss_ok = True
    except ImportError:
        faiss_ok = False

    # Session‑Anzahl
    session_count = 0
    async with _async_session_maker() as session:
        result = await session.execute(text("SELECT COUNT(DISTINCT session_id) FROM interactions"))
        session_count = result.scalar() or 0

    # Patch 112: Hash von config.yaml (Single Source of Truth)
    config_hash = None
    config_path = Path("config.yaml")
    if config_path.exists():
        with open(config_path, "rb") as f:
            config_hash = hashlib.sha256(f.read()).hexdigest()[:8]

    return {
        "faiss_available": faiss_ok,
        "active_sessions": session_count,
        "config_hash": config_hash,
        "modules": {
            "emotional": {"enabled": True, "last_ok": None},
            "nudge": {"enabled": True, "last_ok": None},
            "preparer": {"enabled": True, "last_ok": None},
            "rag": {"enabled": True, "last_ok": None},
        }
    }


# Konfigurationsdateien im Projekt-Root
WHISPER_CLEANER_PATH = Path("whisper_cleaner.json")
FUZZY_DICT_PATH = Path("fuzzy_dictionary.json")
DIALECT_PATH = Path("dialect.json")
SYSTEM_PROMPT_PATH = Path("system_prompt.json")

# HTML-Template mit HTML-Entities für Emojis
ADMIN_HTML = """<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0, user-scalable=yes">
    <title>&#9889; Hel – Zerberus Admin</title>
    <link rel="icon" href="/static/favicon.ico">
    <!-- Patch 200: PWA-Verdrahtung — Manifest, Theme-Color, Apple-Meta, Touch-Icons. -->
    <link rel="manifest" href="/hel/manifest.json">
    <meta name="theme-color" content="#1a1a1a">
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
    <meta name="apple-mobile-web-app-title" content="Hel">
    <link rel="apple-touch-icon" href="/static/pwa/hel-192.png">
    <link rel="apple-touch-icon" sizes="512x512" href="/static/pwa/hel-512.png">
    <!-- Patch 151 (L-001): Gemeinsame Design-Tokens für Nala UND Hel. -->
    <link rel="stylesheet" href="/static/css/shared-design.css">
    <!-- Patch 91: Chart.js 4.4.7 + Zoom-Plugin + hammerjs (für Touch-Pinch-Zoom) -->
    <script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.7/dist/chart.umd.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/hammerjs@2.0.8/hammer.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/chartjs-plugin-zoom@2.0.1/dist/chartjs-plugin-zoom.min.js"></script>
    <script>
        // Patch 90 (N-F09b): Schriftgröße früh laden, vermeidet FOUC
        (function () {
            try {
                var fs = localStorage.getItem('hel_font_size');
                if (fs) document.documentElement.style.setProperty('--hel-font-size-base', fs);
            } catch (_) {}
        })();
        // Patch 99 (H-F01): aktiven Tab früh setzen, vermeidet FOUC
        window.__hel_active_tab = (function () {
            try {
                var t = localStorage.getItem('hel_active_tab');
                return t || 'metrics';
            } catch (_) { return 'metrics'; }
        })();
    </script>
    <style>
        :root { --hel-font-size-base: 15px; }
        * { box-sizing: border-box; }
        body {
            font-family: 'Segoe UI', system-ui, sans-serif;
            background: #1a1a1a;
            color: #e0e0e0;
            margin: 0;
            padding: 20px;
            font-size: var(--hel-font-size-base);
        }
        .container { max-width: 1400px; margin: 0 auto; }
        h1 { color: #ff6b6b; border-bottom: 2px solid #ff6b6b; padding-bottom: 10px; }
        /* Patch 99 (H-F01): Sticky Tab-Leiste ersetzt Akkordeon.
           Die .hel-section-header bleiben im HTML für Rückwärtskompatibilität,
           werden aber per CSS versteckt. .hel-section-body verliert seine
           collapsed/max-height-Mechanik — die gesamte .hel-section wird per
           data-tab + display:none/block getoggelt. */
        .hel-section { margin-bottom: 4px; }
        .hel-section[data-tab]:not(.active) { display: none; }
        .hel-section-header { display: none !important; }
        .hel-section-body { background: #2d2d2d; border-radius: 12px; padding: 20px; overflow: visible; }
        .hel-section-body.collapsed { max-height: none !important; padding: 20px; }

        .hel-tab-nav {
            position: sticky;
            top: 0;
            z-index: 100;
            display: flex;
            gap: 4px;
            overflow-x: auto;
            overflow-y: hidden;
            white-space: nowrap;
            background: #1a1a1a;
            border-bottom: 2px solid #c8941f;
            margin: 0 -20px 14px -20px;
            padding: 6px 14px 0 14px;
            scrollbar-width: none;
            -ms-overflow-style: none;
            -webkit-overflow-scrolling: touch;
        }
        .hel-tab-nav::-webkit-scrollbar { display: none; }
        .hel-tab {
            flex: 0 0 auto;
            min-height: 44px;
            padding: 8px 16px;
            background: transparent;
            color: #c8ccd0;
            border: none;
            border-bottom: 3px solid transparent;
            font-size: 15px;
            font-weight: 600;
            cursor: pointer;
            user-select: none;
            -webkit-tap-highlight-color: transparent;
        }
        .hel-tab:hover, .hel-tab:active { color: #ffd700; }
        .hel-tab.active {
            color: #ffd700;
            border-bottom-color: #ffd700;
        }
        .card { background: #3d3d3d; border-radius: 12px; padding: 20px; margin-bottom: 20px; }
        label { display: block; margin: 15px 0 5px; color: #ffa5a5; }
        select, textarea, input {
            width: 100%;
            padding: 12px;
            background: #3d3d3d;
            border: 1px solid #555;
            color: white;
            border-radius: 8px;
            font-size: var(--hel-font-size-base);
        }
        .hel-section-body { font-size: var(--hel-font-size-base); }
        .hel-section-body table th, .hel-section-body table td { font-size: var(--hel-font-size-base); }
        .font-preset-bar {
            display: flex;
            gap: 8px;
            align-items: center;
            margin: 0 0 16px 0;
            flex-wrap: wrap;
        }
        .font-preset-bar .label { color:#aaa; font-size:0.85em; margin-right:4px; }
        .font-preset-btn {
            min-width: 44px;
            min-height: 44px;
            padding: 8px 12px;
            background: #2d2d2d;
            color: #e0e0e0;
            border: 1px solid #555;
            border-radius: 8px;
            cursor: pointer;
            font-weight: bold;
            margin-top: 0;
            -webkit-tap-highlight-color: transparent;
        }
        .font-preset-btn:active { background: #3d3d3d; }
        .font-preset-btn.active { background: #c8941f; color: #1a1a1a; border-color: #ffd700; }
        /* Patch 90 (H-F02): Whisper-Cleaner Karten */
        .cleaner-list { max-height: 60vh; overflow-y: auto; padding-right: 4px; }
        .cleaner-card {
            background: #2d2d2d;
            border: 1px solid #444;
            border-radius: 10px;
            padding: 12px;
            margin-bottom: 10px;
        }
        .cleaner-card.invalid { border-color: #ff6b6b; box-shadow: 0 0 0 1px rgba(255,107,107,0.4); }
        .cleaner-card-row { display: flex; gap: 8px; align-items: flex-start; flex-wrap: wrap; }
        .cleaner-card-grow { flex: 1 1 220px; min-width: 0; }
        .cleaner-card label {
            display: block;
            color: #ffa5a5;
            font-size: 0.78em;
            margin: 6px 0 3px;
        }
        .cleaner-card input[type="text"] {
            width: 100%;
            padding: 9px 11px;
            background: #1f1f1f;
            border: 1px solid #555;
            color: #f0f0f0;
            border-radius: 6px;
            font-family: 'Consolas', 'Courier New', monospace;
            font-size: 0.95em;
        }
        .cleaner-card.invalid input[name="pattern"] { border-color: #ff6b6b; }
        .cleaner-error { color: #ff6b6b; font-size: 0.8em; margin-top: 4px; min-height: 1em; }
        .cleaner-trash {
            background: #4a1a1a !important;
            color: #ffb3b3 !important;
            border: 1px solid #6a2a2a !important;
            min-width: 44px;
            min-height: 44px;
            padding: 8px 12px !important;
            margin-top: 0 !important;
            border-radius: 8px !important;
            font-weight: normal !important;
        }
        .cleaner-trash:active { background: #6a2a2a !important; }
        .cleaner-section {
            background: #1a1a1a;
            border-left: 3px solid #c8941f;
            border-radius: 6px;
            padding: 10px 12px;
            margin: 14px 0 8px;
            display: flex;
            gap: 8px;
            align-items: center;
        }
        .cleaner-section input[type="text"] {
            flex: 1;
            background: transparent;
            border: none;
            color: #ffd700;
            font-weight: bold;
            font-size: 0.95em;
            padding: 4px;
        }
        .cleaner-section input[type="text"]:focus { background: #252525; outline: 1px solid #555; border-radius: 4px; }
        /* Patch 90 (N-F10): Landscape mit niedrigem Viewport */
        @media (orientation: landscape) and (max-height: 500px) {
            body { padding: 10px; }
            h1 { font-size: 1.2em; padding-bottom: 6px; margin-top: 0; }
            .hel-section-header { height: 40px; font-size: 14px; }
            .hel-section-body { padding: 12px; }
            .card { padding: 12px; margin-bottom: 12px; }
            .font-preset-bar { margin: 0 0 10px 0; }
        }
        button {
            background: #ff6b6b;
            color: black;
            border: none;
            padding: 12px 30px;
            border-radius: 30px;
            font-weight: bold;
            cursor: pointer;
            margin-top: 10px;
            font-size: 16px;
        }
        button:hover, button:active { background: #ff5252; }
        .balance {
            background: #252525;
            padding: 15px;
            border-radius: 8px;
            font-family: monospace;
            font-size: 1.2em;
            margin: 10px 0;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            background: #3d3d3d;
            border-radius: 8px;
            overflow: hidden;
        }
        th, td { padding: 12px; text-align: left; border-bottom: 1px solid #555; }
        th { background: #4d4d4d; }
        .session-item {
            background: #3d3d3d;
            padding: 12px;
            border-radius: 8px;
            margin-bottom: 8px;
            cursor: pointer;
            display: flex;
            justify-content: space-between;
        }
        .session-item:hover, .session-item:active { background: #4d4d4d; }
        /* Patch 95: Per-User-Filter */
        .metric-profile-filter {
            display: flex; align-items: center; gap: 8px; flex-wrap: wrap;
            margin-bottom: 10px;
        }
        .metric-profile-filter label {
            font-size: 13px; color: #c8ccd0;
        }
        .profile-select {
            width: auto; padding: 6px 14px; border-radius: 16px;
            border: 1px solid rgba(240,180,41,0.3);
            background: #2d2d2d; color: #c8ccd0;
            font-size: 13px; cursor: pointer; min-height: 36px;
            -webkit-tap-highlight-color: transparent;
        }
        .profile-select:focus, .profile-select:active {
            border-color: #f0b429; outline: none;
        }
        /* Patch 91: Zeitraum-Chips */
        .metric-timerange { display: flex; gap: 6px; flex-wrap: wrap; margin-bottom: 12px; }
        .time-chip {
            padding: 6px 14px; border-radius: 16px;
            border: 1px solid rgba(240,180,41,0.3);
            background: transparent; color: #c8ccd0;
            font-size: 13px; cursor: pointer; min-height: 36px;
            transition: all 0.15s ease;
        }
        .time-chip:active, .time-chip.active {
            background: #f0b429; color: #1a1a2e; border-color: #f0b429;
        }
        .custom-range-picker {
            display: none; gap: 6px; align-items: center; margin-bottom: 10px;
            padding: 8px; background: rgba(255,255,255,0.03); border-radius: 8px;
        }
        .custom-range-picker.open { display: flex; }
        .custom-range-picker input[type="date"] {
            padding: 4px 8px; border: 1px solid #555; border-radius: 6px;
            background: #1a1a2e; color: #c8ccd0; font-size: 13px; min-height: 36px;
        }
        .custom-range-picker button {
            padding: 6px 12px; border-radius: 6px; border: 1px solid #f0b429;
            background: transparent; color: #f0b429; cursor: pointer; min-height: 36px;
        }
        .zoom-reset-btn {
            padding: 4px 10px; border-radius: 6px; border: 1px solid #555;
            background: transparent; color: #aaa; cursor: pointer;
            font-size: 12px; margin-left: auto;
        }
        .zoom-reset-btn:active { background: #333; color: #fff; }
        .metric-chart-header {
            display: flex; align-items: center; gap: 8px; margin-bottom: 8px;
        }

        /* Patch 91: Neue Metrik-Toggles (Pill-Style) */
        .metric-toggle-pills { display: flex; gap: 6px; flex-wrap: wrap; margin-top: 10px; }
        .metric-toggle-pill {
            display: inline-flex; align-items: center; gap: 5px;
            padding: 5px 10px; border-radius: 14px;
            border: 1px solid rgba(255,255,255,0.15);
            background: transparent; color: #888;
            font-size: 12px; cursor: pointer; min-height: 34px;
            transition: all 0.15s ease;
        }
        .metric-toggle-pill.active { color: #c8ccd0; }
        .metric-toggle-pill:not(.active) { opacity: 0.4; }
        .toggle-dot {
            width: 8px; height: 8px; border-radius: 50%; display: inline-block;
        }
        .toggle-info {
            font-size: 14px; opacity: 0.5; padding: 2px 4px; cursor: help;
        }
        .toggle-info:active { opacity: 1; }

        /* Patch 91: Tabellen-Fix */
        #messagesTable {
            table-layout: fixed; width: 100%;
            font-size: var(--hel-font-size-base, 15px);
        }
        #messagesTable td {
            max-width: 200px; overflow: hidden;
            text-overflow: ellipsis; white-space: nowrap;
        }
        .metric-table-scroll {
            overflow-x: auto; -webkit-overflow-scrolling: touch;
        }
        .metric-details { margin-top: 10px; }
        .metric-details summary {
            cursor: pointer; color: #f0b429;
            padding: 8px 0; font-size: 14px;
            list-style: none;
        }
        .metric-details summary::before { content: '▸ '; }
        .metric-details[open] summary::before { content: '▾ '; }

        .chart-container-p91 {
            position: relative; height: 280px; width: 100%;
        }
        .slider-row { display: flex; align-items: center; gap: 15px; }
        .slider-row input { flex: 1; }
        .value-display { min-width: 50px; text-align: center; background: #252525; padding: 8px; border-radius: 8px; }
        .nav-links { display: flex; flex-wrap: wrap; gap: 12px; margin-top: 10px; }
        .nav-link {
            display: inline-block;
            background: #1a3a5c;
            color: #ffd700;
            border: 1px solid #ffd700;
            padding: 12px 20px;
            border-radius: 8px;
            text-decoration: none;
            font-size: 15px;
            font-weight: bold;
            transition: 0.2s;
        }
        .nav-link:hover, .nav-link:active { background: #ffd700; color: #1a1a1a; }
        .metric-toggles { display: flex; gap: 15px; flex-wrap: wrap; margin-bottom: 10px; }
        .metric-toggle label { display: flex; align-items: center; gap: 6px; cursor: pointer; color: #e0e0e0; }
        /* Patch 205 (Phase 5a Schuld aus P199): RAG-Status-Toast nach Datei-Upload.
           Mobile-first 44px Touch-Target zum Dismiss, fixed bottom-right, fade-in
           via .visible-Klasse (kein DOM-Wegnehmen, damit der Renderer leicht bleibt
           und CSS-Transitions sauber spielen). */
        .rag-toast {
            position: fixed;
            bottom: 20px;
            right: 20px;
            max-width: calc(100vw - 40px);
            min-height: 44px;
            padding: 12px 16px;
            background: #2d2d2d;
            border: 1px solid #c8941f;
            border-radius: 8px;
            color: #e0e0e0;
            font-size: 14px;
            line-height: 1.4;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.5);
            z-index: 1000;
            cursor: pointer;
            opacity: 0;
            transform: translateY(10px);
            transition: opacity 0.2s ease-out, transform 0.2s ease-out;
            pointer-events: none;
            display: flex;
            align-items: center;
            -webkit-tap-highlight-color: transparent;
        }
        .rag-toast.visible {
            opacity: 1;
            transform: translateY(0);
            pointer-events: auto;
        }
        .rag-toast.success { border-color: #4ecdc4; }
        .rag-toast.warn { border-color: #ff6b6b; }
        /* Patch 211 (Phase 5a #11): GPU-Queue-Wartetoast. Erscheint links unten,
           damit er nicht mit dem RAG-Toast (rechts) kollidiert. Format identisch
           — fixed, fade-in, klick-dismiss, 44px Touch-Target. Nur sichtbar wenn
           tatsaechlich Konsumenten warten. */
        .gpu-toast {
            position: fixed;
            bottom: 20px;
            left: 20px;
            max-width: calc(100vw - 40px);
            min-height: 44px;
            padding: 12px 16px;
            background: #2d2d2d;
            border: 1px solid #c8941f;
            border-radius: 8px;
            color: #e0e0e0;
            font-size: 14px;
            line-height: 1.4;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.5);
            z-index: 1000;
            cursor: pointer;
            opacity: 0;
            transform: translateY(10px);
            transition: opacity 0.2s ease-out, transform 0.2s ease-out;
            pointer-events: none;
            display: flex;
            align-items: center;
            -webkit-tap-highlight-color: transparent;
        }
        .gpu-toast.visible {
            opacity: 1;
            transform: translateY(0);
            pointer-events: auto;
        }
    </style>
</head>
<body>
    <div class="container">
        <!-- Patch 149 (B-025): Hel-Header mit eigenem Zahnrad-Settings-Button. -->
        <div class="hel-header" style="display:flex;align-items:center;justify-content:space-between;gap:10px;">
            <h1 style="margin:0;">&#9889; Hel – Admin-Konsole</h1>
            <button type="button" id="helSettingsBtn" onclick="toggleHelSettings()"
                    style="background:transparent;border:1px solid #444;color:#ccc;font-size:1.4em;padding:6px 12px;border-radius:8px;cursor:pointer;min-height:44px;min-width:44px;"
                    title="Einstellungen">⚙️</button>
        </div>
        <!-- Hel Mini-Settings (B-025): Schrift-Slider, gleiche Logik wie Nala -->
        <div id="helSettingsPanel" style="display:none;padding:14px;margin:10px 0;background:rgba(255,255,255,0.04);border:1px solid #333;border-radius:8px;">
            <h3 style="margin:0 0 10px;color:#FFD700;">⚙️ Hel-Einstellungen</h3>
            <div style="display:flex;align-items:center;gap:10px;">
                <label style="flex:0 0 120px;color:#ccc;">UI-Skalierung</label>
                <input type="range" id="helUiScaleSlider" min="0.8" max="1.4" step="0.05" value="1.0"
                       style="flex:1;min-height:30px;" oninput="applyHelUiScale(this.value)">
                <span id="helUiScaleVal" style="flex:0 0 60px;text-align:right;font-family:monospace;color:#888;">1.00×</span>
            </div>
            <button onclick="resetHelUiScale()" style="margin-top:8px;padding:6px 14px;background:#333;color:#ccc;border:1px solid #555;border-radius:6px;cursor:pointer;">↺ Zurücksetzen</button>
        </div>
        <!-- Patch 90 (N-F09b): Schriftgr&#246;&#223;en-Wahl (altes UI als Fallback) -->
        <div class="font-preset-bar" role="group" aria-label="Schriftgr&#246;&#223;e" style="display:none;">
            <span class="label">Schrift:</span>
            <button type="button" class="font-preset-btn" data-size="13px" onclick="setFontSize('13px')">13</button>
            <button type="button" class="font-preset-btn" data-size="15px" onclick="setFontSize('15px')">15</button>
            <button type="button" class="font-preset-btn" data-size="17px" onclick="setFontSize('17px')">17</button>
            <button type="button" class="font-preset-btn" data-size="19px" onclick="setFontSize('19px')">19</button>
        </div>
        <!-- Patch 99 (H-F01): Sticky Tab-Leiste - ersetzt Akkordeon -->
        <nav class="hel-tab-nav" id="helTabNav" role="tablist" aria-label="Hel-Sektionen">
            <button type="button" class="hel-tab active" data-tab="metrics" onclick="activateTab('metrics')">&#128202; Metriken</button>
            <button type="button" class="hel-tab" data-tab="llm" onclick="activateTab('llm')">&#129302; LLM</button>
            <button type="button" class="hel-tab" data-tab="system" onclick="activateTab('system')">&#128172; Prompt</button>
            <button type="button" class="hel-tab" data-tab="gedaechtnis" onclick="activateTab('gedaechtnis')">&#128218; RAG</button>
            <button type="button" class="hel-tab" data-tab="cleaner" onclick="activateTab('cleaner')">&#128295; Cleaner</button>
            <button type="button" class="hel-tab" data-tab="usermgmt" onclick="activateTab('usermgmt')">&#128101; User</button>
            <button type="button" class="hel-tab" data-tab="tests" onclick="activateTab('tests')">&#129514; Tests</button>
            <button type="button" class="hel-tab" data-tab="dialect" onclick="activateTab('dialect')">&#128483;&#65039; Dialekte</button>
            <button type="button" class="hel-tab" data-tab="sysctl" onclick="activateTab('sysctl')">&#128147; System</button>
            <button type="button" class="hel-tab" data-tab="provider" onclick="activateTab('provider')">&#10060; Provider</button>
            <button type="button" class="hel-tab" data-tab="huginn" onclick="activateTab('huginn')">&#128020; Huginn</button>
            <button type="button" class="hel-tab" data-tab="projects" onclick="activateTab('projects')">&#128193; Projekte</button>
            <button type="button" class="hel-tab" data-tab="nav" onclick="activateTab('nav')">&#128279; Links</button>
            <button type="button" class="hel-tab" data-tab="about" onclick="activateTab('about')">&#8505;&#65039; About</button>
        </nav>

        <!-- Metriken (offen) -->
        <div class="hel-section active" data-tab="metrics" id="section-metrics">
          <div class="hel-section-header" onclick="toggleSection('metrics')">
            <span class="section-arrow">&#9660;</span> &#128202; Metriken
          </div>
          <div class="hel-section-body" id="body-metrics">
            <div class="card">
                <h2>Metrik-Verlauf</h2>
                <!-- Patch 95: Per-User-Filter -->
                <div class="metric-profile-filter">
                  <label for="profileSelect">Profil:</label>
                  <select id="profileSelect" class="profile-select">
                    <option value="">Alle Profile</option>
                  </select>
                </div>
                <!-- Patch 91: Zeitraum-Chips -->
                <div class="metric-timerange">
                  <button class="time-chip active" data-days="7">7 Tage</button>
                  <button class="time-chip" data-days="30">30 Tage</button>
                  <button class="time-chip" data-days="90">90 Tage</button>
                  <button class="time-chip" data-days="0">Alles</button>
                  <button class="time-chip" data-days="-1" id="customRangeBtn">Custom</button>
                </div>
                <div class="custom-range-picker" id="customRangePicker">
                  <input type="date" id="metricFrom">
                  <span>&ndash;</span>
                  <input type="date" id="metricTo">
                  <button onclick="applyCustomRange()">Anwenden</button>
                </div>
                <div class="metric-chart-header">
                  <span id="metricCountLabel" style="color:#888;font-size:12px;"></span>
                  <button class="zoom-reset-btn" onclick="if(metricsChart) metricsChart.resetZoom()">Zoom Reset</button>
                </div>
                <div class="chart-container-p91">
                    <canvas id="metricsCanvas"></canvas>
                </div>
                <div class="metric-toggle-pills" id="metricToggles"></div>

                <details class="metric-details">
                  <summary>Datentabelle anzeigen</summary>
                  <div class="metric-table-scroll">
                    <table id="messagesTable">
                      <thead><tr><th>Zeit</th><th>Rolle</th><th>Inhalt</th><th>W&#246;rter</th><th>Sentiment</th><th>Kosten (USD)</th></tr></thead>
                      <tbody></tbody>
                    </table>
                  </div>
                </details>
            </div>
            <div class="card">
                <h2>Gespeicherte Sessions</h2>
                <div id="sessionList"></div>
            </div>
          </div>
        </div>

        <!-- Patch 96: Testreports (H-F04) -->
        <div class="hel-section" data-tab="tests" id="section-tests">
          <div class="hel-section-header" onclick="toggleSection('tests')">
            <span class="section-arrow">&#9660;</span> &#129514; Testreports
          </div>
          <div class="hel-section-body collapsed" id="body-tests">
            <div class="card">
                <h2>Playwright-Reports</h2>
                <p style="color:#c8ccd0; font-size:13px; margin-top:0;">
                    Loki (E2E) &amp; Fenrir (Chaos) — generiert via
                    <code>pytest --html=zerberus/tests/report/full_report.html --self-contained-html</code>.
                </p>
                <button class="font-preset-btn" onclick="window.open('/hel/tests/report', '_blank')">
                    Letzten Report &#246;ffnen
                </button>
                <div id="reportsList" style="margin-top:14px;"></div>
            </div>
          </div>
        </div>

        <!-- LLM & Guthaben -->
        <div class="hel-section" data-tab="llm" id="section-llm">
          <div class="hel-section-header" onclick="toggleSection('llm')">
            <span class="section-arrow">&#9654;</span> LLM &amp; Guthaben
          </div>
          <div class="hel-section-body collapsed" id="body-llm" style="max-height:0;padding:0 20px;">
            <div class="card">
                <h2>OpenRouter Modell & Guthaben</h2>
                <label>Modell auswählen:</label>
                <div style="display:flex;gap:8px;margin-bottom:6px;">
                    <button onclick="sortModels('name')" id="sortByName" style="padding:6px 14px;font-size:13px;background:#1a3a5c;color:#ffd700;border:1px solid #ffd700;border-radius:6px;cursor:pointer;">Nach Name (A&#8211;Z)</button>
                    <button onclick="sortModels('price')" id="sortByPrice" style="padding:6px 14px;font-size:13px;background:#252525;color:#aaa;border:1px solid #555;border-radius:6px;cursor:pointer;">Nach Preis (&#8593;)</button>
                </div>
                <select id="modelSelect" onchange="changeModel()"></select>
                <div class="balance" id="balanceDisplay">Guthaben wird geladen...</div>
            </div>
            <div class="card">
                <h2>LLM-Parameter</h2>
                <label>Temperatur (0–2):</label>
                <div class="slider-row">
                    <button onclick="stepSlider('temperature','tempValue',-0.1,0,2)" style="padding:4px 10px;font-size:16px;background:#333;color:#fff;border:1px solid #555;border-radius:6px;cursor:pointer;min-width:34px;">&#9660;</button>
                    <input type="range" id="temperature" min="0" max="2" step="0.1" value="0.7" oninput="document.getElementById('tempValue').innerText=parseFloat(this.value).toFixed(1)">
                    <button onclick="stepSlider('temperature','tempValue',0.1,0,2)" style="padding:4px 10px;font-size:16px;background:#333;color:#fff;border:1px solid #555;border-radius:6px;cursor:pointer;min-width:34px;">&#9650;</button>
                    <span class="value-display" id="tempValue">0.7</span>
                </div>
                <label>Threshold (W&#246;rter f&#252;r Cloud):</label>
                <div class="slider-row">
                    <button onclick="stepSlider('threshold','threshValue',-0.1,0,100)" style="padding:4px 10px;font-size:16px;background:#333;color:#fff;border:1px solid #555;border-radius:6px;cursor:pointer;min-width:34px;">&#9660;</button>
                    <input type="range" id="threshold" min="1" max="100" step="1" value="10" oninput="document.getElementById('threshValue').innerText=parseFloat(this.value).toFixed(1)">
                    <button onclick="stepSlider('threshold','threshValue',0.1,0,100)" style="padding:4px 10px;font-size:16px;background:#333;color:#fff;border:1px solid #555;border-radius:6px;cursor:pointer;min-width:34px;">&#9650;</button>
                    <span class="value-display" id="threshValue">10</span>
                </div>
                <button onclick="saveLLMConfig()">Einstellungen speichern</button>
                <div id="llmStatus"></div>
            </div>
            <!-- Patch 131: Vision-Modell-Dropdown (nur Vision-fähige Modelle) -->
            <div class="card">
                <h2>&#128065;&#65039; Vision-Modell</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:10px;">Modell f&#252;r Bild-Analyse (Huginn-Photos, k&#252;nftig Nala Bild-Upload). DeepSeek V3.2 hat keinen Vision-Support &#8212; daher separate Wahl.</p>
                <label>Vision-Modell ausw&#228;hlen:</label>
                <select id="visionModelSelect" onchange="markVisionDirty()" style="width:100%;padding:8px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;margin-bottom:10px;"></select>
                <label><input type="checkbox" id="visionEnabled" onchange="markVisionDirty()"> Vision aktiviert</label>
                <label style="display:block;margin-top:10px;">Max. Bildgr&#246;&#223;e (MB):</label>
                <input type="number" id="visionMaxMb" min="1" max="50" value="10" onchange="markVisionDirty()" style="padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;width:80px;">
                <div style="display:flex;gap:10px;margin-top:12px;">
                    <button type="button" onclick="visionSave()" style="padding:10px 18px;min-height:44px;background:#DAA520;color:#111;border:none;border-radius:8px;cursor:pointer;font-weight:bold;">&#128190; Speichern</button>
                    <button type="button" onclick="visionReload()" style="padding:10px 18px;min-height:44px;background:#333;color:#eee;border:1px solid #555;border-radius:8px;cursor:pointer;">&#128260; Neu laden</button>
                </div>
                <div id="visionStatus" style="margin-top:10px;color:#8f8;min-height:1.4em;"></div>
            </div>
          </div>
        </div>

        <!-- Whisper-Cleaner -->
        <div class="hel-section" data-tab="cleaner" id="section-cleaner">
          <div class="hel-section-header" onclick="toggleSection('cleaner')">
            <span class="section-arrow">&#9654;</span> Whisper-Cleaner
          </div>
          <div class="hel-section-body collapsed" id="body-cleaner" style="max-height:0;padding:0 20px;">
            <!-- Patch 149 (B-021): WhisperCleaner-Regeln wurden aus dem UI entfernt.
                 Pflege läuft ausschließlich über whisper_cleaner.json auf dem Server
                 (Regex-Patterns sind zu gefährlich ohne Review).
                 Fuzzy-Dictionary bleibt, weil das nur Stringlisten sind. -->
            <div class="card">
                <h2>Whisper-Cleaner Regeln</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:10px;">
                    Die Regex-Regeln werden nicht mehr in der UI gepflegt — siehe
                    <code>whisper_cleaner.json</code> im Projekt-Root. Änderungen dort
                    werden beim nächsten Server-Neustart geladen.
                </p>
                <div id="cleanerStatus" style="margin-top:10px; min-height:1.4em;color:#888;">
                    UI deaktiviert (Patch 149 / B-021). Pflege nur noch via Config-Datei.
                </div>
            </div>
            <div class="card">
                <h2>Fuzzy-Dictionary bearbeiten</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:10px;">Projektspezifische Begriffe f&#252;r Whisper-Fehlerkorrektur via Fuzzy-Matching (<code>fuzzy_dictionary.json</code>). JSON-Array von Strings.</p>
                <textarea id="fuzzyDictEditor" rows="12" style="font-family: monospace;"></textarea>
                <button onclick="saveFuzzyDict()">Speichern</button>
                <div id="fuzzyDictStatus"></div>
            </div>
          </div>
        </div>

        <!-- Dialekte -->
        <div class="hel-section" data-tab="dialect" id="section-dialect">
          <div class="hel-section-header" onclick="toggleSection('dialect')">
            <span class="section-arrow">&#9654;</span> Dialekte
          </div>
          <div class="hel-section-body collapsed" id="body-dialect" style="max-height:0;padding:0 20px;">
            <div class="card">
                <h2>Dialekt-Mappings</h2>
                <!-- Patch 148 (B-022): Strukturiertes UI statt roher JSON-Blob.
                     Pro Dialekt-Gruppe (berlin/schwäbisch/…) eine Sektion mit
                     "Von → Nach"-Eingaben. Suche filtert live. Neue Einträge oben. -->
                <div style="display:flex;gap:8px;margin-bottom:10px;align-items:center;">
                    <input type="search" id="dialectSearch" placeholder="🔍 Suchen (Von oder Nach)…"
                           oninput="renderDialectGroups()"
                           style="flex:1;padding:8px 12px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                    <button onclick="saveDialectStructured()">💾 Alles speichern</button>
                </div>
                <div id="dialectGroups"></div>
                <div style="margin-top:12px;">
                    <input type="text" id="newGroupName" placeholder="Neue Gruppe (z.B. 'bayern')"
                           style="padding:6px 10px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                    <button onclick="addDialectGroup()">+ Gruppe hinzufügen</button>
                </div>
                <div id="dialectStatus" style="margin-top:10px;"></div>
                <!-- Fallback: Raw-JSON (aufklappbar), falls jemand's doch braucht -->
                <details style="margin-top:16px;">
                    <summary style="cursor:pointer;color:#888;font-size:0.85em;">Raw JSON (nur für Notfall-Bearbeitung)</summary>
                    <textarea id="dialectEditor" rows="10" style="font-family: monospace;width:100%;margin-top:8px;"></textarea>
                    <button onclick="saveDialect()" style="margin-top:4px;">Raw-JSON speichern</button>
                </details>
            </div>
          </div>
        </div>

        <!-- System-Prompt -->
        <div class="hel-section" data-tab="system" id="section-system">
          <div class="hel-section-header" onclick="toggleSection('system')">
            <span class="section-arrow">&#9654;</span> System-Prompt
          </div>
          <div class="hel-section-body collapsed" id="body-system" style="max-height:0;padding:0 20px;">
            <div class="card">
                <h2>System-Prompt (f&#252;r alle Chats)</h2>
                <textarea id="systemPromptEditor" rows="8" style="width:100%; font-family: monospace;"></textarea>
                <button onclick="saveSystemPrompt()">Speichern</button>
                <div id="systemPromptStatus"></div>
            </div>
            <div class="card">
                <h2>&#128101; Profil-&#220;bersicht (readonly)</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">Konfigurierte Profile aus <code>config.yaml</code>. Temperatur <em>null</em> = globale Einstellung aus LLM-Parametern.</p>
                <div id="profilesTable" style="overflow-x:auto;">
                    <table>
                        <thead><tr><th>Profil-Key</th><th>Anzeigename</th><th>Berechtigung</th><th>Modell-Override</th><th>Temperatur</th></tr></thead>
                        <tbody id="profilesTableBody"><tr><td colspan="5" style="color:#888;">Wird geladen&#8230;</td></tr></tbody>
                    </table>
                </div>
                <button onclick="loadProfiles()" style="margin-top:12px; background:#333; border:1px solid #555; color:#ccc;">&#8635; Aktualisieren</button>
            </div>
          </div>
        </div>

        <!-- Ged&#228;chtnis / RAG -->
        <div class="hel-section" data-tab="gedaechtnis" id="section-gedaechtnis">
          <div class="hel-section-header" onclick="toggleSection('gedaechtnis')">
            <span class="section-arrow">&#9654;</span> &#129504; Ged&#228;chtnis / RAG
          </div>
          <div class="hel-section-body collapsed" id="body-gedaechtnis" style="max-height:0;padding:0 20px;">
            <div class="card">
                <h2>&#128196; Dokument hochladen</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">Unterst&#252;tzte Formate: <strong>.txt .md .docx .pdf .json .csv</strong>. Chunk-Gr&#246;&#223;e und Split-Strategie h&#228;ngen von der gew&#228;hlten Kategorie ab (Patch 110).</p>
                <label>Datei ausw&#228;hlen:</label>
                <label for="ragFileInput" id="ragFileLabel" style="display:block; padding:18px; background:#252525; border:2px dashed #555; border-radius:10px; text-align:center; cursor:pointer; margin-bottom:14px; font-size:1em; color:#ccc; touch-action:manipulation;">
                    &#128196; Datei tippen/ausw&#228;hlen (.txt .md .docx .pdf .json .csv)
                </label>
                <input type="file" id="ragFileInput" accept=".txt,.md,.docx,.pdf,.json,.csv" style="position:absolute; width:1px; height:1px; opacity:0; overflow:hidden;" onchange="updateRagFileLabel()">
                <label for="ragCategory" style="margin-top:4px;">Kategorie:</label>
                <select id="ragCategory" class="profile-select" style="width:100%; margin-bottom:14px;">
                    <option value="auto" selected>&#128269; Automatisch erkennen</option>
                    <option value="general">Allgemein</option>
                    <option value="narrative">Narrativ / Prosa</option>
                    <option value="technical">Technisch / Code</option>
                    <option value="personal">Pers&#246;nlich / Tagebuch</option>
                    <option value="lore">Lore / Worldbuilding</option>
                    <option value="reference">Referenz / Nachschlagewerk</option>
                    <option value="system">System / Selbstwissen</option>
                </select>
                <button onclick="uploadRagFile()" style="width:100%; padding:16px; font-size:1.1em; border-radius:12px; touch-action:manipulation;">&#128196; Hochladen &amp; Indizieren</button>
                <div id="ragUploadStatus" style="margin-top:14px; font-size:1.05em; word-break:break-word;"></div>
            </div>
            <div class="card">
                <h2>&#129504; Index-&#220;bersicht</h2>
                <div id="ragIndexInfo" style="background:#252525; padding:15px; border-radius:8px; font-family:monospace; margin-bottom:15px;">Wird geladen...</div>
                <h3 style="color:#ffa5a5; margin-bottom:8px;">Indizierte Dokumente</h3>
                <div id="ragSourcesList" style="background:#252525; padding:15px; border-radius:8px; min-height:40px;"></div>
                <br>
                <button type="button" onclick="clearRagIndex()" style="background:#888;">&#128465; Index komplett leeren</button>
                <div id="ragClearStatus" style="margin-top:8px;"></div>
            </div>
            <!-- Patch 115: Background Memory Extraction -->
            <div class="card">
                <h2>&#129504; Automatische Ged&#228;chtnis-Extraktion</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">
                    L&#228;uft automatisch jede Nacht um 04:30 als Teil des Overnight-Jobs.
                    Liest die User-Nachrichten der letzten 24 h, extrahiert konkrete Fakten
                    via Cloud-LLM und schreibt neue Erkenntnisse ins RAG.
                    Duplikate werden per Similarity-Check ausgesiebt.
                </p>
                <button type="button" id="memoryExtractBtn" onclick="triggerMemoryExtraction()" style="width:100%; padding:14px; font-size:1em; border-radius:10px; touch-action:manipulation;">
                    &#129504; Extraktion jetzt starten
                </button>
                <div id="memoryExtractStatus" style="margin-top:12px; font-size:1em; word-break:break-word;"></div>
            </div>

            <!-- Patch 152 (B-020): Memory-Dashboard ─ extrahierte Fakten sichtbar + editierbar. -->
            <div class="card">
                <h2>&#128218; Memory-Dashboard</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:10px;">
                    Alle von der Overnight-Extraktion gespeicherten Fakten. Suche/Filter,
                    Confidence-Badges, manuell hinzufügen, einzeln löschen.
                </p>
                <!-- Statistik -->
                <div id="memoryStats" style="font-size:0.88em;color:#8aa0c0;margin-bottom:10px;"></div>
                <!-- Such- und Filter-Leiste -->
                <div style="display:flex;gap:8px;margin-bottom:10px;flex-wrap:wrap;">
                    <input type="search" id="memorySearch" placeholder="🔍 Suchen in Subjekt/Fakt…"
                           oninput="renderMemoryTable()"
                           style="flex:1;min-width:200px;padding:6px 10px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                    <select id="memoryCategoryFilter" onchange="renderMemoryTable()"
                            style="padding:6px 10px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                        <option value="">Alle Kategorien</option>
                        <option value="PERSON">PERSON</option>
                        <option value="PREFERENCE">PREFERENCE</option>
                        <option value="FACT">FACT</option>
                        <option value="EVENT">EVENT</option>
                        <option value="SKILL">SKILL</option>
                        <option value="EMOTION">EMOTION</option>
                    </select>
                    <button onclick="loadMemoryDashboard()">🔄 Neu laden</button>
                </div>
                <!-- Manuell hinzufügen -->
                <details style="margin-bottom:10px;">
                    <summary style="cursor:pointer;color:#FFD700;">+ Fakt manuell hinzufügen</summary>
                    <div style="display:flex;gap:6px;margin-top:8px;flex-wrap:wrap;">
                        <select id="newMemoryCategory" style="padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:4px;">
                            <option value="FACT">FACT</option>
                            <option value="PERSON">PERSON</option>
                            <option value="PREFERENCE">PREFERENCE</option>
                            <option value="EVENT">EVENT</option>
                            <option value="SKILL">SKILL</option>
                            <option value="EMOTION">EMOTION</option>
                        </select>
                        <input type="text" id="newMemorySubject" placeholder="Subjekt"
                               style="flex:1;min-width:120px;padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:4px;">
                        <input type="text" id="newMemoryFact" placeholder="Fakt"
                               style="flex:2;min-width:200px;padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:4px;">
                        <button onclick="addMemoryManual()">+ Hinzufügen</button>
                    </div>
                    <div id="newMemoryStatus" style="margin-top:6px;font-size:0.85em;"></div>
                </details>
                <!-- Tabelle -->
                <div id="memoryTableHost"></div>
            </div>
          </div>
        </div>

        <!-- Systemsteuerung -->
        <div class="hel-section" data-tab="sysctl" id="section-sysctl">
          <div class="hel-section-header" onclick="toggleSection('sysctl')">
            <span class="section-arrow">&#9654;</span> &#128147; Systemsteuerung
          </div>
          <div class="hel-section-body collapsed" id="body-sysctl" style="max-height:0;padding:0 20px;">
            <div class="card">
                <h2>&#128147; Pacemaker-Konfiguration</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">Der Pacemaker h&#228;lt den Whisper-Dienst aktiv, indem er regelm&#228;&#223;ig stille WAV-Pings sendet.</p>
                <label>Pacemaker-Laufzeit (Minuten):</label>
                <input type="number" id="pacemakerMinutes" min="1" max="480" value="120" style="width:150px;">
                <p style="color:#888; font-size:0.85em; margin-top:6px;">&#8505;&#65039; &#196;nderungen wirken erst nach Neustart des Servers.</p>
                <button onclick="savePacemakerConfig()">Speichern</button>
                <div id="pacemakerStatus" style="margin-top:8px;"></div>
            </div>
            <div class="card">
                <h2>&#128260; Whisper Docker-Restart (Patch 119)</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">St&#252;ndlicher Auto-Restart des Whisper-Containers l&#228;uft im Hintergrund. Manueller Trigger bei Bedarf:</p>
                <button type="button" onclick="restartWhisperContainer()" style="min-height:44px;">&#128260; Whisper neu starten</button>
                <div id="whisperRestartStatus" style="margin-top:10px; font-size:0.95em;"></div>
            </div>

            <!-- Patch 150 (B-024): Pacemaker-Prozess-Steuerung -->
            <div class="card">
                <h2>&#128147; Pacemaker-Prozesse</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">
                    Granulare Kontrolle über Hintergrund-Prozesse: Intervall, Device (CPU/GPU), Master-Schalter.
                </p>
                <!-- Master -->
                <div class="pacemaker-master" style="display:flex;gap:16px;align-items:center;margin-bottom:12px;padding:10px;background:rgba(255,255,255,0.03);border-radius:6px;">
                    <label style="flex:1;color:#ccc;">
                        <input type="checkbox" id="pacemaker-master" onchange="savePacemakerProcesses()"> 🫀 Pacemaker Master
                    </label>
                    <label style="flex:1;color:#ccc;">
                        <input type="checkbox" id="pacemaker-sync" onchange="savePacemakerProcesses()"> Intervalle synchronisieren
                    </label>
                </div>
                <!-- Process-Liste (von JS befüllt) -->
                <div id="pacemakerProcesses"></div>
                <!-- Activity-Anzeige -->
                <div id="pacemakerActivity" style="margin-top:12px;font-size:0.85em;color:#8aa0c0;"></div>
                <button onclick="savePacemakerProcesses()" style="margin-top:12px;">💾 Prozess-Einstellungen speichern</button>
                <div id="pacemakerProcessesStatus" style="margin-top:8px;"></div>
            </div>
          </div>
        </div>

        <!-- Provider -->
        <div class="hel-section" data-tab="provider" id="section-provider">
          <div class="hel-section-header" onclick="toggleSection('provider')">
            <span class="section-arrow">&#9654;</span> &#10060; Provider-Blacklist
          </div>
          <div class="hel-section-body collapsed" id="body-provider" style="max-height:0;padding:0 20px;">
            <div class="card">
                <h2>OpenRouter Provider-Blacklist</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">
                    Blockierte Provider werden bei jedem Request via <code>provider.ignore</code> ausgeschlossen.
                    <br><span style="color:#ffd700;">&#9888;&#65039; chutes und targon sind standardm&#228;&#223;ig gesperrt</span>
                    (Rate-Limiting bekannt).
                </p>
                <div id="providerList" style="display:flex;flex-wrap:wrap;gap:6px;margin-bottom:14px;"></div>
                <div style="display:flex;flex-wrap:wrap;gap:8px;margin-bottom:10px;align-items:center;">
                    <select id="newProviderSelect" class="zb-select" onchange="onProviderSelectChange()" style="flex:1;min-width:200px;"></select>
                    <input type="text" id="newProviderInput" placeholder="Provider-Name (z.B. deepinfra)" style="flex:1;min-width:180px;display:none;">
                    <button onclick="addProvider()" style="padding:10px 18px;white-space:nowrap;">+ Hinzuf&#252;gen</button>
                </div>
                <button onclick="saveProviderBlacklist()">&#128190; Speichern</button>
                <div id="providerStatus" style="margin-top:8px;color:#4ecdc4;"></div>
            </div>
          </div>
        </div>

        <!-- User-Verwaltung -->
        <div class="hel-section" data-tab="usermgmt" id="section-usermgmt">
          <div class="hel-section-header" onclick="toggleSection('usermgmt')">
            <span class="section-arrow">&#9654;</span> &#128274; User-Verwaltung
          </div>
          <div class="hel-section-body collapsed" id="body-usermgmt" style="max-height:0;padding:0 20px;">
            <div class="card">
                <h2>&#128274; Hash-Tester</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">Pr&#252;ft ob ein Passwort zum gespeicherten Hash passt.</p>
                <label>Profil</label>
                <select id="hashTestProfile"></select>
                <label>Passwort</label>
                <input type="text" id="hashTestPassword" placeholder="Passwort eingeben">
                <button onclick="testHash()" style="margin-top:12px;">&#128269; Testen</button>
                <div id="hashTestResult" style="margin-top:12px; font-size:1.1em; font-weight:bold;"></div>
            </div>
            <div class="card">
                <h2>&#128260; Passwort zur&#252;cksetzen</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">Setzt ein neues Passwort f&#252;r ein Profil (mind. 4 Zeichen).</p>
                <label>Profil</label>
                <select id="resetProfile"></select>
                <label>Neues Passwort</label>
                <input type="text" id="resetPassword" placeholder="Neues Passwort (mind. 4 Zeichen)">
                <button onclick="resetProfilePassword()" style="margin-top:12px;">&#128190; Passwort setzen</button>
                <div id="resetResult" style="margin-top:12px; font-size:1.1em; font-weight:bold;"></div>
            </div>
          </div>
        </div>

        <!-- Patch 127: Huginn (Telegram-Bot) Konfiguration -->
        <div class="hel-section" data-tab="huginn" id="section-huginn">
          <div class="hel-section-header" onclick="toggleSection('huginn')">
            <span class="section-arrow">&#9654;</span> &#128020; Huginn (Telegram-Bot)
          </div>
          <div class="hel-section-body" id="body-huginn">
            <div class="card">
              <h2>&#128020; Huginn &mdash; Bullauge in Zerberus</h2>
              <p style="color:#aaa; font-size:0.92em; margin-top:-6px;">
                Patch 123: Vollwertiger Chat-Partner auf Telegram. Guard (Mistral Small 3) prueft jede Antwort, HitL fuer Code-Ausfuehrung und Gruppenbeitritt.
              </p>
              <div id="huginn-status" style="margin:12px 0; padding:10px; border-radius:8px; background:#1f1f1f;">
                <span id="huginn-status-dot" style="display:inline-block;width:10px;height:10px;border-radius:50%;background:#888;margin-right:6px;"></span>
                <span id="huginn-status-text">Lade...</span>
              </div>

              <div style="display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-top:10px;">
                <label style="display:flex;flex-direction:column;gap:4px;">
                  <span style="color:#DAA520;font-size:0.92em;">Aktiviert</span>
                  <select id="huginn-enabled" style="padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                    <option value="false">Deaktiviert</option>
                    <option value="true">Aktiviert</option>
                  </select>
                </label>
                <label style="display:flex;flex-direction:column;gap:4px;">
                  <span style="color:#DAA520;font-size:0.92em;">Admin Chat-ID</span>
                  <input type="text" id="huginn-admin-chat-id" placeholder="z.B. 123456789" style="padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                </label>
                <label style="display:flex;flex-direction:column;gap:4px;grid-column:span 2;">
                  <span style="color:#DAA520;font-size:0.92em;">Bot-Token (nur setzen wenn neu/geaendert)</span>
                  <input type="password" id="huginn-bot-token" placeholder="wird maskiert angezeigt" style="padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;font-family:monospace;">
                  <span id="huginn-bot-token-masked" style="font-family:monospace;color:#888;font-size:0.88em;margin-top:2px;"></span>
                </label>
                <label style="display:flex;flex-direction:column;gap:4px;">
                  <span style="color:#DAA520;font-size:0.92em;">Modell (OpenRouter)</span>
                  <select id="huginn-model" style="padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                    <option value="deepseek/deepseek-chat">deepseek/deepseek-chat (Default)</option>
                  </select>
                  <span id="huginn-model-price" style="color:#888;font-size:0.85em;"></span>
                </label>
                <label style="display:flex;flex-direction:column;gap:4px;">
                  <span style="color:#DAA520;font-size:0.92em;">Max Response-Laenge</span>
                  <input type="number" id="huginn-max-length" value="4000" min="500" max="4096" style="padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                </label>
                <label style="display:flex;flex-direction:column;gap:4px;grid-column:span 2;">
                  <span style="color:#DAA520;font-size:0.92em;">System-Prompt (Persona) &mdash; Patch 158</span>
                  <textarea id="huginn-system-prompt" rows="12"
                    style="width:100%;font-family:monospace;font-size:13px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;padding:10px;resize:vertical;"></textarea>
                  <small style="color:#888;">Definiert Huginns Charakter und Tonfall. Wird als System-Prompt an das LLM geschickt. Leer lassen &rarr; kein System-Prompt.</small>
                  <button type="button" onclick="huginnResetPrompt()" style="margin-top:4px;padding:6px 12px;font-size:12px;background:#333;color:#eee;border:1px solid #555;border-radius:6px;cursor:pointer;align-self:flex-start;">
                    &#x21a9;&#xfe0f; Standard-Persona wiederherstellen
                  </button>
                </label>
              </div>

              <details style="margin-top:18px;" id="hg-allowlist-details">
                <summary style="cursor:pointer;color:#DAA520;">&#128274; Zugriffskontrolle (Allowlist) &mdash; Patch 181</summary>
                <div style="display:grid;grid-template-columns:1fr;gap:10px;margin-top:10px;padding-left:14px;">
                  <label style="display:flex;flex-direction:column;gap:4px;">
                    Modus:
                    <select id="huginn-allowlist-mode" onchange="huginnAllowlistModeChange()" style="padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                      <option value="open">Offen (alle d&uuml;rfen) &mdash; Default</option>
                      <option value="allowlist">Allowlist (nur freigeschaltete User)</option>
                      <option value="admin_only">Nur Admin</option>
                    </select>
                  </label>
                  <label id="huginn-allowed-users-group" style="display:none;flex-direction:column;gap:4px;">
                    Erlaubte User-IDs (kommasepariert):
                    <input type="text" id="huginn-allowed-users" placeholder="123456789, 987654321" style="padding:6px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;font-family:monospace;">
                    <small style="color:#888;">Telegram-User-IDs. Admin ist immer erlaubt. Leere Liste = alle erlaubt (Safety-Fallback).</small>
                  </label>
                </div>
              </details>

              <details style="margin-top:18px;">
                <summary style="cursor:pointer;color:#DAA520;">&#128172; Gruppen-Verhalten</summary>
                <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-top:10px;padding-left:14px;">
                  <label><input type="checkbox" id="hg-resp-name" checked> Reagiert auf "Huginn" im Text</label>
                  <label><input type="checkbox" id="hg-resp-mention" checked> Reagiert auf @-Mention</label>
                  <label><input type="checkbox" id="hg-resp-reply" checked> Reagiert auf Replies auf seine Messages</label>
                  <label><input type="checkbox" id="hg-autonomous" checked> Autonome Einwuerfe (mit LLM-Validation)</label>
                  <label style="grid-column:span 2;">
                    Cooldown autonomous (Sekunden):
                    <input type="number" id="hg-cooldown" value="300" min="30" max="3600" style="width:100px;padding:4px;background:#121212;color:#eee;border:1px solid #444;border-radius:4px;">
                  </label>
                </div>
              </details>

              <details style="margin-top:12px;">
                <summary style="cursor:pointer;color:#DAA520;">&#128274; HitL (Human-in-the-Loop)</summary>
                <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-top:10px;padding-left:14px;">
                  <label><input type="checkbox" id="hitl-code" checked> HitL vor Code-Ausfuehrung</label>
                  <label><input type="checkbox" id="hitl-group" checked> HitL vor Gruppenbeitritt</label>
                  <label style="grid-column:span 2;">
                    HitL-Timeout (Sekunden):
                    <input type="number" id="hitl-timeout" value="300" min="30" max="3600" style="width:100px;padding:4px;background:#121212;color:#eee;border:1px solid #444;border-radius:4px;">
                  </label>
                </div>
              </details>

              <div style="margin-top:18px;display:flex;gap:10px;flex-wrap:wrap;">
                <button type="button" onclick="huginnSave()" style="padding:10px 18px;min-height:44px;background:#DAA520;color:#111;border:none;border-radius:8px;cursor:pointer;font-weight:bold;">
                  &#128190; Speichern
                </button>
                <button type="button" onclick="huginnReload()" style="padding:10px 18px;min-height:44px;background:#333;color:#eee;border:1px solid #555;border-radius:8px;cursor:pointer;">
                  &#8635; Neu laden
                </button>
              </div>
              <div id="huginn-save-status" style="margin-top:10px;color:#8f8;min-height:1.4em;"></div>
            </div>
          </div>
        </div>

        <!-- Patch 100: About / Meilenstein -->
        <div class="hel-section" data-tab="about" id="section-about">
          <div class="hel-section-header" onclick="toggleSection('about')">
            <span class="section-arrow">&#9654;</span> &#8505;&#65039; About
          </div>
          <div class="hel-section-body" id="body-about">
            <div class="card" style="text-align:center;">
                <img src="/static/pics/Architekt_und_Sonnenblume.png" alt="Architekt und Sonnenblume"
                     style="max-width:100%; max-height:60vh; border-radius:12px; box-shadow:0 4px 20px rgba(0,0,0,0.6); margin-bottom:18px;">
                <h2 style="color:#DAA520; margin-top:0;">&#127995; Patch 100 &ndash; Zerberus Pro 4.0 &#127995;</h2>
                <p style="font-style:italic; color:#ffd700; margin:6px 0 16px;">&bdquo;Das Gebrochene sichtbar machen."</p>
                <p style="line-height:1.6;">
                    Kein Code wurde von Hand geschrieben.<br>
                    Jede Zeile entstand im Dialog zwischen<br>
                    Architekt und Maschine.<br><br>
                    Von Patch 1 bis Patch 100.
                </p>
                <p style="color:#DAA520; line-height:1.8;">
                    &#128021;&#8205;&#129466; Zerberus &middot; &#128049; Nala &middot; &#128081; Hel &middot; &#127801; Rosa<br>
                    &#129418; Loki &middot; &#128058; Fenrir &middot; &#128063;&#65039; Ratatoskr &middot; &#127752; Heimdall
                </p>
                <hr style="border:none; border-top:1px solid #c8941f; margin:22px auto; max-width:60%;">
                <div style="text-align:left; font-family:monospace; color:#c8ccd0; line-height:1.7; display:inline-block;">
                    Version: Patch 100<br>
                    Architektur: FastAPI + SQLite + FAISS + Whisper + Ollama<br>
                    Tests: 34/34 (Loki &amp; Fenrir)<br>
                    RAG: 10/11
                </div>
                <p style="margin-top:18px; color:#aaa; font-size:0.9em;">
                    Entwickelt von Chris mit Claude (Supervisor + Claude Code)<br>
                    F&uuml;r Jojo und Nala &#128049;
                </p>
            </div>
          </div>
        </div>

        <!-- Navigation -->
        <div class="hel-section" data-tab="nav" id="section-nav">
          <div class="hel-section-header" onclick="toggleSection('nav')">
            <span class="section-arrow">&#9654;</span> &#128279; Navigation
          </div>
          <div class="hel-section-body collapsed" id="body-nav" style="max-height:0;padding:0 20px;">
            <div class="card">
                <h2>&#128279; Schnell-Navigation</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:14px;">Alle Links &#246;ffnen in einem neuen Tab.</p>
                <div class="nav-links">
                    <a href="/nala" target="_blank" class="nav-link">&#128172; Nala-Chat</a>
                    <a href="/hel" target="_blank" class="nav-link">&#9889; Hel-Dashboard</a>
                    <a href="/health" target="_blank" class="nav-link">&#10084;&#65039; API Health</a>
                    <a href="/archive/sessions" target="_blank" class="nav-link">&#128196; Session-Archiv</a>
                    <a href="/hel/metrics/latest_with_costs" target="_blank" class="nav-link">&#128202; Metrics Latest</a>
                    <a href="/hel/debug/state" target="_blank" class="nav-link">&#128736;&#65039; Debug State</a>
                    <a href="/docs" target="_blank" class="nav-link">&#128216; API Docs (Swagger)</a>
                </div>
            </div>
          </div>
        </div>

        <!-- Patch 195 (Phase 5a #1): Projekte-Verwaltung -->
        <div class="hel-section" data-tab="projects" id="section-projects">
          <div class="hel-section-header" onclick="toggleSection('projects')">
            <span class="section-arrow">&#9654;</span> &#128193; Projekte
          </div>
          <div class="hel-section-body" id="body-projects">
            <div class="card">
                <h2>&#128193; Projekte</h2>
                <p style="color:#aaa; font-size:0.9em; margin-bottom:12px;">
                    Projekt-Verwaltung (Phase 5a #1). Backend P194: Tabellen + CRUD-Endpoints.
                    Persona-Overlay als Merge-Layer (System &rarr; User &rarr; Projekt).
                </p>
                <div style="display:flex;flex-wrap:wrap;gap:10px;align-items:center;margin-bottom:14px;">
                    <button onclick="openProjectForm()" style="padding:10px 18px;min-height:44px;background:#1a3a5c;color:#ffd700;border:1px solid #ffd700;border-radius:8px;">+ Projekt anlegen</button>
                    <label style="display:inline-flex;align-items:center;gap:6px;color:#aaa;font-size:0.9em;">
                        <input type="checkbox" id="projectsShowArchived" onchange="loadProjects()"> Archivierte anzeigen
                    </label>
                    <button onclick="loadProjects()" style="padding:8px 14px;min-height:44px;">&#128260; Reload</button>
                </div>
                <div id="projectsStatus" style="margin-bottom:10px;color:#4ecdc4;font-size:0.9em;"></div>
                <div style="overflow-x:auto;">
                    <table id="projectsTable" style="width:100%;border-collapse:collapse;">
                        <thead>
                            <tr style="background:#1f1f1f;color:#ffd700;">
                                <th style="text-align:left;padding:8px;">Slug</th>
                                <th style="text-align:left;padding:8px;">Name</th>
                                <th style="text-align:left;padding:8px;">Updated</th>
                                <th style="text-align:left;padding:8px;">Status</th>
                                <th style="text-align:left;padding:8px;">Aktionen</th>
                            </tr>
                        </thead>
                        <tbody id="projectsTableBody"></tbody>
                    </table>
                </div>
            </div>

            <!-- Detail-Panel (zeigt Dateien des aktuell gewaehlten Projekts) -->
            <div class="card" id="projectDetailCard" style="display:none;">
                <h2>&#128196; Dateien <span id="projectDetailName" style="color:#aaa;font-weight:normal;font-size:0.85em;"></span></h2>
                <!-- Patch 196: Drop-Zone + File-Picker -->
                <div id="projectDropZone" data-project-id=""
                     style="border:2px dashed #4ecdc4;border-radius:10px;padding:20px;margin-bottom:12px;text-align:center;color:#aaa;cursor:pointer;transition:background 0.15s;">
                    <div style="font-size:1.5em;color:#4ecdc4;">&#11014;&#65039;</div>
                    <div style="margin-top:6px;">Dateien hierher ziehen oder <span style="color:#ffd700;text-decoration:underline;">klicken zum Auswaehlen</span></div>
                    <div style="margin-top:4px;font-size:0.78em;color:#666;">Max. 50&nbsp;MB pro Datei. Blockiert: .exe, .bat, .sh, ...</div>
                    <input type="file" id="projectFileInput" multiple style="display:none;">
                </div>
                <div id="projectUploadProgress" style="margin-bottom:10px;font-family:monospace;font-size:0.85em;"></div>
                <div id="projectFilesList" style="font-family:monospace;font-size:0.88em;color:#ddd;"></div>
            </div>

            <!-- Form-Overlay (Modal-aehnlich, ohne extra CSS-Lib) -->
            <div id="projectFormOverlay" style="display:none;position:fixed;top:0;left:0;width:100%;height:100%;background:rgba(0,0,0,0.7);z-index:1000;align-items:flex-start;justify-content:center;overflow-y:auto;padding:20px;">
                <div style="background:#2d2d2d;border-radius:12px;padding:20px;max-width:600px;width:100%;margin-top:40px;">
                    <h2 id="projectFormTitle" style="margin-top:0;">Projekt anlegen</h2>
                    <input type="hidden" id="projectFormId">
                    <label style="display:block;margin-top:10px;color:#DAA520;">Name *</label>
                    <input type="text" id="projectFormName" placeholder="z.B. Akte Mutzenbacher" style="width:100%;padding:10px;min-height:44px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                    <label style="display:block;margin-top:10px;color:#DAA520;">Beschreibung</label>
                    <textarea id="projectFormDescription" rows="2" placeholder="Optional, freier Text" style="width:100%;padding:10px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;"></textarea>
                    <label style="display:block;margin-top:10px;color:#DAA520;">Slug-Override <span style="color:#888;font-weight:normal;">(optional, nur bei Anlegen)</span></label>
                    <input type="text" id="projectFormSlug" placeholder="ableitung-aus-name" style="width:100%;padding:10px;min-height:44px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                    <fieldset style="margin-top:14px;border:1px solid #444;border-radius:8px;padding:12px;">
                        <legend style="color:#DAA520;padding:0 6px;">Persona-Overlay</legend>
                        <label style="display:block;color:#aaa;font-size:0.88em;">System-Addendum</label>
                        <textarea id="projectFormSystemAddendum" rows="3" placeholder="Zusaetzlicher System-Prompt-Text (Projekt-Kontext)" style="width:100%;padding:8px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;font-family:monospace;font-size:0.88em;"></textarea>
                        <label style="display:block;margin-top:8px;color:#aaa;font-size:0.88em;">Tone-Hints <span style="color:#888;">(Komma-Liste)</span></label>
                        <input type="text" id="projectFormToneHints" placeholder="z.B. praezise, fachsprachlich, knapp" style="width:100%;padding:8px;min-height:44px;background:#121212;color:#eee;border:1px solid #444;border-radius:6px;">
                    </fieldset>
                    <div style="margin-top:16px;display:flex;gap:10px;flex-wrap:wrap;">
                        <button onclick="saveProjectForm()" style="padding:10px 18px;min-height:44px;background:#1a5c1a;color:#fff;border:1px solid #4ecdc4;border-radius:8px;">&#128190; Speichern</button>
                        <button onclick="closeProjectForm()" style="padding:10px 18px;min-height:44px;">Abbrechen</button>
                    </div>
                    <div id="projectFormStatus" style="margin-top:10px;color:#f88;"></div>
                </div>
            </div>
          </div>
        </div>

    </div>

    <script>
        let currentModel = '';
        let chart;
        let _allModels = [];
        let _currentSort = 'name';

        function sortModels(mode) {
            _currentSort = mode;
            document.getElementById('sortByName').style.background = mode === 'name' ? '#1a3a5c' : '#252525';
            document.getElementById('sortByName').style.color = mode === 'name' ? '#ffd700' : '#aaa';
            document.getElementById('sortByName').style.borderColor = mode === 'name' ? '#ffd700' : '#555';
            document.getElementById('sortByPrice').style.background = mode === 'price' ? '#1a3a5c' : '#252525';
            document.getElementById('sortByPrice').style.color = mode === 'price' ? '#ffd700' : '#aaa';
            document.getElementById('sortByPrice').style.borderColor = mode === 'price' ? '#ffd700' : '#555';
            renderModelSelect(_allModels, currentModel);
        }

        function stepSlider(sliderId, displayId, delta, minVal, maxVal) {
            const slider = document.getElementById(sliderId);
            const display = document.getElementById(displayId);
            let val = Math.round((parseFloat(slider.value) + delta) * 10) / 10;
            if (val < minVal) val = minVal;
            if (val > maxVal) val = maxVal;
            slider.value = val;
            display.innerText = val.toFixed(1);
        }

        // Patch 147 (B-019): Einheitlicher Formatter für ALLE Modell-Dropdowns.
        // Format: "Name — $input/$output/1M" — keine Budget/Premium-Labels mehr.
        // Sortierung standardmäßig nach Input-Preis aufsteigend (günstigste oben).
        function formatModelLabel(name, inputPrice, outputPrice) {
            const disp = name || '?';
            const inNum = Number(inputPrice) || 0;
            const outNum = Number(outputPrice) || 0;
            if (inNum === 0 && outNum === 0) return `${disp} — kostenlos`;
            return `${disp} — $${inNum.toFixed(2)}/$${outNum.toFixed(2)}/1M`;
        }

        function renderModelSelect(models, selectedModel) {
            const sorted = [...models].sort((a, b) => {
                if (_currentSort === 'price') {
                    return parseFloat(a.pricing?.prompt || 0) - parseFloat(b.pricing?.prompt || 0);
                }
                // Patch 147: Default-Sortierung ebenfalls nach Preis aufsteigend.
                return parseFloat(a.pricing?.prompt || 0) - parseFloat(b.pricing?.prompt || 0);
            });
            const select = document.getElementById('modelSelect');
            select.innerHTML = '';
            sorted.forEach(m => {
                const option = document.createElement('option');
                option.value = m.id;
                const inPrice = parseFloat(m.pricing?.prompt || 0) * 1_000_000;
                const outPrice = parseFloat(m.pricing?.completion || 0) * 1_000_000;
                option.textContent = formatModelLabel(m.name || m.id, inPrice, outPrice);
                select.appendChild(option);
            });
            if (selectedModel && select.querySelector(`option[value="${selectedModel}"]`)) {
                select.value = selectedModel;
            }
        }
        let _chartHistory = [];
        let _prevBalance = (() => { try { const v = localStorage.getItem('hel_prevBalance'); return v !== null ? parseFloat(v) : null; } catch(_) { return null; } })();

        // Patch 99 (H-F01): Sticky Tab-Navigation — ersetzt Akkordeon.
        // Jede .hel-section trägt data-tab; genau eine ist .active (display:block),
        // alle anderen sind via CSS versteckt. Lazy-Loads analog altem toggleSection.
        const _HEL_LAZY_LOADED = new Set();
        function activateTab(id) {
            document.querySelectorAll('.hel-section[data-tab]').forEach(s => {
                s.classList.toggle('active', s.dataset.tab === id);
            });
            document.querySelectorAll('.hel-tab').forEach(btn => {
                btn.classList.toggle('active', btn.dataset.tab === id);
            });
            try { localStorage.setItem('hel_active_tab', id); } catch (_) {}
            // Lazy-Load genau einmal pro Sektion
            if (!_HEL_LAZY_LOADED.has(id)) {
                _HEL_LAZY_LOADED.add(id);
                if (id === 'metrics') loadMetrics();
                if (id === 'system') { loadSystemPrompt(); loadProfiles(); }
                if (id === 'gedaechtnis') { loadRagStatus(); loadMemoryDashboard(); }
                if (id === 'sysctl') { loadPacemakerConfig(); loadPacemakerProcesses(); }
                if (id === 'provider') loadProviderBlacklist();
                if (id === 'huginn') huginnReload();
                if (id === 'llm') visionReload();
                if (id === 'projects') loadProjects();
            }
            // Aktiven Tab in die Mitte scrollen, falls Overflow
            const activeBtn = document.querySelector('.hel-tab.active');
            if (activeBtn && activeBtn.scrollIntoView) {
                activeBtn.scrollIntoView({ inline: 'center', block: 'nearest', behavior: 'smooth' });
            }
        }
        // Rückwärtskompat-Alias, falls Altcode noch toggleSection aufruft.
        function toggleSection(id) { activateTab(id); }

        // --- Patch 131: Vision-Modell-Konfiguration ---
        let _visionModels = [];
        let _visionDirty = false;
        function markVisionDirty() { _visionDirty = true; }
        function _visionTierBadge(tier) {
            const map = { budget: '#2f6f2f', mid: '#6f5f2f', premium: '#6f2f4f' };
            const color = map[tier] || '#444';
            return `<span style="background:${color};color:#fff;padding:1px 6px;border-radius:4px;font-size:0.72em;margin-right:6px;">${tier||'?'}</span>`;
        }
        async function visionReload() {
            const statusEl = document.getElementById('visionStatus');
            if (!statusEl) return;
            statusEl.textContent = 'Lade Vision-Modelle...';
            try {
                const [mResp, cResp] = await Promise.all([
                    fetch('/hel/admin/vision/models'),
                    fetch('/hel/admin/vision/config'),
                ]);
                const mData = await mResp.json();
                const cData = await cResp.json();
                _visionModels = mData.models || [];
                // Patch 147 (B-019): Sortierung nach Input-Preis, einheitlicher Formatter,
                // kein [Budget]/[Premium]-Label mehr.
                _visionModels.sort((a, b) => (Number(a.input_price) || 0) - (Number(b.input_price) || 0));
                const sel = document.getElementById('visionModelSelect');
                sel.innerHTML = _visionModels.map(m => {
                    const label = formatModelLabel(m.name, m.input_price, m.output_price);
                    return `<option value="${m.id}" data-tier="${m.tier}">${label}</option>`;
                }).join('');
                sel.value = cData.model || 'qwen/qwen2.5-vl-7b-instruct';
                document.getElementById('visionEnabled').checked = !!cData.enabled;
                document.getElementById('visionMaxMb').value = cData.max_image_size_mb || 10;
                _visionDirty = false;
                statusEl.textContent = `${_visionModels.length} Modelle geladen.`;
            } catch (e) {
                statusEl.textContent = 'Fehler: ' + e;
                statusEl.style.color = '#f88';
            }
        }
        async function visionSave() {
            const statusEl = document.getElementById('visionStatus');
            statusEl.style.color = '#8f8';
            const payload = {
                model: document.getElementById('visionModelSelect').value,
                enabled: document.getElementById('visionEnabled').checked,
                max_image_size_mb: parseInt(document.getElementById('visionMaxMb').value, 10) || 10,
            };
            try {
                const r = await fetch('/hel/admin/vision/config', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(payload),
                });
                if (r.ok) {
                    statusEl.textContent = 'Gespeichert.';
                    _visionDirty = false;
                } else {
                    const err = await r.text();
                    statusEl.textContent = 'Fehler: ' + err;
                    statusEl.style.color = '#f88';
                }
            } catch (e) {
                statusEl.textContent = 'Fehler: ' + e;
                statusEl.style.color = '#f88';
            }
        }

        // --- Patch 127: Huginn (Telegram) Tab ---
        // Patch 158: Default-Persona aus letztem GET-Response cachen, damit
        // der "Standard wiederherstellen"-Button keinen zweiten Request braucht.
        let _huginnDefaultPrompt = '';

        function huginnResetPrompt() {
            const el = document.getElementById('huginn-system-prompt');
            if (!el) return;
            if (!_huginnDefaultPrompt) {
                alert('Standard-Persona noch nicht geladen. Bitte "Neu laden" druecken.');
                return;
            }
            el.value = _huginnDefaultPrompt;
        }

        // Patch 181: User-IDs-Feld nur sichtbar im Allowlist-Modus.
        function huginnAllowlistModeChange() {
            const sel = document.getElementById('huginn-allowlist-mode');
            const grp = document.getElementById('huginn-allowed-users-group');
            if (!sel || !grp) return;
            grp.style.display = (sel.value === 'allowlist') ? 'flex' : 'none';
        }

        async function huginnReload() {
            const statusEl = document.getElementById('huginn-status-text');
            const dotEl = document.getElementById('huginn-status-dot');
            const saveStatus = document.getElementById('huginn-save-status');
            if (saveStatus) saveStatus.textContent = '';
            try {
                const r = await fetch('/hel/admin/huginn/config');
                if (!r.ok) throw new Error('HTTP ' + r.status);
                const cfg = await r.json();
                document.getElementById('huginn-enabled').value = String(cfg.enabled);
                document.getElementById('huginn-admin-chat-id').value = cfg.admin_chat_id || '';
                document.getElementById('huginn-bot-token-masked').textContent =
                    cfg.bot_token_masked ? ('aktueller Token: ' + cfg.bot_token_masked) : '(kein Token gesetzt)';
                document.getElementById('huginn-bot-token').value = '';
                document.getElementById('huginn-max-length').value = cfg.max_response_length || 4000;

                // Patch 158: System-Prompt (Persona) befuellen + Default cachen.
                _huginnDefaultPrompt = cfg.default_system_prompt || '';
                const promptEl = document.getElementById('huginn-system-prompt');
                if (promptEl) {
                    promptEl.value = (cfg.system_prompt != null) ? cfg.system_prompt : '';
                }

                // Modell-Dropdown mit _allModels fuellen wenn verfuegbar.
                // Patch 147 (B-019): Einheitlicher Formatter, Sortierung nach Input-Preis.
                const modelSel = document.getElementById('huginn-model');
                const current = cfg.model || 'deepseek/deepseek-chat';
                if (Array.isArray(_allModels) && _allModels.length > 0) {
                    const sorted = [..._allModels].sort((a, b) =>
                        parseFloat(a.pricing?.prompt || 0) - parseFloat(b.pricing?.prompt || 0)
                    );
                    modelSel.innerHTML = '';
                    for (const m of sorted) {
                        const opt = document.createElement('option');
                        opt.value = m.id;
                        const inP = parseFloat(m.pricing?.prompt || 0) * 1_000_000;
                        const outP = parseFloat(m.pricing?.completion || 0) * 1_000_000;
                        opt.textContent = formatModelLabel(m.name || m.id, inP, outP);
                        if (m.id === current) opt.selected = true;
                        modelSel.appendChild(opt);
                    }
                } else {
                    modelSel.innerHTML = `<option value="${current}" selected>${current}</option>`;
                }

                const gb = cfg.group_behavior || {};
                document.getElementById('hg-resp-name').checked = gb.respond_to_name !== false;
                document.getElementById('hg-resp-mention').checked = gb.respond_to_mention !== false;
                document.getElementById('hg-resp-reply').checked = gb.respond_to_direct_reply !== false;
                document.getElementById('hg-autonomous').checked = gb.autonomous_interjection !== false;
                document.getElementById('hg-cooldown').value = gb.interjection_cooldown_seconds || 300;

                const hitl = cfg.hitl || {};
                document.getElementById('hitl-code').checked = hitl.code_execution !== false;
                document.getElementById('hitl-group').checked = hitl.group_join !== false;
                document.getElementById('hitl-timeout').value = hitl.confirmation_timeout_seconds || 300;

                // Patch 181: Allowlist-Mode + erlaubte User-IDs.
                const modeSel = document.getElementById('huginn-allowlist-mode');
                if (modeSel) {
                    modeSel.value = cfg.allowlist_mode || 'open';
                }
                const usersInput = document.getElementById('huginn-allowed-users');
                if (usersInput) {
                    const arr = Array.isArray(cfg.allowed_users) ? cfg.allowed_users : [];
                    usersInput.value = arr.join(', ');
                }
                huginnAllowlistModeChange();

                if (cfg.enabled) {
                    dotEl.style.background = '#3fb84e';
                    statusEl.textContent = 'Aktiv';
                } else {
                    dotEl.style.background = '#888';
                    statusEl.textContent = 'Deaktiviert';
                }
            } catch (e) {
                statusEl.textContent = 'Fehler: ' + e.message;
                dotEl.style.background = '#e74c3c';
            }
        }

        async function huginnSave() {
            const statusEl = document.getElementById('huginn-save-status');
            statusEl.style.color = '#8f8';
            statusEl.textContent = 'Speichere...';
            // Patch 158: system_prompt aus der Textarea mitnehmen — leerer String
            // bedeutet "keine Persona" und wird bewusst durchgereicht.
            const promptEl = document.getElementById('huginn-system-prompt');
            const payload = {
                enabled: document.getElementById('huginn-enabled').value === 'true',
                admin_chat_id: document.getElementById('huginn-admin-chat-id').value.trim(),
                model: document.getElementById('huginn-model').value,
                max_response_length: parseInt(document.getElementById('huginn-max-length').value, 10) || 4000,
                system_prompt: promptEl ? promptEl.value : '',
                group_behavior: {
                    respond_to_name: document.getElementById('hg-resp-name').checked,
                    respond_to_mention: document.getElementById('hg-resp-mention').checked,
                    respond_to_direct_reply: document.getElementById('hg-resp-reply').checked,
                    autonomous_interjection: document.getElementById('hg-autonomous').checked,
                    interjection_cooldown_seconds: parseInt(document.getElementById('hg-cooldown').value, 10) || 300,
                    interjection_trigger: 'smart'
                },
                hitl: {
                    code_execution: document.getElementById('hitl-code').checked,
                    group_join: document.getElementById('hitl-group').checked,
                    confirmation_timeout_seconds: parseInt(document.getElementById('hitl-timeout').value, 10) || 300
                },
                // Patch 181: Allowlist-Mode + erlaubte User-IDs.
                allowlist_mode: (document.getElementById('huginn-allowlist-mode') || {}).value || 'open',
                allowed_users: (document.getElementById('huginn-allowed-users') || {value:''}).value
                    .split(',').map(s => s.trim()).filter(s => s.length > 0)
            };
            const tokenVal = document.getElementById('huginn-bot-token').value.trim();
            if (tokenVal && !tokenVal.includes('\u2026')) {
                payload.bot_token = tokenVal;
            }
            try {
                const r = await fetch('/hel/admin/huginn/config', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(payload)
                });
                if (!r.ok) throw new Error('HTTP ' + r.status);
                statusEl.textContent = '\u2705 Gespeichert. Server-Neustart fuer Webhook-Wechsel erforderlich.';
                await huginnReload();
            } catch (e) {
                statusEl.style.color = '#f88';
                statusEl.textContent = '\u274c Fehler: ' + e.message;
            }
        }

        // Patch 156: Webhook-Registrierungs-Funktion entfernt — Long-Polling
        // ist Default, Webhook-Setup erfolgt nicht mehr ueber das Hel-UI.

        async function loadModelsAndBalance() {
            const balanceEl = document.getElementById('balanceDisplay');
            try {
                const [modelsRes, balanceRes] = await Promise.all([
                    fetch('/hel/admin/models'),
                    fetch('/hel/admin/balance')
                ]);

                if (!modelsRes.ok) {
                    const err = await modelsRes.json().catch(() => ({}));
                    balanceEl.innerText = 'OpenRouter nicht erreichbar: ' + (err.detail || modelsRes.status);
                    return;
                }
                if (!balanceRes.ok) {
                    const err = await balanceRes.json().catch(() => ({}));
                    balanceEl.innerText = 'OpenRouter nicht erreichbar: ' + (err.detail || balanceRes.status);
                    return;
                }

                const models = await modelsRes.json();
                const balance = await balanceRes.json();

                _allModels = Array.isArray(models) ? models : [];

                const configRes = await fetch('/hel/admin/config');
                const config = await configRes.json();
                currentModel = config.llm?.cloud_model || '';
                renderModelSelect(_allModels, currentModel);

                // Patch 136: Kostenanzeige FIX — zeigt tatsächliche Beträge in EUR statt
                // sinnloser "pro 1M Tokens"-Umrechnung.
                const curBalance = balance.balance != null ? parseFloat(balance.balance) : null;
                const balanceEur = balance.balance_eur != null ? parseFloat(balance.balance_eur) : null;
                const balanceStr = balanceEur != null
                    ? `Kontostand: ${balanceEur.toFixed(2).replace('.', ',')} &euro; <span style="color:#888;font-size:0.85em;">($${curBalance.toFixed(2)})</span>`
                    : 'Kontostand: nicht verf&uuml;gbar';
                const lastCostEur = parseFloat(balance.last_cost_eur || 0);
                const todayEur = parseFloat(balance.today_total_eur || 0);
                const lastLine = `Letzte Anfrage: ${lastCostEur.toFixed(4).replace('.', ',')} &euro;`;
                const todayLine = `Heute gesamt: ${todayEur.toFixed(4).replace('.', ',')} &euro;`;
                balanceEl.innerHTML =
                    `${balanceStr}<br>` +
                    `<span style="font-size:0.9em;color:#ccc;">${lastLine}</span><br>` +
                    `<span style="font-size:0.85em;color:#aaa;">${todayLine}</span>`;

                // Temperatur und Threshold aus config laden
                const tempSlider = document.getElementById('temperature');
                const threshSlider = document.getElementById('threshold');
                tempSlider.value = config.llm?.temperature ?? 0.7;
                threshSlider.value = config.llm?.threshold ?? 10;
                document.getElementById('tempValue').innerText = parseFloat(tempSlider.value).toFixed(1);
                document.getElementById('threshValue').innerText = parseFloat(threshSlider.value).toFixed(1);
            } catch (e) {
                balanceEl.innerText = 'Fehler beim Laden: ' + e.message;
            }
        }

        async function changeModel() {
            const model = document.getElementById('modelSelect').value;
            await fetch('/hel/admin/config', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({ llm: { cloud_model: model } })
            });
            alert('Modell gespeichert');
        }

        async function saveLLMConfig() {
            const temp = parseFloat(document.getElementById('temperature').value);
            const thresh = parseInt(document.getElementById('threshold').value);
            const config = { llm: { temperature: temp, threshold: thresh } };
            const res = await fetch('/hel/admin/config', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify(config)
            });
            if (res.ok) document.getElementById('llmStatus').innerText = '✅ Einstellungen gespeichert';
            else document.getElementById('llmStatus').innerText = '❌ Fehler';
        }

        // ── Patch 90 (H-F02): Whisper-Cleaner als UX-Formular ──
        let _cleanerEntries = []; // Array of {kind:'rule'|'comment', pattern, replacement, comment}

        function _escapeHtml(s) {
            return String(s == null ? '' : s)
                .replace(/&/g, '&amp;').replace(/</g, '&lt;')
                .replace(/>/g, '&gt;').replace(/"/g, '&quot;');
        }

        // Validiert ein Python-Regex grob in JS: strippt fuehrendes (?i)/(?s)/(?m)
        // und versucht new RegExp(). Echte Python-Spezialitaeten (z.B. \1 in pattern)
        // werden nicht erkannt — nur grobe Syntax-Pruefung.
        function _validatePattern(p) {
            if (!p) return 'Pattern darf nicht leer sein';
            let stripped = p;
            let flags = '';
            const m = stripped.match(/^\(\?([imsx]+)\)/);
            if (m) {
                stripped = stripped.slice(m[0].length);
                if (m[1].indexOf('i') >= 0) flags += 'i';
                if (m[1].indexOf('s') >= 0) flags += 's';
                if (m[1].indexOf('m') >= 0) flags += 'm';
            }
            try { new RegExp(stripped, flags); return ''; }
            catch (e) { return 'Ungueltiges Regex: ' + e.message; }
        }

        function renderCleanerList() {
            const host = document.getElementById('cleanerList');
            // Patch 169 (B6): cleanerList-DOM-Element wurde mit Patch 149
            // entfernt (Cleaner-Pflege laeuft jetzt server-seitig ueber
            // whisper_cleaner.json). Falls renderCleanerList trotzdem
            // gerufen wird (z. B. aus loadCleaner beim Boot), ohne Guard
            // crashen wir mit "can't access property innerHTML, host is null".
            if (!host) return;
            const parts = _cleanerEntries.map((e, idx) => {
                if (e.kind === 'comment') {
                    return `
                        <div class="cleaner-section" data-idx="${idx}" data-kind="comment">
                            <span style="color:#c8941f;">&#9776;</span>
                            <input type="text" name="comment" value="${_escapeHtml(e.comment)}"
                                   placeholder="Sektion / Kommentar">
                            <button class="cleaner-trash" onclick="removeCleanerEntry(${idx})" title="Entfernen">&#128465;</button>
                        </div>`;
                }
                const errId = 'err-' + idx;
                return `
                    <div class="cleaner-card" data-idx="${idx}" data-kind="rule">
                        <div class="cleaner-card-row">
                            <div class="cleaner-card-grow">
                                <label>Pattern (Python-Regex)</label>
                                <input type="text" name="pattern" value="${_escapeHtml(e.pattern)}"
                                       oninput="_revalidateRow(${idx})" autocapitalize="off"
                                       autocorrect="off" spellcheck="false">
                                <div class="cleaner-error" id="${errId}"></div>
                            </div>
                            <div class="cleaner-card-grow">
                                <label>Replacement</label>
                                <input type="text" name="replacement" value="${_escapeHtml(e.replacement)}"
                                       autocapitalize="off" autocorrect="off" spellcheck="false">
                            </div>
                            <button class="cleaner-trash" onclick="removeCleanerEntry(${idx})"
                                    title="Regel l&#246;schen">&#128465;</button>
                        </div>
                        <label style="margin-top:8px;">Kommentar (optional)</label>
                        <input type="text" name="comment" value="${_escapeHtml(e.comment)}"
                               placeholder="Kategorie/Notiz">
                    </div>`;
            });
            host.innerHTML = parts.join('') ||
                '<div style="color:#888; padding:12px;">Keine Eintr&#228;ge. &#187;Regel hinzuf&#252;gen&#171; klicken.</div>';
            // Pattern-Validierung beim ersten Render zeigen
            _cleanerEntries.forEach((e, idx) => { if (e.kind === 'rule') _revalidateRow(idx); });
        }

        function _revalidateRow(idx) {
            const card = document.querySelector('.cleaner-card[data-idx="' + idx + '"]');
            if (!card) return;
            const patEl = card.querySelector('input[name="pattern"]');
            const errEl = card.querySelector('.cleaner-error');
            const msg = _validatePattern(patEl.value);
            errEl.textContent = msg;
            card.classList.toggle('invalid', !!msg);
        }

        function _collectCleanerFromDom() {
            const out = [];
            const errors = [];
            const cards = document.querySelectorAll('#cleanerList > [data-idx]');
            cards.forEach((el) => {
                const idx = parseInt(el.dataset.idx, 10);
                if (el.dataset.kind === 'comment') {
                    const c = el.querySelector('input[name="comment"]').value.trim();
                    if (c) out.push({ _comment: c });
                } else {
                    const pat = el.querySelector('input[name="pattern"]').value;
                    const rep = el.querySelector('input[name="replacement"]').value;
                    const com = el.querySelector('input[name="comment"]').value.trim();
                    const err = _validatePattern(pat);
                    if (err) errors.push('Zeile ' + (idx + 1) + ': ' + err);
                    const obj = { pattern: pat, replacement: rep };
                    if (com) obj._comment = com;
                    out.push(obj);
                }
            });
            return { entries: out, errors };
        }

        function addCleanerRule() {
            _cleanerEntries.push({ kind: 'rule', pattern: '', replacement: '', comment: '' });
            renderCleanerList();
            // Fokus aufs neue Pattern-Feld
            const cards = document.querySelectorAll('#cleanerList .cleaner-card');
            const last = cards[cards.length - 1];
            if (last) last.querySelector('input[name="pattern"]').focus();
        }
        function addCleanerComment() {
            _cleanerEntries.push({ kind: 'comment', comment: '' });
            renderCleanerList();
            const sections = document.querySelectorAll('#cleanerList .cleaner-section');
            const last = sections[sections.length - 1];
            if (last) last.querySelector('input[name="comment"]').focus();
        }
        function removeCleanerEntry(idx) {
            const e = _cleanerEntries[idx];
            if (!e) return;
            const label = e.kind === 'comment'
                ? ('Kommentar "' + (e.comment || '') + '"')
                : ('Regel "' + (e.pattern || '') + '"');
            if (!confirm(label + ' wirklich entfernen?')) return;
            _cleanerEntries.splice(idx, 1);
            renderCleanerList();
        }

        async function loadCleaner() {
            // Patch 169 (B6): Wenn das cleanerList-DOM-Element nicht existiert
            // (Patch-149-Aufraeumung), ist nichts zu rendern — kein Fetch noetig.
            if (!document.getElementById('cleanerList')) return;
            try {
                const res = await fetch('/hel/admin/whisper_cleaner');
                const data = await res.json();
                _cleanerEntries = (Array.isArray(data) ? data : []).map((e) => {
                    if (e && typeof e === 'object' && 'pattern' in e) {
                        return {
                            kind: 'rule',
                            pattern: e.pattern || '',
                            replacement: e.replacement == null ? '' : String(e.replacement),
                            comment: e._comment || '',
                        };
                    }
                    return { kind: 'comment', comment: (e && e._comment) || '' };
                });
                renderCleanerList();
            } catch (err) {
                // Patch 169 (B6): cleanerStatus kann auch fehlen — Null-Guard.
                const statusEl = document.getElementById('cleanerStatus');
                if (statusEl) {
                    statusEl.innerHTML =
                        '<span style="color:#ff6b6b;">Laden fehlgeschlagen: ' + _escapeHtml(err.message) + '</span>';
                }
            }
        }

        async function saveCleaner() {
            const status = document.getElementById('cleanerStatus');
            const { entries, errors } = _collectCleanerFromDom();
            if (errors.length) {
                status.innerHTML = '<span style="color:#ff6b6b;">&#10007; '
                    + errors.length + ' ung&#252;ltige Regex-Pattern. Speichern blockiert.</span>';
                return;
            }
            status.innerHTML = '<span style="color:#aaa;">Speichere&#8230;</span>';
            try {
                const res = await fetch('/hel/admin/whisper_cleaner', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(entries),
                });
                if (res.ok) {
                    status.innerHTML = '<span style="color:#4ecdc4;">&#10003; Gespeichert ('
                        + entries.length + ' Eintr&#228;ge)</span>';
                    // DOM-State zur State-Liste zurueck synchronisieren (Reihenfolge unveraendert)
                    _cleanerEntries = entries.map((e) => ({
                        kind: 'pattern' in e ? 'rule' : 'comment',
                        pattern: e.pattern || '',
                        replacement: e.replacement == null ? '' : String(e.replacement),
                        comment: e._comment || '',
                    }));
                } else {
                    const t = await res.text();
                    status.innerHTML = '<span style="color:#ff6b6b;">&#10007; Fehler: '
                        + _escapeHtml(t.slice(0, 200)) + '</span>';
                }
            } catch (err) {
                status.innerHTML = '<span style="color:#ff6b6b;">&#10007; ' + _escapeHtml(err.message) + '</span>';
            }
        }

        async function loadFuzzyDict() {
            const res = await fetch('/hel/admin/fuzzy_dict');
            const data = await res.json();
            document.getElementById('fuzzyDictEditor').value = JSON.stringify(data, null, 2);
        }
        async function saveFuzzyDict() {
            const raw = document.getElementById('fuzzyDictEditor').value;
            let parsed;
            try { parsed = JSON.parse(raw); } catch (e) { alert('Ungültiges JSON'); return; }
            if (!Array.isArray(parsed)) { alert('Muss ein JSON-Array sein, z.B. ["Zerberus", "FastAPI"]'); return; }
            const res = await fetch('/hel/admin/fuzzy_dict', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify(parsed)
            });
            if (res.ok) document.getElementById('fuzzyDictStatus').innerText = '&#10003; Gespeichert';
            else document.getElementById('fuzzyDictStatus').innerText = '&#10007; Fehler';
        }

        // ── Patch 148 (B-022): Strukturierter Dialekt-Editor ──
        // _dialectData ist Arbeitskopie im Speicher, bis der User "Speichern" klickt.
        let _dialectData = {};

        async function loadDialect() {
            const res = await fetch('/hel/admin/dialect');
            const data = await res.json();
            _dialectData = data || {};
            // Raw-JSON füllen (Fallback-Editor)
            const rawEd = document.getElementById('dialectEditor');
            if (rawEd) rawEd.value = JSON.stringify(_dialectData, null, 2);
            renderDialectGroups();
        }

        function renderDialectGroups() {
            const host = document.getElementById('dialectGroups');
            if (!host) return;
            const q = (document.getElementById('dialectSearch')?.value || '').toLowerCase();
            host.innerHTML = '';
            Object.keys(_dialectData).sort().forEach(group => {
                const entries = _dialectData[group] || {};
                // Filter anwenden
                const filteredKeys = Object.keys(entries).filter(k => {
                    if (!q) return true;
                    return k.toLowerCase().includes(q) || String(entries[k]).toLowerCase().includes(q);
                });
                if (q && filteredKeys.length === 0) return;
                const g = document.createElement('div');
                g.className = 'dialect-group';
                g.style.marginBottom = '20px';
                g.style.padding = '10px';
                g.style.border = '1px solid #333';
                g.style.borderRadius = '6px';
                g.style.background = 'rgba(255,255,255,0.02)';
                const h = document.createElement('h3');
                h.textContent = group;
                h.style.margin = '0 0 10px';
                h.style.color = '#FFD700';
                // Patch 170 (B4): Icon-Button statt großer roter Button
                const delGroup = document.createElement('button');
                delGroup.textContent = '🗑️';
                delGroup.title = 'Gruppe löschen';
                delGroup.setAttribute('aria-label', `Gruppe "${group}" löschen`);
                delGroup.style.cssText = 'float:right;width:28px;height:28px;min-width:0;min-height:0;'
                    + 'padding:0;display:inline-flex;align-items:center;justify-content:center;'
                    + 'background:transparent;border:1px solid #555;border-radius:4px;'
                    + 'color:#ccc;cursor:pointer;opacity:0.5;transition:opacity 0.15s,color 0.15s,border-color 0.15s;';
                delGroup.onmouseenter = () => {
                    delGroup.style.opacity = '1';
                    delGroup.style.color = 'var(--zb-danger, #FF6B6B)';
                    delGroup.style.borderColor = 'var(--zb-danger, #FF6B6B)';
                };
                delGroup.onmouseleave = () => {
                    delGroup.style.opacity = '0.5';
                    delGroup.style.color = '#ccc';
                    delGroup.style.borderColor = '#555';
                };
                delGroup.ontouchstart = () => {
                    delGroup.style.opacity = '1';
                    delGroup.style.color = 'var(--zb-danger, #FF6B6B)';
                    delGroup.style.borderColor = 'var(--zb-danger, #FF6B6B)';
                };
                delGroup.onclick = () => {
                    if (confirm(`Gruppe "${group}" wirklich löschen?`)) {
                        delete _dialectData[group];
                        renderDialectGroups();
                    }
                };
                h.appendChild(delGroup);
                g.appendChild(h);
                // Neuer-Eintrag-Button OBEN (Patch 148 fordert "oben, nicht scrollen müssen")
                const addRow = document.createElement('div');
                addRow.style.cssText = 'display:flex;gap:6px;margin-bottom:10px;align-items:center;';
                const fromIn = document.createElement('input');
                fromIn.type = 'text'; fromIn.placeholder = 'Von…';
                fromIn.style.cssText = 'flex:1;padding:4px 8px;background:#121212;color:#eee;border:1px solid #444;border-radius:4px;';
                const arr = document.createElement('span'); arr.textContent = '→'; arr.style.color = '#FFD700';
                const toIn = document.createElement('input');
                toIn.type = 'text'; toIn.placeholder = 'Nach…';
                toIn.style.cssText = 'flex:1;padding:4px 8px;background:#121212;color:#eee;border:1px solid #444;border-radius:4px;';
                const addBtn = document.createElement('button');
                addBtn.textContent = '+';
                addBtn.style.cssText = 'padding:4px 12px;background:#264;border:1px solid #484;color:#fff;border-radius:4px;cursor:pointer;';
                addBtn.onclick = () => {
                    const k = fromIn.value.trim(), v = toIn.value.trim();
                    if (!k || !v) { alert('Beide Felder ausfüllen'); return; }
                    _dialectData[group][k] = v;
                    renderDialectGroups();
                };
                addRow.appendChild(fromIn); addRow.appendChild(arr); addRow.appendChild(toIn); addRow.appendChild(addBtn);
                g.appendChild(addRow);
                // Bestehende Einträge
                filteredKeys.forEach(k => {
                    const row = document.createElement('div');
                    row.style.cssText = 'display:flex;gap:6px;margin-bottom:4px;align-items:center;';
                    const fi = document.createElement('input');
                    fi.type = 'text'; fi.value = k;
                    fi.style.cssText = 'flex:1;padding:4px 8px;background:#1a1a1a;color:#eee;border:1px solid #333;border-radius:4px;';
                    fi.onchange = () => {
                        const newKey = fi.value.trim();
                        if (!newKey || newKey === k) return;
                        _dialectData[group][newKey] = _dialectData[group][k];
                        delete _dialectData[group][k];
                        renderDialectGroups();
                    };
                    const a2 = document.createElement('span'); a2.textContent = '→'; a2.style.color = '#888';
                    const vi = document.createElement('input');
                    vi.type = 'text'; vi.value = entries[k];
                    vi.style.cssText = 'flex:1;padding:4px 8px;background:#1a1a1a;color:#eee;border:1px solid #333;border-radius:4px;';
                    vi.onchange = () => { _dialectData[group][k] = vi.value; };
                    const del = document.createElement('button');
                    del.className = 'delete-entry';
                    del.textContent = '✕';
                    del.style.cssText = 'padding:4px 10px;background:transparent;border:1px solid #844;color:#e88;border-radius:4px;cursor:pointer;';
                    del.onclick = () => {
                        delete _dialectData[group][k];
                        renderDialectGroups();
                    };
                    row.appendChild(fi); row.appendChild(a2); row.appendChild(vi); row.appendChild(del);
                    g.appendChild(row);
                });
                host.appendChild(g);
            });
        }

        function addDialectGroup() {
            const n = document.getElementById('newGroupName');
            const name = (n.value || '').trim().toLowerCase();
            if (!name) { alert('Gruppenname fehlt'); return; }
            if (_dialectData[name]) { alert('Gruppe existiert bereits'); return; }
            _dialectData[name] = {};
            n.value = '';
            renderDialectGroups();
        }

        async function saveDialectStructured() {
            const res = await fetch('/hel/admin/dialect', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify(_dialectData),
            });
            const st = document.getElementById('dialectStatus');
            if (res.ok) {
                st.innerText = '✅ Gespeichert';
                st.style.color = '#8f8';
                // Raw-JSON-Fallback synchronisieren
                const rawEd = document.getElementById('dialectEditor');
                if (rawEd) rawEd.value = JSON.stringify(_dialectData, null, 2);
            } else {
                st.innerText = '❌ Fehler';
                st.style.color = '#f88';
            }
        }

        // Legacy: Raw-JSON-Editor (Fallback)
        async function saveDialect() {
            const raw = document.getElementById('dialectEditor').value;
            try { _dialectData = JSON.parse(raw); } catch (e) { alert('Ungültiges JSON'); return; }
            const res = await fetch('/hel/admin/dialect', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: raw
            });
            if (res.ok) {
                document.getElementById('dialectStatus').innerText = '✅ Gespeichert';
                renderDialectGroups();
            } else {
                document.getElementById('dialectStatus').innerText = '❌ Fehler';
            }
        }

        async function loadSystemPrompt() {
            const res = await fetch('/hel/admin/system_prompt');
            const data = await res.json();
            document.getElementById('systemPromptEditor').value = data.prompt || '';
        }
        async function saveSystemPrompt() {
            const prompt = document.getElementById('systemPromptEditor').value;
            const res = await fetch('/hel/admin/system_prompt', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({ prompt: prompt })
            });
            const result = await res.json();
            document.getElementById('systemPromptStatus').innerText = result.status === 'ok' ? '✅ Gespeichert' : '❌ Fehler';
        }

        async function deleteSession(sessionId) {
            if (!confirm('Chat wirklich löschen?')) return;
            const res = await fetch(`/hel/admin/session/${sessionId}`, { method: 'DELETE' });
            if (res.ok) {
                loadMetrics();
                document.getElementById('systemPromptStatus').innerText = '✅ Session gelöscht';
            } else {
                document.getElementById('systemPromptStatus').innerText = '❌ Fehler beim Löschen';
            }
        }

        // ========== Patch 91: Chart.js Metriken-Dashboard ==========
        let metricsChart = null;
        let _currentTimeRange = 7;

        const METRIC_DEFS = {
            bert_sentiment: { label: 'BERT Sentiment', color: '#f0b429', info: 'Stimmungsanalyse per German BERT (0=negativ, 0.5=neutral, 1=positiv)' },
            rolling_ttr:    { label: 'TTR (Rolling-50)', color: '#4fc3f7', info: 'Type-Token-Ratio über 50 Nachrichten — misst lexikalische Vielfalt (0–1)' },
            shannon_entropy:{ label: 'Shannon Entropy', color: '#81c784', info: 'Informationsdichte der Wortwahl — höher = vielfältigere Sprache' },
            hapax_ratio:    { label: 'Hapax Ratio', color: '#e57373', info: 'Anteil einmalig verwendeter Wörter — höher = mehr einzigartige Begriffe' },
            avg_word_length:{ label: 'Ø Wortlänge', color: '#ba68c8', info: 'Durchschnittliche Wortlänge in Zeichen — Indikator für Wortkomplexität' }
        };

        function renderMetricToggles() {
            const container = document.getElementById('metricToggles');
            if (!container) return;
            container.innerHTML = Object.entries(METRIC_DEFS).map(([key, def]) => `
                <button class="metric-toggle-pill active" id="toggle_${key}" onclick="toggleMetric('${key}')">
                  <span class="toggle-dot" style="background:${def.color}"></span>
                  ${def.label}
                  <span class="toggle-info" onclick="event.stopPropagation(); showMetricInfo('${key}')" title="${def.info}">&#9432;</span>
                </button>
            `).join('');
        }

        function toggleMetric(key) {
            const btn = document.getElementById('toggle_' + key);
            if (btn) btn.classList.toggle('active');
            rebuildChart();
        }

        function showMetricInfo(key) {
            alert(METRIC_DEFS[key].label + '\\n\\n' + METRIC_DEFS[key].info);
        }

        function buildChart(data) {
            const canvas = document.getElementById('metricsCanvas');
            if (!canvas) return;
            const ctx = canvas.getContext('2d');
            if (metricsChart) { metricsChart.destroy(); metricsChart = null; }

            const labels = data.map(r => {
                const d = new Date(r.created_at || r.timestamp);
                return d.toLocaleDateString('de-DE', {day:'2-digit', month:'2-digit'}) + ' '
                     + d.toLocaleTimeString('de-DE', {hour:'2-digit', minute:'2-digit'});
            });

            const activeMetrics = Object.keys(METRIC_DEFS).filter(k => {
                const toggle = document.getElementById('toggle_' + k);
                return toggle ? toggle.classList.contains('active') : true;
            });

            const datasets = activeMetrics.map(key => ({
                label: METRIC_DEFS[key].label,
                data: data.map(r => r[key] ?? null),
                borderColor: METRIC_DEFS[key].color,
                backgroundColor: 'transparent',
                borderWidth: 1.5,
                pointRadius: 0,
                pointHitRadius: 12,
                tension: 0.3,
                spanGaps: true
            }));

            metricsChart = new Chart(ctx, {
                type: 'line',
                data: { labels, datasets },
                options: {
                    responsive: true,
                    maintainAspectRatio: false,
                    interaction: { mode: 'index', intersect: false },
                    plugins: {
                        legend: { display: false },
                        tooltip: {
                            backgroundColor: 'rgba(26,26,46,0.95)',
                            titleColor: '#f0b429', bodyColor: '#c8ccd0',
                            padding: 10, cornerRadius: 8, displayColors: true,
                            callbacks: {
                                label: function(ctx) {
                                    const v = ctx.parsed.y;
                                    return ctx.dataset.label + ': ' + (v == null ? '–' : v.toFixed(3));
                                }
                            }
                        },
                        zoom: {
                            pan: { enabled: true, mode: 'x' },
                            zoom: {
                                pinch: { enabled: true },
                                wheel: { enabled: true },
                                mode: 'x'
                            }
                        }
                    },
                    scales: {
                        x: {
                            ticks: { color: '#888', maxRotation: 45, maxTicksLimit: 12, font: { size: 11 } },
                            grid: { color: 'rgba(255,255,255,0.05)' }
                        },
                        y: {
                            min: 0,
                            ticks: { color: '#888', font: { size: 11 } },
                            grid: { color: 'rgba(255,255,255,0.08)' }
                        }
                    }
                }
            });
        }

        function rebuildChart() {
            if (!_chartHistory || !_chartHistory.length) { buildChart([]); return; }
            buildChart(_chartHistory);
        }

        async function loadMetricsChart(days) {
            if (days === undefined) days = _currentTimeRange;
            _currentTimeRange = days;

            let url = '/hel/metrics/history?limit=500';
            if (days === -1) {
                const from = document.getElementById('metricFrom').value;
                const to = document.getElementById('metricTo').value;
                if (from) url += '&from_date=' + encodeURIComponent(from);
                if (to) url += '&to_date=' + encodeURIComponent(to);
            } else if (days > 0) {
                const from = new Date(Date.now() - days * 86400000).toISOString().slice(0, 10);
                url += '&from_date=' + from;
            }
            // Patch 95: Per-User-Filter
            const pkSel = document.getElementById('profileSelect');
            const pk = pkSel ? pkSel.value : '';
            if (pk) url += '&profile_key=' + encodeURIComponent(pk);

            const res = await fetch(url);
            if (!res.ok) return;
            const body = await res.json();
            const data = (body && body.results) ? body.results : (Array.isArray(body) ? body : []);
            _chartHistory = data;

            const countLabel = document.getElementById('metricCountLabel');
            if (countLabel && body && body.meta) {
                countLabel.textContent = body.meta.count + ' Einträge';
            }

            buildChart(data);

            document.querySelectorAll('.time-chip').forEach(c => c.classList.remove('active'));
            const activeChip = document.querySelector('.time-chip[data-days="' + days + '"]');
            if (activeChip) activeChip.classList.add('active');
        }

        function applyCustomRange() { loadMetricsChart(-1); }

        async function loadMetrics() {
            // Tabelle (Letzte Nachrichten) — unverändert über latest_with_costs
            const res = await fetch('/hel/metrics/latest_with_costs?limit=20');
            const data = res.ok ? await res.json() : [];

            const tbody = document.querySelector('#messagesTable tbody');
            if (tbody) {
                tbody.innerHTML = '';
                data.forEach((msg) => {
                    const tr = document.createElement('tr');
                    const raw = (msg.content || '');
                    const sentenceEnd = raw.search(/[.!?]/);
                    const firstSentence = sentenceEnd > 0 ? raw.substring(0, sentenceEnd + 1) : raw;
                    const truncated = firstSentence.length > 80 ? firstSentence.substring(0, 80) + '\u2026' : firstSentence;
                    const ts = new Date(msg.timestamp);
                    const timeStr = ts.toLocaleDateString('de-DE') + ' ' + ts.toLocaleTimeString('de-DE', {hour:'2-digit',minute:'2-digit'});
                    tr.innerHTML = '<td>' + timeStr + '</td>' +
                                   '<td>' + msg.role + '</td>' +
                                   '<td title="' + raw.substring(0,200).replace(/"/g,'&quot;') + '">' + truncated + '</td>' +
                                   '<td>' + (msg.word_count ?? '') + '</td>' +
                                   '<td>' + (msg.vader_compound ?? '') + '</td>' +
                                   '<td>' + (msg.cost !== null && msg.cost !== undefined ? (msg.cost * 0.92).toFixed(5).replace('.', ',') + ' &euro;' : '') + '</td>';
                    tbody.appendChild(tr);
                });
            }

            // Chart über neuen Pfad laden
            await loadMetricsChart(_currentTimeRange);

            const sessionsRes = await fetch('/hel/admin/sessions');
            const sessions = await sessionsRes.json();
            const sessionDiv = document.getElementById('sessionList');
            sessionDiv.innerHTML = sessions.map(s => `
                <div class="session-item" style="display: flex; justify-content: space-between; align-items: center;">
                    <span style="cursor: pointer;" onclick="exportSession('${s.session_id}')">${s.first_message || 'Neuer Chat'}</span>
                    <span style="display: flex; gap: 10px;">
                        <button onclick="exportSession('${s.session_id}')" style="background: #666; padding: 5px 10px;">&#128196;</button>
                        <button onclick="deleteSession('${s.session_id}')" style="background: #ff6b6b; padding: 5px 10px;">&#128465;</button>
                    </span>
                    <span class="session-date">${new Date(s.created_at).toLocaleString()}</span>
                </div>
            `).join('');
        }

        // updateChart als Kompatibilitäts-Alias (alte Aufrufe)
        function updateChart() { rebuildChart(); }

        // Patch 96: Testreports-Liste laden (H-F04)
        async function loadReportsList() {
            const el = document.getElementById('reportsList');
            if (!el) return;
            try {
                const r = await fetch('/hel/tests/reports');
                if (!r.ok) { el.textContent = 'Fehler beim Laden (' + r.status + ')'; return; }
                const data = await r.json();
                const reports = data.reports || [];
                if (!reports.length) {
                    el.textContent = 'Keine Reports vorhanden. Bitte pytest ausführen.';
                    return;
                }
                // Patch 170 (B5): Einzel-Reports verlinken (full/fenrir/loki).
                const ALLOWED = new Set(['full_report.html', 'fenrir_report.html', 'loki_report.html']);
                const rows = reports.map(f => {
                    const dt = new Date(f.mtime * 1000).toLocaleString('de-DE');
                    const kb = (f.size / 1024).toFixed(1) + ' KB';
                    let link;
                    if (f.name === 'full_report.html') {
                        link = '<a href="/hel/tests/report" target="_blank" style="color:#f0b429;">öffnen</a>';
                    } else if (ALLOWED.has(f.name)) {
                        const stem = f.name.replace(/\.html$/, '');
                        link = '<a href="/hel/tests/report/' + stem + '" target="_blank" style="color:#f0b429;">öffnen</a>';
                    } else {
                        link = '<span style="color:#888;">(Teil des Gesamtreports)</span>';
                    }
                    return '<tr><td style="padding:4px 8px;">' + f.name + '</td>' +
                           '<td style="padding:4px 8px; color:#aaa;">' + dt + '</td>' +
                           '<td style="padding:4px 8px; color:#aaa;">' + kb + '</td>' +
                           '<td style="padding:4px 8px;">' + link + '</td></tr>';
                }).join('');
                el.innerHTML = '<table style="width:100%; font-size:13px; border-collapse:collapse;">' +
                    '<thead><tr style="border-bottom:1px solid #555;">' +
                    '<th style="text-align:left; padding:4px 8px;">Datei</th>' +
                    '<th style="text-align:left; padding:4px 8px;">Stand</th>' +
                    '<th style="text-align:left; padding:4px 8px;">Größe</th>' +
                    '<th style="text-align:left; padding:4px 8px;">Aktion</th>' +
                    '</tr></thead><tbody>' + rows + '</tbody></table>';
            } catch (e) {
                console.warn('[P96] loadReportsList', e);
                el.textContent = 'Fehler beim Laden.';
            }
        }

        // Patch 95: Per-User-Filter — Profile-Liste laden + Change-Handler
        async function loadProfilesList() {
            const sel = document.getElementById('profileSelect');
            if (!sel) return;
            try {
                const res = await fetch('/hel/metrics/profiles');
                if (!res.ok) return;
                const data = await res.json();
                for (const p of (data.profiles || [])) {
                    const opt = document.createElement('option');
                    opt.value = p; opt.textContent = p;
                    sel.appendChild(opt);
                }
            } catch (e) { console.warn('[P95] loadProfilesList', e); }
            sel.addEventListener('change', function() {
                loadMetricsChart(_currentTimeRange);
            });
        }

        // Zeitraum-Chip-Handler
        document.addEventListener('DOMContentLoaded', function() {
            document.querySelectorAll('.time-chip').forEach(function(chip) {
                chip.addEventListener('click', function() {
                    const days = parseInt(chip.dataset.days);
                    const picker = document.getElementById('customRangePicker');
                    if (days === -1) {
                        if (picker) picker.classList.toggle('open');
                    } else {
                        if (picker) picker.classList.remove('open');
                        loadMetricsChart(days);
                    }
                });
            });
            renderMetricToggles();
            loadProfilesList();
            loadReportsList();
            // Patch 99: aktiven Tab setzen (aus localStorage oder Default 'metrics')
            try { activateTab(window.__hel_active_tab || 'metrics'); }
            catch (e) { console.warn('[P99] activateTab', e); }
        });

        async function exportSession(sessionId) {
            window.location.href = `/hel/admin/export/session/${sessionId}`;
        }

        // ========== Profil-Übersicht (Patch 61) ==========
        async function loadProfiles() {
            const tbody = document.getElementById('profilesTableBody');
            if (!tbody) return;
            try {
                const res = await fetch('/hel/admin/profiles');
                if (!res.ok) { tbody.innerHTML = '<tr><td colspan="5" style="color:#ff6b6b;">Fehler beim Laden (' + res.status + ')</td></tr>'; return; }
                const profiles = await res.json();
                if (!profiles.length) { tbody.innerHTML = '<tr><td colspan="5" style="color:#888;">Keine Profile konfiguriert.</td></tr>'; return; }
                tbody.innerHTML = profiles.map(p => `
                    <tr>
                        <td style="font-family:monospace;">${p.key}</td>
                        <td>${p.display_name}</td>
                        <td><span style="background:${p.permission_level === 'admin' ? '#7b1fa2' : p.permission_level === 'user' ? '#1a3a5c' : '#333'}; padding:3px 10px; border-radius:12px; font-size:0.85em;">${p.permission_level}</span></td>
                        <td style="font-family:monospace; font-size:0.85em;">${p.allowed_model || '<span style="color:#888;">global</span>'}</td>
                        <td style="font-family:monospace;">${p.temperature !== null && p.temperature !== undefined ? p.temperature : '<span style="color:#888;">global</span>'}</td>
                    </tr>
                `).join('');
            } catch(e) {
                tbody.innerHTML = '<tr><td colspan="5" style="color:#ff6b6b;">Netzwerkfehler: ' + e.message + '</td></tr>';
            }
        }

        // ========== RAG / Gedächtnis ==========
        // Patch 108: Farbcode pro Kategorie (gleiches Farbschema wie Permission-Badges)
        const RAG_CATEGORY_COLORS = {
            general:   '#555',
            narrative: '#7b1fa2',
            technical: '#1a3a5c',
            personal:  '#b8860b',
            lore:      '#2e7d32',
            reference: '#5d4037',
            system:    '#37474f',
        };

        // Patch 116: Gruppierte Dokumentenliste (eine Card pro Source) + Delete pro Doc
        async function loadRagStatus() {
            try {
                const res = await fetch('/hel/admin/rag/documents');
                if (!res.ok) { document.getElementById('ragIndexInfo').innerText = 'Fehler beim Laden'; return; }
                const data = await res.json();
                const activeInfo = (typeof data.total_chunks === 'number' && typeof data.total_documents === 'number')
                    ? `${data.total_documents} Dokument(e), ${data.total_chunks} Chunk(s) gesamt`
                    : `${(data.documents || []).length} Dokument(e)`;
                document.getElementById('ragIndexInfo').innerText = activeInfo;
                const sourcesDiv = document.getElementById('ragSourcesList');
                const docs = data.documents || [];
                if (docs.length === 0) {
                    sourcesDiv.innerHTML = '<span style="color:#888;">Noch keine Dokumente indiziert.</span>';
                    return;
                }
                sourcesDiv.innerHTML = docs.map(doc => {
                    const color = RAG_CATEGORY_COLORS[doc.category] || '#555';
                    const safeSrc = (doc.source || '').replace(/"/g, '&quot;');
                    const srcDisplay = (doc.source || 'unbekannt').replace(/</g, '&lt;');
                    return `
                        <div class="rag-doc-card" style="background:#2a2a2a; border:1px solid #444; border-radius:10px; padding:12px 14px; margin-bottom:10px;">
                            <div style="display:flex; justify-content:space-between; align-items:center; gap:10px; flex-wrap:wrap;">
                                <div style="flex:1; min-width:0; word-break:break-word;">
                                    <strong style="font-size:1.02em;">\uD83D\uDCC4 ${srcDisplay}</strong>
                                    <span style="background:${color}; padding:2px 8px; border-radius:10px; font-size:0.8em; margin-left:8px;">${doc.category || 'general'}</span>
                                </div>
                                <button type="button" onclick="deleteRagDocument('${safeSrc}')" style="background:#8B0000; color:#fff; padding:8px 14px; font-size:0.9em; border-radius:8px; min-height:44px; touch-action:manipulation;">&#128465; L\u00f6schen</button>
                            </div>
                            <div style="color:#aaa; font-size:0.85em; margin-top:6px;">
                                ${doc.chunk_count} Chunk(s) &middot; ${doc.total_words.toLocaleString('de-DE')} W\u00f6rter
                            </div>
                        </div>
                    `;
                }).join('');
            } catch(e) {
                document.getElementById('ragIndexInfo').innerText = 'Netzwerkfehler: ' + e.message;
            }
        }

        async function deleteRagDocument(source) {
            if (!source) return;
            if (!confirm(`Dokument "${source}" wirklich l\u00f6schen? Alle Chunks dieser Quelle werden entfernt.`)) return;
            try {
                const res = await fetch('/hel/admin/rag/document?source=' + encodeURIComponent(source), { method: 'DELETE' });
                const data = await res.json().catch(() => ({}));
                if (res.ok) {
                    loadRagStatus();
                } else {
                    alert('Fehler: HTTP ' + res.status + ' \u2013 ' + (data.detail || JSON.stringify(data)));
                }
            } catch(e) {
                alert('Netzwerkfehler: ' + e.message);
            }
        }

        // ── Patch 152 (B-020): Memory-Dashboard ──
        let _memoryList = [];

        async function loadMemoryDashboard() {
            try {
                const [lstR, statR] = await Promise.all([
                    fetch('/hel/admin/memory/list?limit=500'),
                    fetch('/hel/admin/memory/stats'),
                ]);
                if (lstR.ok) {
                    const body = await lstR.json();
                    _memoryList = Array.isArray(body) ? body : (body.memories || []);
                }
                if (statR.ok) {
                    const s = await statR.json();
                    const host = document.getElementById('memoryStats');
                    if (host) {
                        const last = s.last_extraction || s.last_extraction_at || '–';
                        const cats = s.categories || s.category_counts || {};
                        host.textContent = `${s.total || _memoryList.length} Fakten · ${Object.keys(cats).length || 0} Kategorien · Letzte Extraktion: ${last}`;
                    }
                }
            } catch (e) {
                const host = document.getElementById('memoryStats');
                if (host) host.textContent = 'Fehler: ' + e.message;
            }
            renderMemoryTable();
        }

        function _confidenceBadge(conf) {
            const c = Number(conf) || 0;
            let color = '#f88';
            if (c >= 0.9) color = '#8f8';
            else if (c >= 0.7) color = '#ff8';
            return `<span style="background:${color};color:#111;padding:2px 6px;border-radius:4px;font-size:0.78em;font-weight:700;">${c.toFixed(2)}</span>`;
        }

        function renderMemoryTable() {
            const host = document.getElementById('memoryTableHost');
            if (!host) return;
            const q = (document.getElementById('memorySearch')?.value || '').toLowerCase();
            const catFilter = (document.getElementById('memoryCategoryFilter')?.value || '').toUpperCase();
            const filtered = _memoryList.filter(m => {
                if (catFilter && (m.category || '').toUpperCase() !== catFilter) return false;
                if (!q) return true;
                const hay = [m.subject, m.fact, m.category].filter(Boolean).join(' ').toLowerCase();
                return hay.includes(q);
            });
            if (filtered.length === 0) {
                host.innerHTML = '<div style="color:#888;padding:20px;text-align:center;">Keine Einträge</div>';
                return;
            }
            const rows = filtered.map(m => `
                <tr data-id="${m.id}">
                    <td style="padding:6px;color:#FFD700;font-weight:600;">${m.category || '–'}</td>
                    <td style="padding:6px;">${(m.subject || '').replace(/</g,'&lt;')}</td>
                    <td style="padding:6px;">${(m.fact || '').replace(/</g,'&lt;')}</td>
                    <td style="padding:6px;text-align:center;">${_confidenceBadge(m.confidence)}</td>
                    <td style="padding:6px;color:#888;font-size:0.8em;">${(m.extracted_at || m.created_at || '').slice(0,16)}</td>
                    <td style="padding:6px;"><button onclick="deleteMemory(${m.id})" style="padding:2px 8px;background:transparent;border:1px solid #844;color:#e88;border-radius:4px;cursor:pointer;">✕</button></td>
                </tr>
            `).join('');
            host.innerHTML = `
                <table style="width:100%;border-collapse:collapse;font-size:0.88em;">
                    <thead>
                        <tr style="border-bottom:1px solid #333;color:#8aa0c0;text-align:left;">
                            <th style="padding:6px;">Kategorie</th>
                            <th style="padding:6px;">Subjekt</th>
                            <th style="padding:6px;">Fakt</th>
                            <th style="padding:6px;text-align:center;">Conf.</th>
                            <th style="padding:6px;">Extrahiert</th>
                            <th style="padding:6px;"></th>
                        </tr>
                    </thead>
                    <tbody>${rows}</tbody>
                </table>
            `;
        }

        async function deleteMemory(id) {
            if (!confirm('Diesen Fakt wirklich löschen?')) return;
            try {
                const r = await fetch('/hel/admin/memory/' + id, { method: 'DELETE' });
                if (r.ok) {
                    _memoryList = _memoryList.filter(m => m.id !== id);
                    renderMemoryTable();
                } else {
                    alert('Fehler: HTTP ' + r.status);
                }
            } catch (e) { alert('Fehler: ' + e.message); }
        }

        async function addMemoryManual() {
            const cat = document.getElementById('newMemoryCategory').value;
            const subj = document.getElementById('newMemorySubject').value.trim();
            const fact = document.getElementById('newMemoryFact').value.trim();
            const st = document.getElementById('newMemoryStatus');
            if (!subj || !fact) { st.textContent = 'Subjekt und Fakt sind Pflichtfelder.'; st.style.color = '#f88'; return; }
            try {
                const r = await fetch('/hel/admin/memory/add', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({ category: cat, subject: subj, fact: fact, confidence: 1.0 }),
                });
                if (r.ok) {
                    st.textContent = '✅ Gespeichert';
                    st.style.color = '#8f8';
                    document.getElementById('newMemorySubject').value = '';
                    document.getElementById('newMemoryFact').value = '';
                    await loadMemoryDashboard();
                } else {
                    st.textContent = '❌ HTTP ' + r.status;
                    st.style.color = '#f88';
                }
            } catch (e) {
                st.textContent = '❌ ' + e.message;
                st.style.color = '#f88';
            }
        }

        // Patch 115: Manueller Trigger f\u00fcr Background Memory Extraction
        async function triggerMemoryExtraction() {
            const btn = document.getElementById('memoryExtractBtn');
            const status = document.getElementById('memoryExtractStatus');
            if (!btn || !status) return;
            if (!confirm('Ged\u00e4chtnis-Extraktion jetzt starten? Das liest alle User-Nachrichten der letzten 24h und kann 30\u201390 s dauern.')) return;
            btn.disabled = true;
            btn.style.opacity = '0.6';
            status.innerHTML = '\u23F3 Extraktion l\u00e4uft\u2026 (bitte warten)';
            status.style.color = '#ffd700';
            try {
                const res = await fetch('/hel/admin/memory/extract', { method: 'POST' });
                const data = await res.json().catch(() => ({}));
                if (res.ok) {
                    const errors = (data.errors && data.errors.length) ? ' \u2014 Fehler: ' + data.errors.join(', ') : '';
                    status.innerHTML = `\u2705 Fertig: <strong>${data.extracted || 0}</strong> Fakten extrahiert, <strong>${data.indexed || 0}</strong> neu indiziert, ${data.skipped || 0} Duplikate${errors}`;
                    status.style.color = '#4ecdc4';
                    loadRagStatus();
                } else {
                    status.innerHTML = '\u274C HTTP ' + res.status + ': ' + (data.detail || JSON.stringify(data));
                    status.style.color = '#ff6b6b';
                }
            } catch(e) {
                status.innerHTML = '\u274C Netzwerkfehler: ' + e.message;
                status.style.color = '#ff6b6b';
            } finally {
                btn.disabled = false;
                btn.style.opacity = '1';
            }
        }

        // Patch 61: Dateiname im Label anzeigen nach Auswahl (Mobile-Feedback)
        function updateRagFileLabel() {
            const input = document.getElementById('ragFileInput');
            const label = document.getElementById('ragFileLabel');
            if (input.files && input.files.length > 0) {
                label.innerHTML = '\u2705 <strong>' + input.files[0].name + '</strong>';
                label.style.borderColor = '#4ecdc4';
                label.style.color = '#4ecdc4';
            } else {
                label.innerHTML = '\uD83D\uDCC4 Datei tippen/ausw\u00e4hlen (.txt oder .docx)';
                label.style.borderColor = '#555';
                label.style.color = '#ccc';
            }
        }

        async function uploadRagFile() {
            const input = document.getElementById('ragFileInput');
            const status = document.getElementById('ragUploadStatus');
            if (!input.files || input.files.length === 0) {
                status.innerHTML = '\u274C Bitte zuerst eine Datei ausw\u00e4hlen.';
                status.style.color = '#ff6b6b';
                return;
            }
            const file = input.files[0];
            const formData = new FormData();
            formData.append('file', file);
            const categorySelect = document.getElementById('ragCategory');
            const category = categorySelect ? categorySelect.value : 'auto';
            formData.append('category', category);
            status.innerHTML = '\u23F3 Wird hochgeladen und indiziert\u2026';
            status.style.color = '#ffd700';
            try {
                const res = await fetch('/hel/admin/rag/upload', { method: 'POST', body: formData });
                const data = await res.json().catch(() => ({}));
                if (res.ok) {
                    const autoFlag = data.auto_detected ? ' (auto)' : '';
                    const catTag = data.category ? ` &middot; Kategorie: <strong>${data.category}</strong>${autoFlag}` : '';
                    status.innerHTML = `\u2705 <strong>${data.chunks_indexed} Chunks indiziert</strong> aus <em>${data.filename}</em>${catTag}`;
                    status.style.color = '#4ecdc4';
                    loadRagStatus();
                    // Label zurücksetzen
                    input.value = '';
                    updateRagFileLabel();
                } else {
                    // Patch 61: HTTP-Status + detail aus Response anzeigen
                    const detail = data.detail || JSON.stringify(data) || 'Unbekannter Fehler';
                    status.innerHTML = `\u274C HTTP ${res.status}: ${detail}`;
                    status.style.color = '#ff6b6b';
                }
            } catch(e) {
                status.innerHTML = '\u274C Netzwerkfehler: ' + e.message;
                status.style.color = '#ff6b6b';
            }
        }

        async function clearRagIndex() {
            if (!confirm('Wirklich den gesamten RAG-Index l\u00f6schen? Das kann nicht r\u00fcckg\u00e4ngig gemacht werden.')) return;
            const res = await fetch('/hel/admin/rag/clear', { method: 'DELETE' });
            const data = await res.json();
            document.getElementById('ragClearStatus').innerText =
                res.ok ? '\u2705 Index geleert.' : '\u274C Fehler: ' + (data.detail || '');
            if (res.ok) loadRagStatus();
        }

        // ========== Systemsteuerung / Pacemaker ==========
        // ── Patch 150 (B-024): Pacemaker-Prozess-Steuerung ──
        const PACEMAKER_DEFAULT_PROCESSES = [
            { key: 'sentiment',  name: 'Sentiment-Analyse',  enabled: true,  interval_min: 4,  device: 'cuda' },
            { key: 'memory',     name: 'Memory-Extraktion',  enabled: true,  interval_min: 4,  device: 'cuda' },
            { key: 'db_dedup',   name: 'DB-Deduplizierung',  enabled: true,  interval_min: 30, device: 'cpu'  },
            { key: 'whisper_ping', name: 'Whisper-Ping',     enabled: true,  interval_min: 4,  device: 'cpu'  },
        ];
        let _pacemakerProcesses = null;

        async function loadPacemakerProcesses() {
            try {
                const r = await fetch('/hel/admin/pacemaker/processes');
                if (r.ok) {
                    const data = await r.json();
                    _pacemakerProcesses = {
                        master: data.master !== false,
                        sync:   !!data.sync,
                        processes: Array.isArray(data.processes) && data.processes.length
                            ? data.processes
                            : PACEMAKER_DEFAULT_PROCESSES,
                        current_activity: data.current_activity || null,
                    };
                } else {
                    _pacemakerProcesses = { master: true, sync: false, processes: PACEMAKER_DEFAULT_PROCESSES };
                }
            } catch (_) {
                _pacemakerProcesses = { master: true, sync: false, processes: PACEMAKER_DEFAULT_PROCESSES };
            }
            renderPacemakerProcesses();
        }

        function renderPacemakerProcesses() {
            const host = document.getElementById('pacemakerProcesses');
            if (!host || !_pacemakerProcesses) return;
            const master = _pacemakerProcesses.master;
            const sync = _pacemakerProcesses.sync;
            document.getElementById('pacemaker-master').checked = master;
            document.getElementById('pacemaker-sync').checked = sync;
            host.innerHTML = '';
            _pacemakerProcesses.processes.forEach((p, idx) => {
                const row = document.createElement('div');
                row.className = 'pacemaker-process';
                row.style.cssText = 'display:flex;gap:10px;align-items:center;padding:8px 10px;margin-bottom:6px;background:rgba(255,255,255,0.02);border:1px solid #2a2a2a;border-radius:6px;flex-wrap:wrap;';
                const name = document.createElement('span');
                name.className = 'process-name';
                name.textContent = p.name;
                name.style.cssText = 'flex:1 1 180px;font-weight:600;color:#eee;';
                const enabled = document.createElement('label');
                enabled.style.cssText = 'color:#ccc;font-size:0.88em;';
                enabled.innerHTML = `<input type="checkbox" ${p.enabled ? 'checked' : ''} ${master ? '' : 'disabled'}> Aktiv`;
                enabled.querySelector('input').addEventListener('change', (e) => {
                    _pacemakerProcesses.processes[idx].enabled = e.target.checked;
                });
                const status = document.createElement('span');
                status.className = 'process-status';
                status.textContent = (master && p.enabled) ? '🟢' : '⚪';
                status.style.flex = '0 0 24px';
                // Interval-Slider
                const sliderWrap = document.createElement('div');
                sliderWrap.style.cssText = 'flex:2 1 220px;display:flex;align-items:center;gap:6px;';
                const slider = document.createElement('input');
                slider.type = 'range'; slider.min = '1'; slider.max = '60'; slider.value = String(p.interval_min);
                slider.className = 'interval-slider';
                slider.style.flex = '1';
                const lbl = document.createElement('span');
                lbl.className = 'interval-label';
                lbl.style.cssText = 'flex:0 0 60px;text-align:right;font-family:monospace;color:#888;';
                lbl.textContent = p.interval_min + ' min';
                slider.addEventListener('input', () => {
                    const v = parseInt(slider.value, 10);
                    lbl.textContent = v + ' min';
                    if (_pacemakerProcesses.sync) {
                        _pacemakerProcesses.processes.forEach((pp, j) => {
                            pp.interval_min = v;
                            const other = host.querySelectorAll('.interval-slider')[j];
                            if (other && other !== slider) other.value = String(v);
                            const otherLbl = host.querySelectorAll('.interval-label')[j];
                            if (otherLbl) otherLbl.textContent = v + ' min';
                        });
                    } else {
                        _pacemakerProcesses.processes[idx].interval_min = v;
                    }
                });
                sliderWrap.appendChild(slider); sliderWrap.appendChild(lbl);
                // Device-Select
                const devSel = document.createElement('select');
                devSel.className = 'device-select';
                devSel.style.cssText = 'padding:4px 8px;background:#121212;color:#eee;border:1px solid #444;border-radius:4px;';
                ['cpu', 'cuda'].forEach(dev => {
                    const opt = document.createElement('option');
                    opt.value = dev; opt.textContent = dev === 'cpu' ? 'CPU' : 'GPU';
                    if (dev === p.device) opt.selected = true;
                    devSel.appendChild(opt);
                });
                devSel.addEventListener('change', () => {
                    _pacemakerProcesses.processes[idx].device = devSel.value;
                });
                row.appendChild(name);
                row.appendChild(enabled);
                row.appendChild(status);
                row.appendChild(sliderWrap);
                row.appendChild(devSel);
                host.appendChild(row);
            });
            // Activity-Display
            const act = document.getElementById('pacemakerActivity');
            if (act) {
                if (_pacemakerProcesses.current_activity) {
                    act.textContent = 'Aktuell: ' + _pacemakerProcesses.current_activity;
                } else {
                    act.textContent = '';
                }
            }
        }

        async function savePacemakerProcesses() {
            if (!_pacemakerProcesses) return;
            _pacemakerProcesses.master = document.getElementById('pacemaker-master').checked;
            _pacemakerProcesses.sync = document.getElementById('pacemaker-sync').checked;
            const st = document.getElementById('pacemakerProcessesStatus');
            try {
                const r = await fetch('/hel/admin/pacemaker/processes', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(_pacemakerProcesses),
                });
                if (r.ok) {
                    st.textContent = '✅ Gespeichert';
                    st.style.color = '#8f8';
                    renderPacemakerProcesses();
                } else {
                    st.textContent = '❌ ' + (await r.text());
                    st.style.color = '#f88';
                }
            } catch (e) {
                st.textContent = '❌ ' + e.message;
                st.style.color = '#f88';
            }
        }

        async function loadPacemakerConfig() {
            try {
                const res = await fetch('/hel/admin/pacemaker/config');
                const data = await res.json();
                document.getElementById('pacemakerMinutes').value = data.keep_alive_minutes ?? 120;
            } catch(e) {}
        }

        async function savePacemakerConfig() {
            const minutes = parseInt(document.getElementById('pacemakerMinutes').value);
            const res = await fetch('/hel/admin/pacemaker/config', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({ keep_alive_minutes: minutes })
            });
            const data = await res.json();
            document.getElementById('pacemakerStatus').innerText =
                res.ok ? '\u2705 Gespeichert. Wirkt nach Neustart.' : '\u274C Fehler: ' + (data.detail || '');
        }

        // Patch 119: Whisper manuell neu starten
        async function restartWhisperContainer() {
            const status = document.getElementById('whisperRestartStatus');
            status.textContent = '\u23f3 Starte Container neu ...';
            status.style.color = '#ffd700';
            try {
                const res = await fetch('/hel/admin/whisper/restart', { method: 'POST' });
                const data = await res.json();
                if (res.ok && data.restart_success && data.post_restart_healthy) {
                    status.textContent = '\u2705 Restart erfolgreich, Whisper antwortet wieder.';
                    status.style.color = '#4ecdc4';
                } else if (res.ok && data.restart_success) {
                    status.textContent = '\u26a0\ufe0f Container neu gestartet, aber Health-Check fehlgeschlagen.';
                    status.style.color = '#ff9800';
                } else {
                    status.textContent = '\u274c Restart fehlgeschlagen: ' + (data.detail || JSON.stringify(data));
                    status.style.color = '#f44336';
                }
            } catch (e) {
                status.textContent = '\u274c Fehler: ' + e.message;
                status.style.color = '#f44336';
            }
        }

        // ========== Provider-Blacklist (Patch 63 / Patch 170 B3) ==========
        // Statisches Superset bekannter OpenRouter-Provider (Stand April 2026).
        // Pragmatik: nicht alle sind immer verfügbar, aber für eine Blacklist OK.
        const KNOWN_PROVIDERS = [
            'azure', 'aws bedrock', 'google cloud vertex', 'together', 'fireworks',
            'lepton', 'avian', 'lambda', 'anyscale', 'modal', 'replicate', 'octoai',
            'deepinfra', 'mancer', 'lynn', 'infermatic', 'sf compute', 'cloudflare',
            'featherless', 'targon', 'chutes', 'novita', 'parasail'
        ];
        let currentBlacklist = [];

        async function loadProviderBlacklist() {
            const res = await fetch('/hel/admin/provider_blacklist');
            const data = await res.json();
            currentBlacklist = data.blacklist || [];
            buildProviderSelect();
            renderProviderList();
        }

        function buildProviderSelect() {
            const sel = document.getElementById('newProviderSelect');
            if (!sel) return;
            const blacklisted = new Set(currentBlacklist.map(p => p.toLowerCase()));
            const opts = ['<option value="">— Provider wählen —</option>'];
            KNOWN_PROVIDERS.forEach(p => {
                if (blacklisted.has(p)) return;
                opts.push(`<option value="${p}">${p}</option>`);
            });
            opts.push('<option value="__custom__">Benutzerdefiniert…</option>');
            sel.innerHTML = opts.join('');
            const inp = document.getElementById('newProviderInput');
            if (inp) { inp.style.display = 'none'; inp.value = ''; }
        }

        function onProviderSelectChange() {
            const sel = document.getElementById('newProviderSelect');
            const inp = document.getElementById('newProviderInput');
            if (!sel || !inp) return;
            inp.style.display = (sel.value === '__custom__') ? '' : 'none';
            if (sel.value === '__custom__') inp.focus();
        }

        function renderProviderList() {
            const container = document.getElementById('providerList');
            if (!container) return;
            if (currentBlacklist.length === 0) {
                container.innerHTML = '<span style="color:#888;">Keine Provider blockiert.</span>';
                return;
            }
            // Chips: kompakt, inline, wrap auf Mobile
            container.innerHTML = currentBlacklist.map(p => {
                const safe = String(p).replace(/"/g, '&quot;').replace(/</g, '&lt;');
                return `<span class="provider-chip" style="display:inline-flex;align-items:center;gap:4px;`
                     + `padding:2px 8px;max-height:32px;background:#2a2a2a;`
                     + `border:1px solid #444;border-radius:var(--zb-radius-sm,4px);`
                     + `font-size:0.9em;line-height:1.4;">`
                     + `<span>${safe}</span>`
                     + `<button onclick="removeProvider('${safe}')" `
                     + `title="Entfernen" `
                     + `style="background:transparent;border:0;color:#e88;cursor:pointer;`
                     + `padding:0 2px;font-size:1em;line-height:1;min-height:0;min-width:0;">&#10005;</button>`
                     + `</span>`;
            }).join('');
        }

        function addProvider() {
            const sel = document.getElementById('newProviderSelect');
            const inp = document.getElementById('newProviderInput');
            const status = document.getElementById('providerStatus');
            let name = '';
            if (sel && sel.value && sel.value !== '__custom__') {
                name = sel.value.trim().toLowerCase();
            } else if (sel && sel.value === '__custom__' && inp) {
                name = inp.value.trim().toLowerCase();
            }
            if (!name) {
                if (status) status.textContent = 'Bitte Provider auswählen oder eingeben.';
                return;
            }
            if (currentBlacklist.includes(name)) {
                if (status) status.textContent = 'Bereits in der Liste.';
                return;
            }
            currentBlacklist.push(name);
            buildProviderSelect();
            renderProviderList();
            if (status) status.textContent = '';
        }

        function removeProvider(name) {
            currentBlacklist = currentBlacklist.filter(p => p !== name);
            buildProviderSelect();
            renderProviderList();
        }

        async function saveProviderBlacklist() {
            const res = await fetch('/hel/admin/provider_blacklist', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({ blacklist: currentBlacklist })
            });
            const data = await res.json();
            document.getElementById('providerStatus').textContent =
                data.status === 'ok' ? '\u2705 Gespeichert.' : '\u274C Fehler beim Speichern.';
        }

        // Patch 90 (N-F09b): Schriftgröße-Wahl
        function setFontSize(px) {
            document.documentElement.style.setProperty('--hel-font-size-base', px);
            try { localStorage.setItem('hel_font_size', px); } catch (_) {}
            markActiveFontPreset();
        }
        function markActiveFontPreset() {
            var current = (getComputedStyle(document.documentElement)
                .getPropertyValue('--hel-font-size-base') || '').trim();
            document.querySelectorAll('.font-preset-btn').forEach(function (b) {
                b.classList.toggle('active', b.dataset.size === current);
            });
        }
        markActiveFontPreset();

        // ── Patch 149 (B-025): Hel Mini-Settings mit Schrift-Slider ──
        function toggleHelSettings() {
            const p = document.getElementById('helSettingsPanel');
            if (!p) return;
            p.style.display = (p.style.display === 'none' || !p.style.display) ? 'block' : 'none';
        }
        function applyHelUiScale(val) {
            const f = parseFloat(val);
            if (!isFinite(f) || f < 0.5 || f > 2.0) return;
            document.documentElement.style.setProperty('--ui-scale', String(f));
            document.documentElement.style.setProperty('--hel-font-size-base', (16 * f) + 'px');
            try { localStorage.setItem('hel_ui_scale', String(f)); } catch(_) {}
            const v = document.getElementById('helUiScaleVal');
            if (v) v.textContent = f.toFixed(2) + '×';
            markActiveFontPreset();
        }
        function resetHelUiScale() {
            document.documentElement.style.removeProperty('--ui-scale');
            document.documentElement.style.removeProperty('--hel-font-size-base');
            try { localStorage.removeItem('hel_ui_scale'); } catch(_) {}
            const s = document.getElementById('helUiScaleSlider');
            const v = document.getElementById('helUiScaleVal');
            if (s) s.value = '1.0';
            if (v) v.textContent = '1.00×';
        }
        (function restoreHelUiScale() {
            try {
                const stored = localStorage.getItem('hel_ui_scale');
                if (!stored) return;
                const f = parseFloat(stored);
                if (!isFinite(f) || f < 0.5 || f > 2.0) return;
                document.documentElement.style.setProperty('--ui-scale', String(f));
                document.documentElement.style.setProperty('--hel-font-size-base', (16 * f) + 'px');
                const s = document.getElementById('helUiScaleSlider');
                const v = document.getElementById('helUiScaleVal');
                if (s) s.value = String(f);
                if (v) v.textContent = f.toFixed(2) + '×';
            } catch(_) {}
        })();

        loadModelsAndBalance();
        loadCleaner();
        loadFuzzyDict();
        loadDialect();
        // Patch 85: Metriken laden da Sektion standardmäßig offen
        loadMetrics();

        // ---- User-Verwaltung (Patch 83) ----
        async function loadUserProfiles() {
            try {
                const res = await fetch('/hel/admin/profiles');
                const profiles = await res.json();
                const selTest = document.getElementById('hashTestProfile');
                const selReset = document.getElementById('resetProfile');
                selTest.innerHTML = '';
                selReset.innerHTML = '';
                profiles.forEach(p => {
                    const o1 = document.createElement('option');
                    o1.value = p.key;
                    o1.textContent = p.display_name + ' (' + p.key + ')';
                    selTest.appendChild(o1);
                    selReset.appendChild(o1.cloneNode(true));
                });
            } catch(e) { console.error('Profile laden fehlgeschlagen:', e); }
        }
        loadUserProfiles();

        async function testHash() {
            const profile = document.getElementById('hashTestProfile').value;
            const password = document.getElementById('hashTestPassword').value;
            const el = document.getElementById('hashTestResult');
            if (!password) { el.innerHTML = '<span style="color:#ffa500;">Bitte Passwort eingeben</span>'; return; }
            try {
                const res = await fetch('/hel/admin/auth/test-hash', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({ profile, password })
                });
                const data = await res.json();
                if (!data.hash_exists) {
                    el.innerHTML = '<span style="color:#ffa500;">\u26a0\ufe0f Kein Hash gespeichert</span>';
                } else if (data.match) {
                    el.innerHTML = '<span style="color:#4ecdc4;">\u2705 Passwort stimmt!</span>';
                } else {
                    el.innerHTML = '<span style="color:#ff6b6b;">\u274c Falsches Passwort</span>';
                }
            } catch(e) { el.innerHTML = '<span style="color:#ff6b6b;">Fehler: ' + e.message + '</span>'; }
        }

        async function resetProfilePassword() {
            const profile = document.getElementById('resetProfile').value;
            const password = document.getElementById('resetPassword').value;
            const el = document.getElementById('resetResult');
            if (!password || password.length < 4) {
                el.innerHTML = '<span style="color:#ffa500;">Mind. 4 Zeichen!</span>';
                return;
            }
            if (!confirm('Passwort f\u00fcr "' + profile + '" wirklich \u00e4ndern?')) return;
            try {
                const res = await fetch('/hel/admin/auth/reset-password', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({ profile, password })
                });
                const data = await res.json();
                if (data.success) {
                    el.innerHTML = '<span style="color:#4ecdc4;">\u2705 Passwort f\u00fcr ' + profile + ' gesetzt!</span>';
                } else {
                    el.innerHTML = '<span style="color:#ff6b6b;">\u274c ' + (data.detail || 'Fehler') + '</span>';
                }
            } catch(e) { el.innerHTML = '<span style="color:#ff6b6b;">Fehler: ' + e.message + '</span>'; }
        }

        // ==================== Patch 195 (Phase 5a #1): Projekte-UI ====================
        let _projectsCache = [];

        function _projectStatusBadge(isArchived) {
            if (isArchived) {
                return '<span style="background:#5c2f2f;color:#f88;padding:2px 8px;border-radius:4px;font-size:0.78em;">archiviert</span>';
            }
            return '<span style="background:#1a5c1a;color:#4ecdc4;padding:2px 8px;border-radius:4px;font-size:0.78em;">aktiv</span>';
        }

        function _projectFmtDate(iso) {
            if (!iso) return '-';
            try {
                const d = new Date(iso);
                return d.toLocaleString('de-DE', { dateStyle: 'short', timeStyle: 'short' });
            } catch(_) { return iso; }
        }

        function _escapeHtml(s) {
            if (s === null || s === undefined) return '';
            return String(s)
                .replace(/&/g, '&amp;')
                .replace(/</g, '&lt;')
                .replace(/>/g, '&gt;')
                .replace(/"/g, '&quot;')
                .replace(/'/g, '&#39;');
        }

        async function loadProjects() {
            const tbody = document.getElementById('projectsTableBody');
            const status = document.getElementById('projectsStatus');
            if (!tbody || !status) return;
            const includeArchived = document.getElementById('projectsShowArchived').checked;
            status.textContent = 'Lade Projekte...';
            status.style.color = '#aaa';
            try {
                const r = await fetch('/hel/admin/projects?include_archived=' + (includeArchived ? 'true' : 'false'));
                if (!r.ok) throw new Error('HTTP ' + r.status);
                const data = await r.json();
                _projectsCache = data.projects || [];
                if (_projectsCache.length === 0) {
                    tbody.innerHTML = '<tr><td colspan="5" style="padding:14px;color:#888;text-align:center;">Noch keine Projekte. &laquo;+ Projekt anlegen&raquo; oben klicken.</td></tr>';
                } else {
                    tbody.innerHTML = _projectsCache.map(p => {
                        const archiveBtn = p.is_archived
                            ? '<button onclick="unarchiveProject(' + p.id + ')" style="padding:6px 10px;min-height:36px;">&#9851; Reaktivieren</button>'
                            : '<button onclick="archiveProject(' + p.id + ')" style="padding:6px 10px;min-height:36px;">&#128451;&#65039; Archivieren</button>';
                        return '<tr style="border-bottom:1px solid #333;">'
                            + '<td style="padding:8px;font-family:monospace;color:#4ecdc4;">' + _escapeHtml(p.slug) + '</td>'
                            + '<td style="padding:8px;cursor:pointer;color:#ffd700;" onclick="loadProjectFiles(' + p.id + ')" title="Dateien anzeigen">' + _escapeHtml(p.name) + '</td>'
                            + '<td style="padding:8px;color:#aaa;font-size:0.88em;">' + _projectFmtDate(p.updated_at) + '</td>'
                            + '<td style="padding:8px;">' + _projectStatusBadge(p.is_archived) + '</td>'
                            + '<td style="padding:8px;display:flex;gap:6px;flex-wrap:wrap;">'
                                + '<button onclick="editProject(' + p.id + ')" style="padding:6px 10px;min-height:36px;">&#9999;&#65039; Edit</button>'
                                + archiveBtn
                                + '<button onclick="deleteProject(' + p.id + ')" style="padding:6px 10px;min-height:36px;background:#5c2f2f;color:#f88;">&#128465;&#65039; Loeschen</button>'
                            + '</td>'
                            + '</tr>';
                    }).join('');
                }
                status.textContent = _projectsCache.length + ' Projekt(e) geladen.';
                status.style.color = '#4ecdc4';
            } catch (e) {
                status.textContent = 'Fehler: ' + e.message;
                status.style.color = '#f88';
            }
        }

        function openProjectForm(project) {
            document.getElementById('projectFormTitle').textContent = project ? ('Projekt bearbeiten: ' + project.slug) : 'Projekt anlegen';
            document.getElementById('projectFormId').value = project ? project.id : '';
            document.getElementById('projectFormName').value = project ? (project.name || '') : '';
            document.getElementById('projectFormDescription').value = project ? (project.description || '') : '';
            const slugInput = document.getElementById('projectFormSlug');
            slugInput.value = project ? (project.slug || '') : '';
            slugInput.disabled = !!project; // Slug ist immutable, nur bei Anlage editierbar
            const overlay = (project && project.persona_overlay) || { system_addendum: '', tone_hints: [] };
            document.getElementById('projectFormSystemAddendum').value = overlay.system_addendum || '';
            document.getElementById('projectFormToneHints').value = (overlay.tone_hints || []).join(', ');
            document.getElementById('projectFormStatus').textContent = '';
            document.getElementById('projectFormOverlay').style.display = 'flex';
        }

        function closeProjectForm() {
            document.getElementById('projectFormOverlay').style.display = 'none';
        }

        async function saveProjectForm() {
            const id = document.getElementById('projectFormId').value;
            const name = document.getElementById('projectFormName').value.trim();
            const description = document.getElementById('projectFormDescription').value.trim();
            const slug = document.getElementById('projectFormSlug').value.trim();
            const systemAddendum = document.getElementById('projectFormSystemAddendum').value;
            const toneHintsRaw = document.getElementById('projectFormToneHints').value;
            const statusEl = document.getElementById('projectFormStatus');
            if (!name) {
                statusEl.textContent = 'Name darf nicht leer sein.';
                statusEl.style.color = '#f88';
                return;
            }
            const toneHints = toneHintsRaw.split(',').map(s => s.trim()).filter(Boolean);
            const personaOverlay = (systemAddendum || toneHints.length > 0)
                ? { system_addendum: systemAddendum, tone_hints: toneHints }
                : null;
            const payload = { name, description, persona_overlay: personaOverlay };
            try {
                let r;
                if (id) {
                    r = await fetch('/hel/admin/projects/' + id, {
                        method: 'PATCH',
                        headers: {'Content-Type': 'application/json'},
                        body: JSON.stringify(payload),
                    });
                } else {
                    if (slug) payload.slug = slug;
                    r = await fetch('/hel/admin/projects', {
                        method: 'POST',
                        headers: {'Content-Type': 'application/json'},
                        body: JSON.stringify(payload),
                    });
                }
                if (!r.ok) {
                    const errBody = await r.json().catch(() => ({}));
                    throw new Error(errBody.detail || ('HTTP ' + r.status));
                }
                statusEl.textContent = 'Gespeichert.';
                statusEl.style.color = '#4ecdc4';
                closeProjectForm();
                await loadProjects();
            } catch (e) {
                statusEl.textContent = 'Fehler: ' + e.message;
                statusEl.style.color = '#f88';
            }
        }

        function editProject(id) {
            const project = _projectsCache.find(p => p.id === id);
            if (!project) return;
            openProjectForm(project);
        }

        async function archiveProject(id) {
            if (!confirm('Projekt archivieren? Bleibt erhalten, wird ausgeblendet.')) return;
            try {
                const r = await fetch('/hel/admin/projects/' + id + '/archive', { method: 'POST' });
                if (!r.ok) throw new Error('HTTP ' + r.status);
                await loadProjects();
            } catch (e) {
                document.getElementById('projectsStatus').textContent = 'Fehler: ' + e.message;
                document.getElementById('projectsStatus').style.color = '#f88';
            }
        }

        async function unarchiveProject(id) {
            try {
                const r = await fetch('/hel/admin/projects/' + id + '/unarchive', { method: 'POST' });
                if (!r.ok) throw new Error('HTTP ' + r.status);
                await loadProjects();
            } catch (e) {
                document.getElementById('projectsStatus').textContent = 'Fehler: ' + e.message;
                document.getElementById('projectsStatus').style.color = '#f88';
            }
        }

        async function deleteProject(id) {
            const project = _projectsCache.find(p => p.id === id);
            const label = project ? (project.slug + ' (' + project.name + ')') : ('#' + id);
            if (!confirm('Projekt ' + label + ' UNWIDERRUFLICH loeschen? Datei-Metadaten werden mitgeloescht (Bytes bleiben im Storage).')) return;
            try {
                const r = await fetch('/hel/admin/projects/' + id, { method: 'DELETE' });
                if (!r.ok) throw new Error('HTTP ' + r.status);
                document.getElementById('projectDetailCard').style.display = 'none';
                await loadProjects();
            } catch (e) {
                document.getElementById('projectsStatus').textContent = 'Fehler: ' + e.message;
                document.getElementById('projectsStatus').style.color = '#f88';
            }
        }

        async function loadProjectFiles(projectId) {
            const project = _projectsCache.find(p => p.id === projectId);
            const card = document.getElementById('projectDetailCard');
            const list = document.getElementById('projectFilesList');
            const nameEl = document.getElementById('projectDetailName');
            if (!card || !list) return;
            card.style.display = 'block';
            nameEl.textContent = project ? '— ' + project.slug : '';
            list.innerHTML = '<span style="color:#888;">Lade...</span>';
            // Patch 196: Drop-Zone fuer dieses Projekt verdrahten
            _setupProjectFileDrop(projectId);
            try {
                const r = await fetch('/hel/admin/projects/' + projectId + '/files');
                if (!r.ok) throw new Error('HTTP ' + r.status);
                const data = await r.json();
                const files = data.files || [];
                if (files.length === 0) {
                    list.innerHTML = '<span style="color:#888;">Keine Dateien. Drop-Zone oben benutzen.</span>';
                } else {
                    // P203b-Hotfix: data-Attribute + Event-Delegation statt inline-onclick.
                    // Die alte Variante mit ',\\'' + replace(/'/g, "\\\\'") + '\\''
                    // hatte falsches Python-Quote-Escaping — das ergab in JS einen
                    // Syntax-Fehler ("Unexpected string"), der den GANZEN <script>-
                    // Block ungueltig machte und damit alle Hel-UI-Klicks blockierte.
                    list.innerHTML = files.map(f => {
                        const kb = (f.size_bytes / 1024).toFixed(1);
                        return '<div style="padding:6px 0;border-bottom:1px solid #2a2a2a;display:flex;justify-content:space-between;align-items:center;gap:8px;flex-wrap:wrap;">'
                            + '<div style="flex:1;min-width:0;">'
                                + '<span style="color:#4ecdc4;">' + _escapeHtml(f.relative_path) + '</span> '
                                + '<span style="color:#888;">(' + kb + ' KB, ' + _escapeHtml(f.mime_type || 'unknown') + ')</span>'
                            + '</div>'
                            + '<button type="button" class="proj-file-delete-btn"'
                                + ' data-project-id="' + projectId + '"'
                                + ' data-file-id="' + f.id + '"'
                                + ' data-relative-path="' + _escapeHtml(f.relative_path) + '"'
                                + ' style="padding:4px 8px;min-height:36px;background:#5c2f2f;color:#f88;border:1px solid #844;border-radius:4px;font-size:0.85em;">'
                                + '&#128465;&#65039; Loeschen</button>'
                            + '</div>';
                    }).join('');
                    list.querySelectorAll('.proj-file-delete-btn').forEach(btn => {
                        btn.addEventListener('click', () => {
                            const pid = parseInt(btn.dataset.projectId, 10);
                            const fid = parseInt(btn.dataset.fileId, 10);
                            const rel = btn.dataset.relativePath || '';
                            deleteProjectFile(pid, fid, rel);
                        });
                    });
                }
            } catch (e) {
                list.innerHTML = '<span style="color:#f88;">Fehler: ' + e.message + '</span>';
            }
        }

        // ==================== /Patch 195 ====================

        // ==================== Patch 196 (Phase 5a #4): Datei-Upload-UI ====================
        let _projectDropWired = false;

        function _setupProjectFileDrop(projectId) {
            const zone = document.getElementById('projectDropZone');
            const input = document.getElementById('projectFileInput');
            if (!zone || !input) return;
            zone.dataset.projectId = String(projectId);
            input.dataset.projectId = String(projectId);
            if (_projectDropWired) return;
            _projectDropWired = true;

            zone.addEventListener('click', () => input.click());
            input.addEventListener('change', () => {
                const pid = parseInt(input.dataset.projectId, 10);
                if (input.files && input.files.length > 0) {
                    _uploadProjectFiles(pid, input.files);
                    input.value = '';
                }
            });

            ['dragenter', 'dragover'].forEach(ev => {
                zone.addEventListener(ev, (e) => {
                    e.preventDefault();
                    e.stopPropagation();
                    zone.style.background = '#1a3a3a';
                });
            });
            ['dragleave', 'drop'].forEach(ev => {
                zone.addEventListener(ev, (e) => {
                    e.preventDefault();
                    e.stopPropagation();
                    zone.style.background = '';
                });
            });
            zone.addEventListener('drop', (e) => {
                const pid = parseInt(zone.dataset.projectId, 10);
                if (e.dataTransfer && e.dataTransfer.files.length > 0) {
                    _uploadProjectFiles(pid, e.dataTransfer.files);
                }
            });
        }

        // Patch 205 (Phase 5a Schuld aus P199): RAG-Status-Toast.
        // Reason-Mapping: kurze de-DE-Strings fuer die Codes aus
        // ``zerberus/core/projects_rag.py::index_project_file``. Falls ein
        // Backend-Code unbekannt ist (zukuenftiger Code), faellt das Mapping
        // auf "uebersprungen" zurueck.
        const _RAG_REASON_LABELS = {
            'rag_disabled': 'RAG aus',
            'too_large': 'zu gross',
            'binary': 'Binärdatei',
            'empty': 'leere Datei',
            'no_chunks': 'kein Inhalt',
            'embed_failed': 'Embed-Fehler',
            'file_not_found': 'Datei nicht gefunden',
            'project_not_found': 'Projekt nicht gefunden',
            'bytes_missing': 'Bytes fehlen',
            'exception': 'Indizierungs-Fehler'
        };

        function _showRagToast(rag) {
            const el = document.getElementById('ragToast');
            if (!el || !rag) return;
            let text;
            let klass;
            if (rag.skipped) {
                const label = _RAG_REASON_LABELS[rag.reason] || 'übersprungen';
                klass = 'warn';
                text = '⚠ Datei nicht indiziert: ' + label;
            } else {
                const n = (typeof rag.chunks === 'number') ? rag.chunks : 0;
                klass = 'success';
                text = '📚 ' + n + ' Chunks indiziert';
            }
            // textContent ist XSS-immun — Reason-Strings stammen aus dem
            // statischen Mapping, niemals vom Server-Pfad direkt.
            el.textContent = text;
            el.classList.remove('warn', 'success');
            el.classList.add(klass, 'visible');
            if (_showRagToast._t) clearTimeout(_showRagToast._t);
            _showRagToast._t = setTimeout(function () {
                el.classList.remove('visible');
            }, 3500);
            el.onclick = function () {
                if (_showRagToast._t) clearTimeout(_showRagToast._t);
                el.classList.remove('visible');
            };
        }

        function _uploadProjectFiles(projectId, fileList) {
            const progress = document.getElementById('projectUploadProgress');
            if (!progress) return;
            const files = Array.from(fileList);
            let done = 0;
            progress.innerHTML = files.map((f, i) =>
                '<div id="proj-up-' + i + '" style="padding:2px 0;color:#aaa;">' + _escapeHtml(f.name) + ' — wartet...</div>'
            ).join('');

            const uploadOne = (f, idx) => new Promise((resolve) => {
                const row = document.getElementById('proj-up-' + idx);
                const fd = new FormData();
                fd.append('file', f, f.name);
                const xhr = new XMLHttpRequest();
                xhr.open('POST', '/hel/admin/projects/' + projectId + '/files', true);
                xhr.upload.addEventListener('progress', (ev) => {
                    if (ev.lengthComputable && row) {
                        const pct = Math.round((ev.loaded / ev.total) * 100);
                        row.innerHTML = _escapeHtml(f.name) + ' — ' + pct + '%';
                        row.style.color = '#4ecdc4';
                    }
                });
                xhr.onload = () => {
                    if (row) {
                        if (xhr.status >= 200 && xhr.status < 300) {
                            row.innerHTML = '&#10004; ' + _escapeHtml(f.name) + ' — fertig';
                            row.style.color = '#4ecdc4';
                            // Patch 205: RAG-Status nach Erfolgs-Render als Toast zeigen.
                            // Fail-quiet: kaputter Body bricht den Upload-Loop nicht ab.
                            try {
                                const body = JSON.parse(xhr.responseText);
                                if (body && body.rag) _showRagToast(body.rag);
                            } catch (_) {}
                        } else {
                            let detail = 'HTTP ' + xhr.status;
                            try {
                                const body = JSON.parse(xhr.responseText);
                                if (body && body.detail) detail = body.detail;
                            } catch (_) {}
                            row.innerHTML = '&#10006; ' + _escapeHtml(f.name) + ' — ' + _escapeHtml(detail);
                            row.style.color = '#f88';
                        }
                    }
                    resolve();
                };
                xhr.onerror = () => {
                    if (row) {
                        row.innerHTML = '&#10006; ' + _escapeHtml(f.name) + ' — Netzwerk-Fehler';
                        row.style.color = '#f88';
                    }
                    resolve();
                };
                xhr.send(fd);
            });

            // Sequenziell, damit der Server nicht mit parallelen Uploads
            // ueberrannt wird und die Progress-Anzeige nachvollziehbar bleibt.
            (async () => {
                for (let i = 0; i < files.length; i++) {
                    await uploadOne(files[i], i);
                    done++;
                }
                await loadProjectFiles(projectId);
            })();
        }

        async function deleteProjectFile(projectId, fileId, label) {
            if (!confirm('Datei "' + label + '" loeschen? Bytes werden entfernt, wenn sie nirgends sonst referenziert sind.')) return;
            try {
                const r = await fetch('/hel/admin/projects/' + projectId + '/files/' + fileId, { method: 'DELETE' });
                if (!r.ok) {
                    const errBody = await r.json().catch(() => ({}));
                    throw new Error(errBody.detail || ('HTTP ' + r.status));
                }
                await loadProjectFiles(projectId);
            } catch (e) {
                const list = document.getElementById('projectFilesList');
                if (list) list.innerHTML = '<span style="color:#f88;">Loeschen fehlgeschlagen: ' + _escapeHtml(e.message) + '</span>';
            }
        }
        // ==================== /Patch 196 ====================
    </script>
    <!-- Patch 205: RAG-Status-Toast (wird nach jedem Datei-Upload kurz eingeblendet).
         Position via CSS fixed (bottom-right). Inhalt wird per textContent gesetzt
         (XSS-immun). Klick + Auto-Timeout dismissen via .visible-Klasse. -->
    <div id="ragToast" class="rag-toast" role="status" aria-live="polite"></div>
    <!-- Patch 211: GPU-Queue-Wartetoast. Polling alle 4s gegen /v1/gpu/status,
         eingeblendet sobald mind. ein Konsument in der Queue wartet. textContent
         (XSS-immun) — Consumer-Namen sind Whitelist (whisper/gemma/embedder/reranker). -->
    <div id="gpuToast" class="gpu-toast" role="status" aria-live="polite"></div>
    <script>
        // Patch 211: GPU-Queue-Polling. Eine Funktion, lazy-startet, kein
        // doppeltes Interval, sauberer Cleanup beim Page-Unload.
        (function () {
            const POLL_INTERVAL_MS = 4000;
            const CONSUMER_LABEL = {
                'whisper': 'Whisper',
                'gemma': 'Gemma',
                'embedder': 'Embedder',
                'reranker': 'Reranker'
            };
            let _timer = null;

            function _renderGpuToast(status) {
                const el = document.getElementById('gpuToast');
                if (!el) return;
                const waiters = (status && Array.isArray(status.waiters)) ? status.waiters : [];
                if (waiters.length === 0) {
                    el.classList.remove('visible');
                    return;
                }
                const first = waiters[0] || {};
                const name = CONSUMER_LABEL[first.consumer] || (first.consumer || 'GPU-Konsument');
                let msg = '⏳ GPU wartet auf ' + name;
                if (waiters.length > 1) {
                    msg += ' — Position ' + waiters.length + ' in Queue';
                }
                // textContent ist XSS-immun. Consumer-Namen kommen aus Whitelist.
                el.textContent = msg;
                el.classList.add('visible');
            }

            function _pollOnce() {
                fetch('/v1/gpu/status', { method: 'GET', cache: 'no-store' })
                    .then(function (r) { return r.ok ? r.json() : null; })
                    .then(function (data) { _renderGpuToast(data); })
                    .catch(function () { /* fail-quiet */ });
            }

            function _start() {
                if (_timer !== null) return;
                _pollOnce();
                _timer = setInterval(_pollOnce, POLL_INTERVAL_MS);
            }

            function _stop() {
                if (_timer !== null) {
                    clearInterval(_timer);
                    _timer = null;
                }
            }

            window.addEventListener('load', _start);
            window.addEventListener('beforeunload', _stop);

            // Klick auf den Toast = Soft-Dismiss (kommt beim naechsten Poll wieder).
            const el = document.getElementById('gpuToast');
            if (el) {
                el.addEventListener('click', function () {
                    el.classList.remove('visible');
                });
            }
        })();
    </script>
    <!-- Patch 200: Service-Worker-Registrierung (PWA App-Shell-Cache). -->
    <script>
        if ('serviceWorker' in navigator) {
            window.addEventListener('load', function () {
                navigator.serviceWorker.register('/hel/sw.js', { scope: '/hel/' })
                    .catch(function (err) { console.warn('[PWA-200] SW-Reg fehlgeschlagen:', err); });
            });
        }
    </script>
</body>
</html>
"""

def load_json_or_empty(path):
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {} if path == DIALECT_PATH else []

@router.get("/", response_class=HTMLResponse)
async def hel_interface():
    return _sanitize_unicode(ADMIN_HTML)

# Admin-Endpunkte
@router.get("/admin/models")
async def get_models():
    """Gibt verfügbare OpenRouter-Modelle als Array zurück."""
    api_key = os.getenv("OPENROUTER_API_KEY", "")
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.get(
                "https://openrouter.ai/api/v1/models",
                headers={"Authorization": f"Bearer {api_key}"},
            )
            resp.raise_for_status()
            data = resp.json()
            return data.get("data", [])
    except httpx.HTTPStatusError as e:
        raise HTTPException(
            status_code=502,
            detail=f"OpenRouter nicht erreichbar: {e.response.status_code}",
        )
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"OpenRouter nicht erreichbar: {e}")


@router.get("/admin/balance")
async def get_balance():
    """Gibt OpenRouter-Guthaben, letzte Anfrage-Kosten und Heute-Summe zurück.

    Patch 136: `last_cost_usd` und `today_total_usd` sind tatsächliche USD-Beträge
    (NICHT pro 1M Tokens, wie es die alte Frontend-Anzeige interpretiert hatte).
    Zusätzlich wird der EUR-Betrag über einen statischen Umrechnungskurs
    berechnet (keine externe API-Abfrage — Wechselkurse sind für diese Zwecke
    ausreichend approximativ).
    """
    api_key = os.getenv("OPENROUTER_API_KEY", "")
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.get(
                "https://openrouter.ai/api/v1/auth/key",
                headers={"Authorization": f"Bearer {api_key}"},
            )
            resp.raise_for_status()
            data = resp.json().get("data", {})
        last_cost = await get_last_cost()
        limit_usd = data.get("limit_usd")
        usage = data.get("usage") or 0.0
        balance = float(limit_usd) - float(usage) if limit_usd is not None else None

        # Heute gesamt (Patch 136)
        today_total = await _get_today_total_cost()

        # EUR-Umrechnung (statischer Kurs, Stand 2026-04: 1 USD ≈ 0.92 EUR)
        usd_to_eur = 0.92
        return {
            "balance": balance,
            "balance_eur": round(balance * usd_to_eur, 4) if balance is not None else None,
            "last_cost": last_cost,  # Alias (backward-compat)
            "last_cost_usd": last_cost,
            "last_cost_eur": round(last_cost * usd_to_eur, 6),
            "today_total_usd": today_total,
            "today_total_eur": round(today_total * usd_to_eur, 4),
            "fx_usd_to_eur": usd_to_eur,
        }
    except httpx.HTTPStatusError as e:
        # Patch 136: Bei fehlendem API-Key oder Netzfehler Zero-Struktur zurückgeben
        # statt 502 — Frontend soll Kosten trotzdem anzeigen (aus lokaler DB)
        last_cost = 0.0
        today_total = 0.0
        try:
            last_cost = await get_last_cost()
            today_total = await _get_today_total_cost()
        except Exception:
            pass
        return {
            "balance": None,
            "balance_eur": None,
            "last_cost": last_cost,
            "last_cost_usd": last_cost,
            "last_cost_eur": round(last_cost * 0.92, 6),
            "today_total_usd": today_total,
            "today_total_eur": round(today_total * 0.92, 4),
            "fx_usd_to_eur": 0.92,
            "error": f"openrouter_http_{e.response.status_code}",
        }
    except Exception as e:
        last_cost = 0.0
        today_total = 0.0
        try:
            last_cost = await get_last_cost()
            today_total = await _get_today_total_cost()
        except Exception:
            pass
        return {
            "balance": None,
            "balance_eur": None,
            "last_cost": last_cost,
            "last_cost_usd": last_cost,
            "last_cost_eur": round(last_cost * 0.92, 6),
            "today_total_usd": today_total,
            "today_total_eur": round(today_total * 0.92, 4),
            "fx_usd_to_eur": 0.92,
            "error": f"openrouter_unreachable: {type(e).__name__}",
        }


async def _get_today_total_cost() -> float:
    """Patch 136: Summiert alle Kosten-Einträge ab 00:00 des aktuellen Tages."""
    from zerberus.core.database import _async_session_maker
    from sqlalchemy import text as sa_text
    if _async_session_maker is None:
        return 0.0
    try:
        async with _async_session_maker() as session:
            row = (await session.execute(sa_text(
                "SELECT COALESCE(SUM(cost), 0) FROM costs "
                "WHERE timestamp >= date('now', 'start of day')"
            ))).scalar()
            return float(row or 0.0)
    except Exception:
        return 0.0

@router.get("/admin/huginn/config")
async def get_huginn_config():
    """Patch 123: Liefert die Telegram/Huginn-Konfiguration fuer den Hel-Tab.
    Der bot_token wird maskiert zurueckgegeben."""
    settings = get_settings()
    mod_cfg = dict(settings.modules.get("telegram", {}) or {})
    # Token maskieren - nie im Klartext ans Frontend
    token = mod_cfg.get("bot_token", "") or ""
    if token and token != "YOUR_TELEGRAM_BOT_TOKEN":
        masked = f"{token[:4]}…{token[-4:]}" if len(token) > 8 else "***"
    else:
        masked = ""
    # Patch 158: system_prompt (Persona). Wenn nicht gesetzt → Default.
    # Leerer String bleibt leer (User kann Persona bewusst entfernen).
    raw_prompt = mod_cfg.get("system_prompt")
    if raw_prompt is None:
        system_prompt = DEFAULT_HUGINN_PROMPT
    else:
        system_prompt = str(raw_prompt)
    return {
        "enabled": bool(mod_cfg.get("enabled", False)),
        "bot_token_masked": masked,
        "admin_chat_id": mod_cfg.get("admin_chat_id", ""),
        "allowed_group_ids": mod_cfg.get("allowed_group_ids", []),
        "model": mod_cfg.get("model", "deepseek/deepseek-chat"),
        "max_response_length": mod_cfg.get("max_response_length", 4000),
        "mode": mod_cfg.get("mode", "polling"),
        "system_prompt": system_prompt,
        "default_system_prompt": DEFAULT_HUGINN_PROMPT,
        "group_behavior": mod_cfg.get("group_behavior", {}),
        "hitl": mod_cfg.get("hitl", {}),
        # Patch 181: Allowlist-Modus + erlaubte User-IDs.
        "allowlist_mode": mod_cfg.get("allowlist_mode", "open"),
        "allowed_users": mod_cfg.get("allowed_users", []),
    }


@router.post("/admin/huginn/config")
@invalidates_settings  # Patch 156: Cache nach YAML-Write neu laden
async def post_huginn_config(request: Request):
    """Patch 123: Speichert Huginn-Config in config.yaml.
    Akzeptiert nur definierte Felder - kein blindes YAML-Update."""
    import yaml as _yaml
    data = await request.json()
    config_path = Path("config.yaml")
    if not config_path.exists():
        raise HTTPException(404, "config.yaml nicht gefunden")

    with open(config_path, "r", encoding="utf-8") as f:
        cfg = _yaml.safe_load(f) or {}

    modules = cfg.setdefault("modules", {})
    tg = modules.setdefault("telegram", {})

    # Nur erlaubte Felder durchreichen
    # Patch 158: mode + system_prompt (Persona) dazu.
    # Patch 181: allowlist_mode + allowed_users dazu.
    for key in (
        "enabled", "admin_chat_id", "allowed_group_ids",
        "model", "max_response_length", "mode", "system_prompt",
    ):
        if key in data:
            tg[key] = data[key]
    if "allowlist_mode" in data:
        mode_val = str(data["allowlist_mode"] or "open").strip().lower()
        if mode_val in {"open", "allowlist", "admin_only"}:
            tg["allowlist_mode"] = mode_val
    if "allowed_users" in data:
        # Akzeptiert Liste oder kommaseparierten String. Filter auf Ints.
        raw_list = data["allowed_users"]
        if isinstance(raw_list, str):
            raw_list = [p.strip() for p in raw_list.split(",") if p.strip()]
        elif not isinstance(raw_list, list):
            raw_list = []
        cleaned: list[int] = []
        for item in raw_list:
            try:
                cleaned.append(int(str(item).strip()))
            except (TypeError, ValueError):
                continue
        tg["allowed_users"] = cleaned
    if "bot_token" in data and data["bot_token"]:
        # Token nur schreiben wenn explizit gesetzt (nicht bei maskierten "…")
        if "…" not in str(data["bot_token"]):
            tg["bot_token"] = data["bot_token"]
    if "group_behavior" in data and isinstance(data["group_behavior"], dict):
        gb = tg.setdefault("group_behavior", {})
        for k in (
            "respond_to_name", "respond_to_mention", "respond_to_direct_reply",
            "autonomous_interjection", "interjection_cooldown_seconds",
            "interjection_trigger",
        ):
            if k in data["group_behavior"]:
                gb[k] = data["group_behavior"][k]
    if "hitl" in data and isinstance(data["hitl"], dict):
        hitl = tg.setdefault("hitl", {})
        for k in ("code_execution", "group_join", "confirmation_timeout_seconds"):
            if k in data["hitl"]:
                hitl[k] = data["hitl"][k]

    temp_fd, temp_path = tempfile.mkstemp(dir=config_path.parent, suffix=".tmp")
    try:
        with os.fdopen(temp_fd, "w", encoding="utf-8") as f:
            _yaml.safe_dump(cfg, f, allow_unicode=True, sort_keys=False)
        os.replace(temp_path, config_path)
    except Exception as e:
        try:
            os.unlink(temp_path)
        except OSError:
            pass
        raise HTTPException(500, f"Speichern fehlgeschlagen: {e}")

    return {"status": "ok"}


# ═══════════════════════════════════════════════════════════════════════════
#  Patch 134: DB-Deduplizierung (Overnight-Job + manueller Trigger)
# ═══════════════════════════════════════════════════════════════════════════

@router.post("/admin/dedup/scan")
async def post_dedup_scan():
    """Dry-Run: scannt, meldet nur — kein Schreibzugriff."""
    from zerberus.utils.db_dedup import deduplicate_interactions
    result = await deduplicate_interactions(dry_run=True, do_backup=False)
    return result


@router.post("/admin/dedup/execute")
async def post_dedup_execute():
    """Execute: scannt + soft-deletet Duplikate. Backup wird automatisch erzeugt."""
    from zerberus.utils.db_dedup import deduplicate_interactions
    result = await deduplicate_interactions(dry_run=False, do_backup=True)
    return result


# ═══════════════════════════════════════════════════════════════════════════
#  Patch 132: Memory-Store-Endpoints (strukturierter Store, ergänzt FAISS)
# ═══════════════════════════════════════════════════════════════════════════

@router.get("/admin/memory/list")
async def get_memory_list(category: str | None = None, limit: int = 100):
    """Gibt die Liste der aktiven Memories zurück, optional nach Kategorie gefiltert."""
    from zerberus.core.database import _async_session_maker
    from sqlalchemy import text as sa_text

    if _async_session_maker is None:
        return {"memories": [], "note": "db_not_initialized"}

    where_parts = ["is_active = 1"]
    params: dict = {"limit": int(max(1, min(500, limit)))}
    if category:
        where_parts.append("category = :cat")
        params["cat"] = category.lower()

    sql = (
        f"SELECT id, category, fact, confidence, source_tag, extracted_at "
        f"FROM memories WHERE {' AND '.join(where_parts)} "
        f"ORDER BY extracted_at DESC LIMIT :limit"
    )
    try:
        async with _async_session_maker() as session:
            rows = (await session.execute(sa_text(sql), params)).fetchall()
    except Exception as e:
        return {"memories": [], "error": str(e)[:200]}

    return {
        "memories": [
            {
                "id": r[0],
                "category": r[1],
                "fact": r[2],
                "confidence": r[3],
                "source_tag": r[4],
                "extracted_at": str(r[5]) if r[5] else None,
            }
            for r in rows
        ],
        "count": len(rows),
    }


@router.get("/admin/memory/stats")
async def get_memory_stats():
    """Aggregierte Zahlen: Gesamtzahl + pro Kategorie + letzte Extraktion."""
    from zerberus.core.database import _async_session_maker
    from sqlalchemy import text as sa_text

    if _async_session_maker is None:
        return {"total": 0, "by_category": {}}
    try:
        async with _async_session_maker() as session:
            total = (await session.execute(sa_text(
                "SELECT COUNT(*) FROM memories WHERE is_active = 1"
            ))).scalar() or 0
            rows = (await session.execute(sa_text(
                "SELECT category, COUNT(*) FROM memories WHERE is_active = 1 GROUP BY category"
            ))).fetchall()
            last = (await session.execute(sa_text(
                "SELECT MAX(extracted_at) FROM memories WHERE is_active = 1"
            ))).scalar()
    except Exception as e:
        return {"total": 0, "by_category": {}, "error": str(e)[:200]}

    return {
        "total": int(total),
        "by_category": {str(r[0]): int(r[1]) for r in rows},
        "last_extraction": str(last) if last else None,
    }


@router.post("/admin/memory/add")
async def post_memory_add(request: Request):
    """Manueller Memory-Insert (nicht vom Extractor). Nutzt nur den Structured-Store."""
    from zerberus.core.database import _async_session_maker
    from sqlalchemy import text as sa_text

    data = await request.json()
    category = str(data.get("category", "personal")).strip().lower()
    fact = str(data.get("fact", "")).strip()
    confidence = float(data.get("confidence", 1.0))
    if not fact:
        raise HTTPException(400, "fact darf nicht leer sein")
    if _async_session_maker is None:
        raise HTTPException(503, "db_not_initialized")

    try:
        async with _async_session_maker() as session:
            result = await session.execute(sa_text(
                "INSERT INTO memories "
                "(category, fact, confidence, source_tag, extracted_at, is_active) "
                "VALUES (:cat, :fact, :conf, 'manual', datetime('now'), 1)"
            ), {"cat": category, "fact": fact, "conf": confidence})
            await session.commit()
            new_id = int(result.lastrowid) if hasattr(result, "lastrowid") else None
    except Exception as e:
        raise HTTPException(500, f"insert failed: {e}")

    return {"status": "ok", "id": new_id}


@router.delete("/admin/memory/{memory_id}")
async def delete_memory(memory_id: int):
    """Soft-Delete: setzt is_active=0. Physischer Eintrag bleibt."""
    from zerberus.core.database import _async_session_maker
    from sqlalchemy import text as sa_text

    if _async_session_maker is None:
        raise HTTPException(503, "db_not_initialized")
    try:
        async with _async_session_maker() as session:
            await session.execute(sa_text(
                "UPDATE memories SET is_active = 0 WHERE id = :id"
            ), {"id": int(memory_id)})
            await session.commit()
    except Exception as e:
        raise HTTPException(500, f"delete failed: {e}")
    return {"status": "ok", "id": memory_id}


# ═══════════════════════════════════════════════════════════════════════════
#  Patch 194 (Phase 5a #1): Projekt-CRUD
#  Admin-only via verify_admin (Hel Basic-Auth). Tabellen liegen in
#  bunker_memory.db (Decision 1, 2026-05-01). Hel-first vor Nala-Integration
#  (Decision 2). Persona-Overlay als Merge-Layer (Decision 3).
# ═══════════════════════════════════════════════════════════════════════════


@router.get("/admin/projects")
async def list_projects_endpoint(include_archived: bool = False):
    from zerberus.core import projects_repo
    items = await projects_repo.list_projects(include_archived=include_archived)
    return {"projects": items, "count": len(items)}


@router.post("/admin/projects")
async def create_project_endpoint(request: Request):
    from zerberus.core import projects_repo, projects_template
    data = await request.json()
    name = str(data.get("name", "")).strip()
    if not name:
        raise HTTPException(400, "name darf nicht leer sein")
    description = str(data.get("description", "")).strip()
    persona_overlay = data.get("persona_overlay")
    slug_override = data.get("slug")
    if persona_overlay is not None and not isinstance(persona_overlay, dict):
        raise HTTPException(400, "persona_overlay muss ein Objekt sein")
    try:
        project = await projects_repo.create_project(
            name=name,
            description=description,
            persona_overlay=persona_overlay,
            slug=slug_override,
        )
    except ValueError as e:
        raise HTTPException(400, str(e))

    # Patch 198: Skelett-Files (ZERBERUS_<SLUG>.md + README.md) im Storage
    # anlegen + als project_files registrieren. Idempotent — vorhandene
    # Pfade werden nicht ueberschrieben. Best-Effort: Fehler beim
    # Materialisieren brechen das Anlegen NICHT ab (Projekt-Eintrag steht,
    # Templates lassen sich notfalls nachgenerieren).
    settings = get_settings()
    template_files: list[dict] = []
    if settings.projects.auto_template:
        try:
            template_files = await projects_template.materialize_template(
                project, _projects_storage_base()
            )
        except Exception as e:
            logger.exception(f"[TEMPLATE-198] materialize failed for slug={project['slug']}: {e}")
    return {"status": "ok", "project": project, "template_files": template_files}


@router.get("/admin/projects/{project_id}")
async def get_project_endpoint(project_id: int):
    from zerberus.core import projects_repo
    project = await projects_repo.get_project(project_id)
    if project is None:
        raise HTTPException(404, "Projekt nicht gefunden")
    return {"project": project}


@router.patch("/admin/projects/{project_id}")
async def update_project_endpoint(project_id: int, request: Request):
    from zerberus.core import projects_repo
    data = await request.json()
    kwargs: dict = {}
    if "name" in data:
        kwargs["name"] = str(data["name"])
    if "description" in data:
        kwargs["description"] = str(data["description"])
    if "persona_overlay" in data:
        overlay = data["persona_overlay"]
        if overlay is not None and not isinstance(overlay, dict):
            raise HTTPException(400, "persona_overlay muss ein Objekt oder null sein")
        kwargs["persona_overlay"] = overlay
    try:
        project = await projects_repo.update_project(project_id, **kwargs)
    except ValueError as e:
        raise HTTPException(400, str(e))
    if project is None:
        raise HTTPException(404, "Projekt nicht gefunden")
    return {"status": "ok", "project": project}


@router.post("/admin/projects/{project_id}/archive")
async def archive_project_endpoint(project_id: int):
    from zerberus.core import projects_repo
    project = await projects_repo.archive_project(project_id)
    if project is None:
        raise HTTPException(404, "Projekt nicht gefunden")
    return {"status": "ok", "project": project}


@router.post("/admin/projects/{project_id}/unarchive")
async def unarchive_project_endpoint(project_id: int):
    from zerberus.core import projects_repo
    project = await projects_repo.unarchive_project(project_id)
    if project is None:
        raise HTTPException(404, "Projekt nicht gefunden")
    return {"status": "ok", "project": project}


@router.delete("/admin/projects/{project_id}")
async def delete_project_endpoint(project_id: int):
    """Harte Loeschung — kaskadiert ueber project_files. Nicht reversibel."""
    from zerberus.core import projects_repo
    # Patch 199 (Phase 5a #3): Slug VOR dem Delete merken, damit wir den
    # RAG-Index-Ordner danach gezielt entfernen koennen (ohne Slug kein Pfad).
    project_pre = await projects_repo.get_project(project_id)
    deleted = await projects_repo.delete_project(project_id)
    if not deleted:
        raise HTTPException(404, "Projekt nicht gefunden")
    if project_pre and get_settings().projects.rag_enabled:
        try:
            from zerberus.core import projects_rag

            projects_rag.remove_project_index(project_pre["slug"], _projects_storage_base())
        except Exception as rag_err:
            logger.exception(
                f"[RAG-199] remove_project_index failed slug={project_pre['slug']}: {rag_err}"
            )
    # Patch 203a: Workspace-Ordner mitkippen — wipe_workspace ist
    # idempotent, falls workspace_enabled erst spaeter aktiviert wurde
    # und der Ordner gar nicht existiert. Sicherheits-Check (Pfad endet
    # auf ``_workspace``) liegt im Helper.
    if project_pre and get_settings().projects.workspace_enabled:
        try:
            from zerberus.core import projects_workspace

            projects_workspace.wipe_workspace(
                projects_workspace.workspace_root_for(
                    project_pre["slug"], _projects_storage_base()
                )
            )
        except Exception as ws_err:
            logger.exception(
                f"[WORKSPACE-203] wipe_workspace failed slug={project_pre['slug']}: {ws_err}"
            )
    return {"status": "ok"}


@router.get("/admin/projects/{project_id}/files")
async def list_project_files_endpoint(project_id: int):
    from zerberus.core import projects_repo
    project = await projects_repo.get_project(project_id)
    if project is None:
        raise HTTPException(404, "Projekt nicht gefunden")
    files = await projects_repo.list_files(project_id)
    return {"files": files, "count": len(files)}


# ═══════════════════════════════════════════════════════════════════════════
#  Patch 196 (Phase 5a #4): Datei-Upload-Endpoints
#  Multipart-Upload + Delete mit SHA-Dedup. Bytes liegen unter
#  ``<projects.data_dir>/projects/<slug>/<sha[:2]>/<sha>`` (Konvention aus
#  P194). Extension-Blacklist + Max-Size kommen aus ProjectsConfig.
# ═══════════════════════════════════════════════════════════════════════════


def _projects_storage_base() -> Path:
    """Storage-Wurzel aus den Settings — als Path. Caller liest erst hier,
    damit ``data_dir`` zur Laufzeit umgebogen werden kann (Tests setzen
    monkeypatch auf diese Funktion statt auf die Settings).
    """
    settings = get_settings()
    return Path(settings.projects.data_dir)


async def _store_uploaded_bytes(target: Path, data: bytes) -> None:
    """Schreibt ``data`` atomar nach ``target`` — erst tempfile in den
    Ziel-Ordner, dann ``os.replace``. Verhindert halb-geschriebene Files
    bei Crash mitten im Schreiben (z.B. Server-Kill).
    """
    target.parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(dir=str(target.parent), prefix=".upload_", suffix=".tmp")
    try:
        with os.fdopen(fd, "wb") as fh:
            fh.write(data)
        os.replace(tmp_path, str(target))
    except Exception:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
        raise


def _cleanup_storage_path(path: Path, base: Path) -> None:
    """Loescht Bytes + leere Parent-Ordner bis zum ``base``-Anker.
    Best-Effort — Fehler werden geschluckt, weil das Repo-DB-Delete bereits
    durch ist und der Aufrufer einen 200 sehen soll.
    """
    try:
        if path.exists():
            path.unlink()
    except OSError:
        return
    parent = path.parent
    try:
        base_resolved = base.resolve()
    except OSError:
        return
    while True:
        try:
            if parent.resolve() == base_resolved:
                return
            if not parent.exists():
                return
            if any(parent.iterdir()):
                return
            parent.rmdir()
            parent = parent.parent
        except OSError:
            return


@router.post("/admin/projects/{project_id}/files")
async def upload_project_file_endpoint(project_id: int, file: UploadFile = File(...)):
    """Multipart-Upload. SHA256-Dedup: existiert derselbe Inhalt schon im
    Storage, wird nur der Metadaten-Eintrag neu angelegt — die Bytes
    werden nicht doppelt geschrieben.
    """
    from zerberus.core import projects_repo

    project = await projects_repo.get_project(project_id)
    if project is None:
        raise HTTPException(404, "Projekt nicht gefunden")

    raw_filename = file.filename or ""
    try:
        relative_path = projects_repo.sanitize_relative_path(raw_filename)
    except ValueError as e:
        raise HTTPException(400, str(e))

    settings = get_settings()
    blocked = settings.projects.blocked_extensions
    if projects_repo.is_extension_blocked(relative_path, blocked):
        raise HTTPException(400, f"Datei-Endung blockiert: {relative_path}")

    max_bytes = settings.projects.max_upload_bytes
    data = await file.read()
    if len(data) == 0:
        raise HTTPException(400, "Leere Datei abgewiesen")
    if len(data) > max_bytes:
        raise HTTPException(
            413,
            f"Datei zu gross ({len(data)} bytes > {max_bytes} bytes Limit)",
        )

    sha256 = projects_repo.compute_sha256(data)
    base = _projects_storage_base()
    target = projects_repo.storage_path_for(project["slug"], sha256, base)

    if not target.exists():
        await _store_uploaded_bytes(target, data)

    try:
        registered = await projects_repo.register_file(
            project_id=project_id,
            relative_path=relative_path,
            sha256=sha256,
            size_bytes=len(data),
            storage_path=str(target),
            mime_type=file.content_type or None,
        )
    except Exception as e:
        # UNIQUE-Verletzung (selber Pfad bereits im Projekt) → 409
        msg = str(e).lower()
        if "unique" in msg or "integrity" in msg:
            raise HTTPException(409, f"Datei mit Pfad '{relative_path}' existiert bereits")
        raise

    # Patch 199 (Phase 5a #3): Datei in den Projekt-RAG-Index aufnehmen.
    # Best-Effort: Indexing-Fehler brechen den Upload nicht ab — die Datei
    # ist in der DB + im Storage, nur der Index-Eintrag fehlt. Kann via
    # Re-Upload oder kuenftigem Reindex-CLI nachgezogen werden.
    rag_status: dict = {"chunks": 0, "skipped": True, "reason": "rag_disabled"}
    if settings.projects.rag_enabled:
        try:
            from zerberus.core import projects_rag

            rag_status = await projects_rag.index_project_file(
                project_id=project_id,
                file_id=registered["id"],
                base_dir=base,
            )
        except Exception as rag_err:
            logger.exception(
                f"[RAG-199] index_project_file failed slug={project['slug']} "
                f"file_id={registered['id']}: {rag_err}"
            )
            rag_status = {"chunks": 0, "skipped": True, "reason": "exception"}

    # Patch 203a (Phase 5a #5, Vorbereitung): Datei in den Projekt-Workspace
    # spiegeln (Hardlink mit Copy-Fallback). Vorbereitung fuer die Code-
    # Execution-Pipeline. Best-Effort: Workspace-Fehler brechen den Upload
    # nicht ab — Source-of-Truth ist der SHA-Storage + DB.
    if settings.projects.workspace_enabled:
        try:
            from zerberus.core import projects_workspace

            await projects_workspace.materialize_file_async(
                project_id=project_id,
                file_id=registered["id"],
                base_dir=base,
            )
        except Exception as ws_err:
            logger.exception(
                f"[WORKSPACE-203] materialize_file_async failed slug={project['slug']} "
                f"file_id={registered['id']}: {ws_err}"
            )

    return {"status": "ok", "file": registered, "rag": rag_status}


@router.delete("/admin/projects/{project_id}/files/{file_id}")
async def delete_project_file_endpoint(project_id: int, file_id: int):
    """Loescht Metadaten-Eintrag und ggf. Bytes. Bytes bleiben erhalten,
    wenn ein anderer ``project_files``-Eintrag (auch in einem anderen
    Projekt) denselben sha256 referenziert.
    """
    from zerberus.core import projects_repo

    file_meta = await projects_repo.get_file(file_id)
    if file_meta is None or file_meta["project_id"] != project_id:
        raise HTTPException(404, "Datei nicht gefunden")

    # Patch 203a: Slug VOR dem DB-Delete merken — wird unten fuer den
    # Workspace-Remove gebraucht (Workspace-Pfad ist slug-keyed).
    project_pre = await projects_repo.get_project(project_id)

    other_refs = await projects_repo.count_sha_references(
        file_meta["sha256"], exclude_file_id=file_id
    )
    deleted = await projects_repo.delete_file(file_id)
    if not deleted:
        raise HTTPException(404, "Datei nicht gefunden")

    if other_refs == 0 and file_meta.get("storage_path"):
        base = _projects_storage_base()
        _cleanup_storage_path(Path(file_meta["storage_path"]), base)

    # Patch 199 (Phase 5a #3): Index-Eintraege fuer diese Datei rauswerfen,
    # damit der RAG-Block im Chat keinen "stale"-Treffer liefert. Best-
    # Effort: Fehler werden geloggt, aber der Delete-Status bleibt 200 —
    # der DB-Eintrag ist schon weg.
    rag_removed = 0
    settings = get_settings()
    if settings.projects.rag_enabled:
        try:
            from zerberus.core import projects_rag

            rag_removed = await projects_rag.remove_file_from_index(
                project_id=project_id,
                file_id=file_id,
                base_dir=_projects_storage_base(),
            )
        except Exception as rag_err:
            logger.exception(
                f"[RAG-199] remove_file_from_index failed project_id={project_id} "
                f"file_id={file_id}: {rag_err}"
            )

    # Patch 203a: Workspace-Eintrag entfernen, damit der naechste Sandbox-
    # Run keine geloeschte Datei mehr sieht. Best-Effort: Fehler werden
    # geloggt, Status bleibt 200.
    if settings.projects.workspace_enabled and project_pre is not None:
        try:
            from zerberus.core import projects_workspace

            await projects_workspace.remove_file_async(
                project_slug=project_pre["slug"],
                relative_path=file_meta["relative_path"],
                base_dir=_projects_storage_base(),
            )
        except Exception as ws_err:
            logger.exception(
                f"[WORKSPACE-203] remove_file_async failed project_id={project_id} "
                f"file_id={file_id}: {ws_err}"
            )

    return {
        "status": "ok",
        "bytes_removed": other_refs == 0,
        "rag_chunks_removed": rag_removed,
    }


# ═══════════════════════════════════════════════════════════════════════════
#  Patch 131: Vision-Modell-Endpoints (nur Vision-fähige Modelle)
# ═══════════════════════════════════════════════════════════════════════════

@router.get("/admin/vision/models")
async def get_vision_models_list():
    """Gibt die gefilterte Liste der Vision-fähigen Modelle zurück."""
    from zerberus.core.vision_models import get_vision_models
    return {"models": get_vision_models()}


@router.get("/admin/vision/config")
async def get_vision_config_endpoint():
    """Liefert den aktuellen `vision:`-Block aus config.yaml."""
    import yaml as _yaml
    config_path = Path("config.yaml")
    if not config_path.exists():
        return {"enabled": False, "model": "", "max_image_size_mb": 10}
    with open(config_path, "r", encoding="utf-8") as f:
        cfg = _yaml.safe_load(f) or {}
    vision = cfg.get("vision", {}) or {}
    return {
        "enabled": bool(vision.get("enabled", True)),
        "model": vision.get("model", "qwen/qwen2.5-vl-7b-instruct"),
        "max_image_size_mb": int(vision.get("max_image_size_mb", 10)),
        "supported_formats": vision.get("supported_formats", ["jpg", "jpeg", "png", "gif", "webp"]),
    }


@router.post("/admin/vision/config")
@invalidates_settings  # Patch 156: Cache nach YAML-Write neu laden
async def post_vision_config(request: Request):
    """Speichert `vision:`-Block in config.yaml. Nur Vision-Modelle akzeptiert."""
    import yaml as _yaml
    from zerberus.core.vision_models import is_vision_model

    data = await request.json()
    model = data.get("model", "")
    if model and not is_vision_model(model):
        raise HTTPException(400, f"Modell '{model}' ist nicht in der Vision-Registry")

    config_path = Path("config.yaml")
    if not config_path.exists():
        raise HTTPException(404, "config.yaml nicht gefunden")

    with open(config_path, "r", encoding="utf-8") as f:
        cfg = _yaml.safe_load(f) or {}

    vision = cfg.setdefault("vision", {})
    for key in ("enabled", "model", "max_image_size_mb", "supported_formats"):
        if key in data:
            vision[key] = data[key]

    temp_fd, temp_path = tempfile.mkstemp(dir=config_path.parent, suffix=".tmp")
    try:
        with os.fdopen(temp_fd, "w", encoding="utf-8") as f:
            _yaml.safe_dump(cfg, f, allow_unicode=True, sort_keys=False)
        os.replace(temp_path, config_path)
    except Exception as e:
        try:
            os.unlink(temp_path)
        except OSError:
            pass
        raise HTTPException(500, f"Speichern fehlgeschlagen: {e}")

    return {"status": "ok"}


@router.get("/admin/whisper_cleaner")
async def get_whisper_cleaner():
    return JSONResponse(content=load_json_or_empty(WHISPER_CLEANER_PATH))

@router.post("/admin/whisper_cleaner")
async def post_whisper_cleaner(request: Request):
    data = await request.json()
    temp_fd, temp_path = tempfile.mkstemp(dir=WHISPER_CLEANER_PATH.parent, suffix=".tmp")
    try:
        with os.fdopen(temp_fd, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        os.replace(temp_path, WHISPER_CLEANER_PATH)
    except Exception as e:
        os.unlink(temp_path)
        raise e
    return {"status": "ok"}

@router.get("/admin/fuzzy_dict")
async def get_fuzzy_dict():
    return JSONResponse(content=load_json_or_empty(FUZZY_DICT_PATH))

@router.post("/admin/fuzzy_dict")
async def post_fuzzy_dict(request: Request):
    data = await request.json()
    if not isinstance(data, list):
        raise HTTPException(status_code=400, detail="fuzzy_dictionary.json muss ein JSON-Array sein")
    temp_fd, temp_path = tempfile.mkstemp(dir=FUZZY_DICT_PATH.parent, suffix=".tmp")
    try:
        with os.fdopen(temp_fd, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        os.replace(temp_path, FUZZY_DICT_PATH)
    except Exception as e:
        os.unlink(temp_path)
        raise e
    return {"status": "ok"}

@router.get("/admin/dialect")
async def get_dialect():
    return JSONResponse(content=load_json_or_empty(DIALECT_PATH))

@router.post("/admin/dialect")
async def post_dialect(request: Request):
    data = await request.json()
    temp_fd, temp_path = tempfile.mkstemp(dir=DIALECT_PATH.parent, suffix=".tmp")
    try:
        with os.fdopen(temp_fd, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        os.replace(temp_path, DIALECT_PATH)
    except Exception as e:
        os.unlink(temp_path)
        raise e
    return {"status": "ok"}

YAML_CONFIG_PATH = Path("config.yaml")


def _yaml_replace_scalar(section_path: list[str], key: str, new_value) -> bool:
    """Patch 105: Ersetzt einen Skalar-Wert in config.yaml in-place.

    Einfache Line-basierte Ersetzung, die Kommentare und Formatierung
    erhält — yaml.safe_dump würde alles neu serialisieren und Kommentare
    zerstören. Findet den Key unter dem vorgegebenen Section-Pfad
    (z.B. ['legacy', 'models']) und ersetzt nur die eine Zeile.

    Gibt True zurück wenn ersetzt, False wenn Key/Section nicht gefunden.
    """
    if not YAML_CONFIG_PATH.exists():
        return False
    lines = YAML_CONFIG_PATH.read_text(encoding="utf-8").splitlines(keepends=True)
    indent_stack: list[tuple[str, int]] = []
    target_depth = len(section_path)
    for i, raw in enumerate(lines):
        stripped = raw.lstrip(" ")
        if not stripped or stripped.startswith("#"):
            continue
        indent = len(raw) - len(stripped)
        while indent_stack and indent_stack[-1][1] >= indent:
            indent_stack.pop()
        if ":" not in stripped:
            continue
        line_key = stripped.split(":", 1)[0].strip()
        current_path = [k for k, _ in indent_stack]
        if current_path == section_path and line_key == key:
            prefix = " " * indent
            temp_fd, temp_path = tempfile.mkstemp(dir=YAML_CONFIG_PATH.parent, suffix=".tmp")
            lines[i] = f"{prefix}{key}: {new_value}\n"
            try:
                with os.fdopen(temp_fd, "w", encoding="utf-8") as f:
                    f.writelines(lines)
                os.replace(temp_path, YAML_CONFIG_PATH)
            except Exception:
                try:
                    os.unlink(temp_path)
                except OSError:
                    pass
                raise
            return True
        after_colon = stripped.split(":", 1)[1].strip()
        if not after_colon or after_colon.startswith("#"):
            indent_stack.append((line_key, indent))
    return False


@router.get("/admin/config")
async def get_config():
    """Patch 105: Liest aus config.yaml (Single Source of Truth).

    Antwortform `{"llm": {...}}` bleibt abwärtskompatibel mit dem Hel-UI-JS
    (admin/hel.py ~L992-1017), das `config.llm?.cloud_model` auswertet.
    """
    settings = get_settings()
    return JSONResponse(content={
        "llm": {
            "cloud_model": settings.legacy.models.cloud_model,
            "temperature": settings.legacy.settings.ai_temperature,
            "threshold": settings.legacy.settings.threshold_length,
        }
    })


@router.post("/admin/config")
async def post_config(request: Request):
    """Patch 105: Schreibt nach config.yaml — nicht mehr in config.json.

    Hintergrund: vorheriger Split-Brain (Hel schrieb config.json, LLMService
    las config.yaml) führte dazu, dass in Hel ausgewählte Modelle nie
    wirksam wurden. Kommentar-erhaltendes Line-Replace, damit die
    handgepflegten Patch-Notizen in config.yaml nicht verloren gehen.
    """
    data = await request.json()
    llm = data.get("llm", {}) if isinstance(data, dict) else {}
    changed: list[str] = []
    if "cloud_model" in llm:
        if _yaml_replace_scalar(["legacy", "models"], "cloud_model", str(llm["cloud_model"])):
            changed.append("cloud_model")
    if "temperature" in llm:
        try:
            temp_val = float(llm["temperature"])
        except (TypeError, ValueError):
            raise HTTPException(status_code=400, detail="temperature must be a float")
        if _yaml_replace_scalar(["legacy", "settings"], "ai_temperature", temp_val):
            changed.append("temperature")
    if "threshold" in llm:
        try:
            thr_val = int(llm["threshold"])
        except (TypeError, ValueError):
            raise HTTPException(status_code=400, detail="threshold must be an int")
        if _yaml_replace_scalar(["legacy", "settings"], "threshold_length", thr_val):
            changed.append("threshold")
    if changed:
        reload_settings()
    return {"status": "ok", "reloaded": bool(changed), "changed": changed}

@router.get("/admin/sessions")
async def get_sessions():
    sessions = await get_all_sessions(limit=50)
    return sessions

@router.get("/admin/export/session/{session_id}")
async def export_session(session_id: str):
    messages = await get_session_messages(session_id)
    if not messages:
        raise HTTPException(404, "Session nicht gefunden")
    return messages

@router.delete("/admin/session/{session_id}")
async def delete_session_admin(session_id: str):
    from zerberus.app.routers.archive import delete_session_endpoint as archive_delete
    return await archive_delete(session_id)

@router.get("/admin/system_prompt")
async def get_system_prompt():
    with open(SYSTEM_PROMPT_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

@router.post("/admin/system_prompt")
async def post_system_prompt(request: Request):
    data = await request.json()
    temp_fd, temp_path = tempfile.mkstemp(dir=SYSTEM_PROMPT_PATH.parent, suffix=".tmp")
    try:
        with os.fdopen(temp_fd, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        os.replace(temp_path, SYSTEM_PROMPT_PATH)
    except Exception as e:
        os.unlink(temp_path)
        raise e
    return {"status": "ok"}

@router.get("/metrics/latest_with_costs")
async def metrics_latest_with_costs(limit: int = 20, session_id: str = None):
    """
    Metriken pro Nachricht, angereichert um Kosten.

    Patch 146 (B-018): Voller Nachrichten-Text wird gekürzt auf 50 Zeichen +
    Truncation-Marker. Metriken-Tab soll Zahlen/Charts zeigen, nicht Rohtext —
    eine einzelne lange LLM-Antwort hatte zuvor mehrere Bildschirme gefüllt.
    """
    messages = await get_latest_metrics(limit, session_id)
    ids = [m['id'] for m in messages if m.get('role') == 'assistant']
    costs = await get_message_costs(ids)
    for m in messages:
        m['cost'] = costs.get(m['id'], None)
        # Patch 146 (B-018): Content auf max. 50 Zeichen truncaten.
        raw = m.get('content') or ''
        if len(raw) > 50:
            m['content'] = raw[:50] + '…'
        m['content_truncated'] = len(raw) > 50
        m['content_original_length'] = len(raw)
    return messages

@router.get("/metrics/summary")
async def metrics_summary(session_id: str = None):
    return await get_metrics_summary(session_id)


# ============================================================
# RAG-Upload (Säule 1 – Patch 56)
# ============================================================

CONFIG_YAML_PATH = Path("config.yaml")


# Kapitelgrenzen-Regex: Prolog / Akt I–VII / Epilog / Glossar am Zeilenanfang
_CHAPTER_RE = re.compile(
    r'(?m)(?=^(?:Prolog|Epilog|Glossar|Akt\s+[IVXLC]+)\b)',
    re.IGNORECASE,
)

# Patch 110: Markdown-Header-Regex (# / ## / ### / …) am Zeilenanfang.
# Für `technical`-Category: Abschnitte an Headern hart trennen.
_MD_HEADER_RE = re.compile(r'(?m)(?=^#{1,6}\s+\S)')

# Patch 110: Chunking-Konfigurationen pro Category.
# - chunk_size / overlap / min_chunk_words in Wörtern (Einheit aus Patch 75).
# - split: "chapter" = _CHAPTER_RE, "markdown" = _MD_HEADER_RE, "none" = nur Wortfenster.
# `narrative`/`lore` behalten die bestehenden 800/160/120-Werte (stabil mit
# bge-reranker-v2-m3, siehe Patch 89). Andere Categories bekommen eigene Profile.
CHUNK_CONFIGS: dict[str, dict] = {
    "narrative":  {"chunk_size": 800, "overlap": 160, "min_chunk_words": 120, "split": "chapter"},
    "lore":       {"chunk_size": 800, "overlap": 160, "min_chunk_words": 120, "split": "chapter"},
    "technical":  {"chunk_size": 500, "overlap": 100, "min_chunk_words": 80,  "split": "markdown"},
    "reference":  {"chunk_size": 300, "overlap": 60,  "min_chunk_words": 50,  "split": "none"},
    "personal":   {"chunk_size": 400, "overlap": 80,  "min_chunk_words": 80,  "split": "chapter"},
    "general":    {"chunk_size": 800, "overlap": 160, "min_chunk_words": 120, "split": "chapter"},
    # Patch 178: ``system`` ist die Selbstwissen-Kategorie fuer Huginn — Dokumente
    # ueber das Zerberus-System selbst, die Huginn ueber den RAG-Filter (nur
    # ``system``-Chunks) ausschliesslich abrufen darf. Markdown-Splits + kleinere
    # min_chunk_words als technical (50 statt 80), damit kurze, prägnante Sektionen
    # (z.B. "Was Zerberus NICHT ist", "Mythologische Namen" Tabelle) eigene Chunks
    # bekommen statt vom Residual-Merge in den Vorgänger geschluckt zu werden.
    "system":     {"chunk_size": 200, "overlap": 40,  "min_chunk_words": 30,  "split": "none"},
}


def _chunk_text(
    text: str,
    chunk_size: int | None = None,
    overlap: int | None = None,
    min_chunk_words: int | None = None,
    category: str = "general",
) -> tuple[list[str], int]:
    """Zerlegt Text in Chunks mit optional harten Splits an Kapitel-/Header-Grenzen.

    Patch 110: `category` wählt ein Profil aus `CHUNK_CONFIGS` (chunk_size, overlap,
    min_chunk_words, split-Strategie). Explizite Argumente überschreiben das Profil —
    nützlich für Tests oder config.yaml-Overrides.

    Split-Strategien:
    - "chapter" (narrative/lore/personal/general): harte Splits an Prolog/Akt/Epilog/Glossar
    - "markdown" (technical): harte Splits an Markdown-Headern (# bis ######)
    - "none" (reference): nur Wortfenster, keine harten Splits

    Kapitel-/Header-Grenzen werden nie überlappt — Overlap nur innerhalb eines Abschnitts.
    Einheit: Wörter (nicht Token, nicht Zeichen) — Patch 75.

    Patch 88 (Fix A): Post-Processing-Pass merged Residual-Chunks unter
    `min_chunk_words` in den Vorgänger-Chunk. Kurze Tails (z.B. 5w/16w/64w)
    kapern bei normalisierten MiniLM-Embeddings sonst Rang 1.
    Rückgabe: (chunks, merged_count)
    """
    cfg = CHUNK_CONFIGS.get(category, CHUNK_CONFIGS["general"])
    chunk_size = chunk_size if chunk_size is not None else cfg["chunk_size"]
    overlap = overlap if overlap is not None else cfg["overlap"]
    min_chunk_words = min_chunk_words if min_chunk_words is not None else cfg["min_chunk_words"]
    split_strategy = cfg["split"]

    if split_strategy == "chapter":
        sections = _CHAPTER_RE.split(text)
    elif split_strategy == "markdown":
        sections = _MD_HEADER_RE.split(text)
    else:
        sections = [text]
    sections = [s.strip() for s in sections if s.strip()]

    chunks: list[str] = []
    for section in sections:
        words = section.split()
        if not words:
            continue
        i = 0
        step = max(1, chunk_size - overlap)
        while i < len(words):
            chunk_words = words[i:i + chunk_size]
            chunks.append(" ".join(chunk_words))
            i += step

    # Patch 88: Residual-Tail-Merge
    merged_count = 0
    if len(chunks) > 1 and min_chunk_words > 0:
        cleaned: list[str] = []
        for chunk in chunks:
            wc = len(chunk.split())
            if wc < min_chunk_words and cleaned:
                # an Vorgänger anhängen
                cleaned[-1] = cleaned[-1] + "\n\n" + chunk
                merged_count += 1
            else:
                cleaned.append(chunk)
        # Edge-Case: Erster Chunk war zu kurz → in Nachfolger einbetten
        if cleaned and len(cleaned[0].split()) < min_chunk_words and len(cleaned) > 1:
            cleaned[1] = cleaned[0] + "\n\n" + cleaned[1]
            cleaned.pop(0)
            merged_count += 1
        chunks = cleaned

    return chunks, merged_count


# Patch 108: Erlaubte Kategorien für RAG-Upload. "general" ist Default/Fallback
# für Altdaten ohne category-Feld. Patch 111: "auto" triggert Extension-basierte
# Detection (_detect_category), landet aber nie als echte Category in der Metadata.
_RAG_CATEGORIES = {
    "general", "narrative", "technical", "personal", "lore", "reference",
    # Patch 178: ``system`` = Selbstwissen-Doku ueber Zerberus, ueber die
    # Huginn fundiert Auskunft geben darf (Category-Filter im Telegram-Router
    # laesst nur diese Kategorie an den LLM-Kontext durch).
    "system",
}


# Patch 111: Extension-basierte Auto-Detection. Entscheidet wenn
# User "auto" (neues Default) oder "general" gewählt hat. Content-basierte
# Detection (LLM) kommt in Phase 4.
_EXTENSION_CATEGORY_MAP: dict[str, str] = {
    ".py":   "technical",
    ".js":   "technical",
    ".jsx":  "technical",
    ".mjs":  "technical",
    ".cjs":  "technical",
    ".ts":   "technical",
    ".tsx":  "technical",
    ".html": "technical",
    ".htm":  "technical",
    ".css":  "technical",
    ".scss": "technical",
    ".sass": "technical",
    ".sql":  "technical",
    ".json": "technical",
    ".yaml": "technical",
    ".yml":  "technical",
    ".md":   "technical",
    ".csv":  "reference",
    ".pdf":  "general",
    ".txt":  "general",
    ".docx": "general",
}


def _detect_category(filename: str, user_category: str) -> tuple[str, bool]:
    """Auto-Detection wenn User 'auto' oder 'general' gewählt hat.

    Returns:
        (resolved_category, auto_detected) — auto_detected=True wenn die
        Detection das Ergebnis bestimmt hat (nicht User-Override).
    """
    user_category = (user_category or "auto").strip().lower()
    if user_category not in ("auto", "general"):
        return user_category, False

    suffix = Path(filename).suffix.lower()
    detected = _EXTENSION_CATEGORY_MAP.get(suffix, "general")
    logger.warning(
        f"[CAT-111] Auto-Detection: {filename} (suffix={suffix!r}) → {detected} "
        f"[user wählte {user_category!r}]"
    )
    return detected, True


@router.post("/admin/rag/upload")
async def rag_upload(
    file: UploadFile = File(...),
    category: str = Form("auto"),
):
    """Lädt .txt, .md, .docx oder .pdf hoch, zerlegt in Chunks und indiziert sie im FAISS-Index.

    Patch 108: `category` (Form-Feld) wird pro Chunk als Metadata gespeichert.
    Patch 111: `category="auto"` (neues Default) oder `"general"` triggert
    Extension-basierte Auto-Detection (siehe `_detect_category`).
    Unbekannte Werte fallen stillschweigend auf "general" zurück.
    """
    from zerberus.modules.rag.router import (
        _ensure_init, _encode, _add_to_index, RAG_AVAILABLE
    )
    settings = get_settings()
    mod_cfg = settings.modules.get("rag", {})
    if not mod_cfg.get("enabled", False):
        raise HTTPException(503, "RAG-Modul ist in config.yaml deaktiviert.")
    if not RAG_AVAILABLE:
        raise HTTPException(503, "RAG-Abhängigkeiten (faiss, sentence-transformers) nicht installiert.")

    filename = file.filename or "unbekannt"
    suffix = Path(filename).suffix.lower()
    raw_category = (category or "auto").strip().lower()

    # Patch 111: Auto-Detection bei "auto"/"general"
    category, auto_detected = _detect_category(filename, raw_category)

    if category not in _RAG_CATEGORIES:
        logger.warning(f"[RAG-108] Unbekannte Kategorie {category!r}, fallback auf 'general'")
        category = "general"
        auto_detected = False

    raw_bytes = await file.read()

    # Patch 122: Code-Dateien (.py/.js/.ts/.html/.css/.sql/...) werden als UTF-8-Text
    # geladen und später durch den AST/Regex-Chunker geschickt. Prose-Dateien laufen
    # weiter durch den klassischen _chunk_text.
    _CODE_SUFFIXES = {
        ".py", ".js", ".jsx", ".mjs", ".cjs", ".ts", ".tsx",
        ".html", ".htm", ".css", ".scss", ".sass", ".sql",
        ".yaml", ".yml",
    }

    if suffix in (".txt", ".md"):
        text = raw_bytes.decode("utf-8", errors="replace")
    elif suffix in _CODE_SUFFIXES:
        text = raw_bytes.decode("utf-8", errors="replace")
    elif suffix == ".docx":
        if not _DOCX_OK:
            raise HTTPException(422, "python-docx nicht installiert. Bitte 'pip install python-docx' ausführen.")
        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            raise HTTPException(422, f"DOCX konnte nicht gelesen werden: {e}")
    elif suffix == ".pdf":
        if not _PDF_OK:
            raise HTTPException(422, "pdfplumber nicht installiert. Bitte 'pip install pdfplumber' ausführen.")
        try:
            pages = []
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
            text = "\n\n".join(pages)
        except Exception as e:
            raise HTTPException(422, f"PDF konnte nicht gelesen werden: {e}")
    elif suffix == ".json":
        # Patch 110: JSON als pretty-printed Text indizieren — erhält Struktur,
        # bleibt als Klartext für Embedding-Modell lesbar.
        try:
            data = json.loads(raw_bytes.decode("utf-8", errors="replace"))
            text = json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            raise HTTPException(422, f"JSON konnte nicht geparst werden: {e}")
    elif suffix == ".csv":
        # Patch 110: CSV als Text indizieren — Header-Zeile zuerst, dann jede Daten-
        # zeile mit "Header: Wert"-Mapping. Hilft Embedding-Modell, Spaltenbezüge
        # zu lernen. Leere Zellen werden ausgelassen.
        import csv as _csv
        from io import StringIO as _StringIO
        try:
            raw_text = raw_bytes.decode("utf-8", errors="replace")
            reader = _csv.reader(_StringIO(raw_text))
            rows = list(reader)
            if not rows:
                raise HTTPException(400, "CSV-Datei ist leer.")
            header = rows[0]
            lines = [", ".join(header)]
            for row in rows[1:]:
                parts = [f"{header[i]}: {v}" for i, v in enumerate(row) if i < len(header) and v.strip()]
                if parts:
                    lines.append("; ".join(parts))
            text = "\n".join(lines)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(422, f"CSV konnte nicht gelesen werden: {e}")
    else:
        raise HTTPException(
            422,
            f"Nicht unterstütztes Format '{suffix}'. Erlaubt: .txt, .md, .docx, .pdf, .json, .csv, "
            f".py, .js/.jsx/.ts/.tsx, .html, .css/.scss, .yaml/.yml, .sql"
        )

    text = text.strip()
    if not text:
        raise HTTPException(400, "Datei ist leer oder enthält keinen lesbaren Text.")

    # Patch 122: Code-Chunker-Weiche. Für bekannte Code-Extensions versuchen wir
    # den semantischen Chunker (AST/Regex). Liefert er etwas zurück, überspringen
    # wir die klassische Prose-Pipeline und speichern die reichhaltige Metadata.
    from zerberus.modules.rag.code_chunker import (
        chunk_code, is_code_file, describe_chunker,
    )

    code_chunks: list[dict] = []
    chunker_strategy = "prose"
    if is_code_file(filename):
        try:
            code_chunks = chunk_code(text, filename)
        except Exception as e:
            logger.warning(f"[CHUNK-122] Code-Chunker für {filename} fehlgeschlagen: {e} — Fallback Prose")
            code_chunks = []
        if code_chunks:
            chunker_strategy = describe_chunker(filename)

    if code_chunks:
        logger.info(
            f"[RAG-Chunking-122] Doc={filename}, Category={category}, "
            f"Strategy={chunker_strategy}, Chunks={len(code_chunks)}"
        )
        await _ensure_init(settings)
        indexed = 0
        chunk_preview: list[str] = []
        for chunk in code_chunks:
            content = chunk.get("content", "")
            meta = chunk.get("metadata", {})
            vec = await asyncio.to_thread(_encode, content)
            word_count = len(content.split())
            extra_meta = {
                "source": filename,
                "word_count": word_count,
                "category": category,
                "chunk_type": meta.get("chunk_type"),
                "name": meta.get("name"),
                "language": meta.get("language"),
                "start_line": meta.get("start_line"),
                "end_line": meta.get("end_line"),
                "chunker_strategy": chunker_strategy,
            }
            await asyncio.to_thread(_add_to_index, vec, content, extra_meta, settings)
            indexed += 1
            if len(chunk_preview) < 8:
                label = meta.get("name") or meta.get("chunk_type") or "chunk"
                chunk_preview.append(f"{meta.get('chunk_type','?')}:{label}")

        logger.info(
            f"✅ RAG-Upload (Code) abgeschlossen: {indexed}/{len(code_chunks)} Chunks aus "
            f"'{filename}' indiziert (Strategy: {chunker_strategy}, Kategorie: {category})"
        )
        return {
            "status": "ok",
            "filename": filename,
            "category": category,
            "auto_detected": auto_detected,
            "chunks_indexed": indexed,
            "merged_residuals": 0,
            "chunker_strategy": chunker_strategy,
            "chunk_preview": chunk_preview,
        }

    # Patch 110: Chunking-Strategie pro Category (CHUNK_CONFIGS). chunk_size / overlap /
    # min_chunk_words kommen aus dem Profil. config.yaml `modules.rag.min_chunk_words`
    # überschreibt den Category-Default, falls gesetzt (Rückwärtskompat Patch 88).
    rag_cfg_full = settings.modules.get("rag", {})
    min_words_override = rag_cfg_full.get("min_chunk_words")
    chunks, merged_count = _chunk_text(
        text,
        category=category,
        min_chunk_words=int(min_words_override) if min_words_override is not None else None,
    )
    cfg_used = CHUNK_CONFIGS.get(category, CHUNK_CONFIGS["general"])
    if not chunks:
        raise HTTPException(400, "Kein Text nach dem Chunking übrig.")

    logger.info(
        f"[RAG-Chunking] Doc={filename}, Category={category}, Chunks={len(chunks)} (nach Merge), "
        f"merged={merged_count} kurze Residuals, "
        f"profile={cfg_used['chunk_size']}/{cfg_used['overlap']}/{cfg_used['min_chunk_words']} "
        f"split={cfg_used['split']}"
    )

    await _ensure_init(settings)

    indexed = 0
    for chunk in chunks:
        vec = await asyncio.to_thread(_encode, chunk)
        word_count = len(chunk.split())
        await asyncio.to_thread(
            _add_to_index, vec, chunk,
            {"source": filename, "word_count": word_count, "category": category,
             "chunker_strategy": "prose"},
            settings,
        )
        indexed += 1

    logger.info(
        f"✅ RAG-Upload abgeschlossen: {indexed}/{len(chunks)} Chunks aus '{filename}' "
        f"indiziert (Kategorie: {category})"
    )
    return {
        "status": "ok",
        "filename": filename,
        "category": category,
        "auto_detected": auto_detected,
        "chunks_indexed": indexed,
        "merged_residuals": merged_count,
        "chunker_strategy": "prose",
        "chunk_preview": [],
    }


@router.get("/admin/rag/status")
async def rag_status():
    """Gibt aktuelle Index-Größe und Liste der indizierten Quellen zurück.

    Patch 108: Zusätzlich `sources_meta` mit (source, category) pro Chunk.
    `sources` bleibt aus Backward-Compat erhalten (nur Dateinamen).

    Patch 169 (B2): Lazy-Init des RAG-Moduls hier explizit ansteuern. Vorher
    konnten ``_index`` und ``_metadata`` als ``None``/``[]`` gelesen werden,
    weil die globalen Variablen erst beim ersten Search/Upload aus den
    On-Disk-Dateien rehydriert wurden. Folge: Hel zeigte „0 Dokumente" bis
    zum ersten Schreibvorgang, danach plötzlich der vollständige Bestand.
    """
    from zerberus.modules.rag.router import _ensure_init, _index, _metadata
    settings = get_settings()
    rag_cfg = settings.modules.get("rag", {}) or {}
    if rag_cfg.get("enabled", False):
        try:
            await _ensure_init(settings)
        except Exception as e:
            logger.warning("[RAG-169] Lazy-Init fehlgeschlagen: %s", e)
    # Re-Import nach Init: die Modul-Globals wurden ggf. gerade befuellt.
    from zerberus.modules.rag.router import _index, _metadata  # noqa: F811
    total = _index.ntotal if _index is not None else 0
    # Patch 116: Soft-deleted Chunks aus der Listen-Ansicht ausblenden.
    visible = [m for m in _metadata if m.get("deleted") is not True]
    sources = [m.get("source", "unbekannt") for m in visible]
    sources_meta = [
        {
            "source": m.get("source", "unbekannt"),
            "category": m.get("category", "general"),
        }
        for m in visible
    ]
    logger.info(
        "[RAG-169] Index-Status: %d Chunks, %d aktive, %d Quellen",
        total, len(visible), len({s for s in sources}),
    )
    return {
        "total_chunks": total,
        "active_chunks": len(visible),
        "sources": sources,
        "sources_meta": sources_meta,
    }


@router.delete("/admin/rag/clear")
async def rag_clear():
    """Leert den FAISS-Index und metadata.json komplett."""
    from zerberus.modules.rag.router import _reset_sync, RAG_AVAILABLE
    settings = get_settings()
    mod_cfg = settings.modules.get("rag", {})
    if not mod_cfg.get("enabled", False):
        raise HTTPException(503, "RAG-Modul ist in config.yaml deaktiviert.")
    if not RAG_AVAILABLE:
        raise HTTPException(503, "RAG-Abhängigkeiten nicht installiert.")
    await asyncio.to_thread(_reset_sync, settings)
    return {"status": "ok", "message": "RAG-Index geleert."}


@router.post("/admin/rag/reindex")
async def rag_reindex():
    """Baut den FAISS-Index aus den gespeicherten Metadaten-Chunks neu auf.
    Gewählte Option: Endpunkt statt Auto-Clear beim Start — kein ungewollter
    Datenverlust beim Neustart; User triggert Reindex bewusst nach Chunk-Änderung.
    Sinnvoll nach: Embedding-Modell-Wechsel oder Chunk-Parameter-Änderung
    (sofern alter Index dann gelöscht und Dokumente neu hochgeladen wurden).
    """
    from zerberus.modules.rag.router import (
        _reset_sync, _encode, _add_to_index, _metadata as current_meta, RAG_AVAILABLE
    )
    settings = get_settings()
    mod_cfg = settings.modules.get("rag", {})
    if not mod_cfg.get("enabled", False):
        raise HTTPException(503, "RAG-Modul ist in config.yaml deaktiviert.")
    if not RAG_AVAILABLE:
        raise HTTPException(503, "RAG-Abhängigkeiten nicht installiert.")

    # Kopie der aktuellen Chunks sichern bevor Reset
    # Patch 116: Soft-deleted Chunks beim Reindex weglassen (physische Bereinigung).
    chunks_to_reindex = [
        (m.get("text", ""), {k: v for k, v in m.items() if k not in ("text", "deleted")})
        for m in current_meta
        if m.get("text", "").strip() and m.get("deleted") is not True
    ]

    if not chunks_to_reindex:
        return {"status": "ok", "message": "Index war leer — nichts zu reindizieren.", "reindexed": 0}

    logger.info(f"🔄 RAG-Reindex gestartet: {len(chunks_to_reindex)} Chunks werden neu eingebettet")

    # Index leeren
    await asyncio.to_thread(_reset_sync, settings)

    # Alle Chunks neu einbetten und einfügen
    reindexed = 0
    for text, meta in chunks_to_reindex:
        vec = await asyncio.to_thread(_encode, text)
        await asyncio.to_thread(_add_to_index, vec, text, meta, settings)
        reindexed += 1

    logger.info(f"✅ RAG-Reindex abgeschlossen: {reindexed} Chunks neu indiziert")
    return {"status": "ok", "reindexed": reindexed}


# ============================================================
# Patch 116: Gruppierte Dokumentenliste + Soft-Delete pro Source
# ============================================================

@router.get("/admin/rag/documents")
async def rag_documents():
    """Gibt Dokumente gruppiert nach `source` zurück — ein Eintrag pro Datei.

    Patch 169 (B2): wie ``rag_status`` mit explizitem Lazy-Init, damit der
    Hel-RAG-Tab auch direkt nach Server-Start korrekte Zahlen zeigt.
    """
    from zerberus.modules.rag.router import _ensure_init
    settings = get_settings()
    rag_cfg = settings.modules.get("rag", {}) or {}
    if rag_cfg.get("enabled", False):
        try:
            await _ensure_init(settings)
        except Exception as e:
            logger.warning("[RAG-169] Lazy-Init fehlgeschlagen: %s", e)
    from zerberus.modules.rag.router import _index, _metadata
    total_chunks = _index.ntotal if _index is not None else 0
    grouped: dict[str, dict] = {}
    for meta in _metadata:
        if meta.get("deleted") is True:
            continue
        src = meta.get("source", "unbekannt")
        cat = meta.get("category", "general")
        wc = int(meta.get("word_count", 0) or 0)
        if src not in grouped:
            grouped[src] = {
                "source": src,
                "category": cat,
                "chunk_count": 0,
                "total_words": 0,
                "categories": {},
            }
        grouped[src]["chunk_count"] += 1
        grouped[src]["total_words"] += wc
        grouped[src]["categories"][cat] = grouped[src]["categories"].get(cat, 0) + 1

    documents = []
    for src, info in grouped.items():
        cats = info.pop("categories")
        # Category für Badge: häufigste Kategorie (Fallback: general)
        if cats:
            info["category"] = max(cats.items(), key=lambda kv: kv[1])[0]
        documents.append(info)

    documents.sort(key=lambda d: d["source"].lower())
    return {
        "documents": documents,
        "total_chunks": total_chunks,
        "total_documents": len(documents),
    }


@router.delete("/admin/rag/document")
async def rag_document_delete(source: str):
    """Soft-Delete aller Chunks einer Source. FAISS-Index bleibt intakt;
    Chunks mit `deleted: true` werden beim Reindex weggelassen und in
    `/admin/rag/documents` nicht mehr gelistet.

    P213-pre-Hotfix: Lazy-Init explizit anstossen — analog zu ``rag_documents``
    und ``rag_status`` (P169). Vorher iterierte der Endpunkt direkt ueber das
    Modul-Globale ``_metadata``, das beim allerersten Request nach Server-
    Start noch leer war (Hydrierung passiert in ``_ensure_init``). Folge:
    Sync-Tool bekam 404, der nachgelagerte UPLOAD-Schritt triggerte dann
    das Lazy-Init und addierte die neuen Chunks zu den vom Disk geladenen
    alten — Duplikate im Index.
    """
    from zerberus.modules.rag.router import (
        _ensure_init, _metadata, _resolve_paths, _save_metadata, RAG_AVAILABLE,
    )
    settings = get_settings()
    src = (source or "").strip()
    if not src:
        raise HTTPException(400, "source darf nicht leer sein.")

    rag_cfg = settings.modules.get("rag", {}) or {}
    if rag_cfg.get("enabled", False) and RAG_AVAILABLE:
        try:
            await _ensure_init(settings)
        except Exception as e:
            logger.warning("[RAG-213-pre] Lazy-Init im Delete fehlgeschlagen: %s", e)
    # Re-Import nach Init: das Modul-Globale wurde ggf. gerade befuellt.
    from zerberus.modules.rag.router import _metadata  # noqa: F811

    affected = 0
    for meta in _metadata:
        if meta.get("source") == src and not meta.get("deleted"):
            meta["deleted"] = True
            affected += 1

    if affected == 0:
        raise HTTPException(404, f"Keine Chunks mit source={src!r} gefunden.")

    _, meta_path = _resolve_paths(settings)
    _save_metadata(_metadata, meta_path)
    logger.warning(f"[RAG-116] Soft-Delete: {affected} Chunk(s) aus {src!r} markiert")
    return {"status": "ok", "source": src, "chunks_removed": affected}


# ============================================================
# Patch 115: Background Memory Extraction — manueller Trigger
# ============================================================

@router.post("/admin/memory/extract")
async def memory_extract():
    """Manueller Trigger für Background Memory Extraction.
    Läuft synchron (wartet auf Ergebnis). Dupliziert die Logik nicht —
    ruft `extract_memories()` aus dem Memory-Modul auf.
    """
    from zerberus.modules.memory.extractor import extract_memories
    settings = get_settings()
    mem_cfg = settings.modules.get("memory", {}) or {}
    result = await extract_memories(mem_cfg)
    return {"status": "ok", **result}


# ============================================================
# Pacemaker-Konfiguration (Säule 2 – Patch 56)
# ============================================================

# ---------------------------------------------------------------------------
# Patch 150 (B-024): Pacemaker-Prozess-Steuerung.
# Granulare Kontrolle pro Prozess (Sentiment/Memory/Dedup/…): Intervall, Device,
# Master-Schalter. Persistiert in config.yaml → modules.pacemaker_processes.
# Der eigentliche Scheduler muss die Config lesen — dieser Patch legt die
# Infrastruktur an, die Scheduler-Integration kann in einem Folge-Patch ergänzen.
# ---------------------------------------------------------------------------

PACEMAKER_DEFAULT_PROCESSES = [
    {"key": "sentiment",    "name": "Sentiment-Analyse",  "enabled": True,  "interval_min": 4,  "device": "cuda"},
    {"key": "memory",       "name": "Memory-Extraktion",  "enabled": True,  "interval_min": 4,  "device": "cuda"},
    {"key": "db_dedup",     "name": "DB-Deduplizierung",  "enabled": True,  "interval_min": 30, "device": "cpu"},
    {"key": "whisper_ping", "name": "Whisper-Ping",       "enabled": True,  "interval_min": 4,  "device": "cpu"},
]


@router.get("/admin/pacemaker/processes")
async def get_pacemaker_processes():
    """Liefert Pacemaker-Prozess-Konfiguration. Fallback: Default-Liste."""
    try:
        with open(CONFIG_YAML_PATH, "r", encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
    except Exception:
        cfg = {}
    pm_cfg = (cfg.get("modules", {}) or {}).get("pacemaker_processes", {}) or {}
    processes = pm_cfg.get("processes") if isinstance(pm_cfg.get("processes"), list) else None
    return {
        "master": pm_cfg.get("master", True),
        "sync":   pm_cfg.get("sync", False),
        "processes": processes or PACEMAKER_DEFAULT_PROCESSES,
        "current_activity": pm_cfg.get("current_activity"),  # wird vom Scheduler gesetzt
    }


@router.post("/admin/pacemaker/processes")
@invalidates_settings  # Patch 156: Cache nach YAML-Write neu laden
async def post_pacemaker_processes(request: Request):
    """Speichert Pacemaker-Prozess-Konfiguration in config.yaml."""
    data = await request.json()
    try:
        with open(CONFIG_YAML_PATH, "r", encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
    except Exception:
        cfg = {}
    cfg.setdefault("modules", {})["pacemaker_processes"] = {
        "master": bool(data.get("master", True)),
        "sync":   bool(data.get("sync", False)),
        "processes": data.get("processes", PACEMAKER_DEFAULT_PROCESSES),
    }
    with open(CONFIG_YAML_PATH, "w", encoding="utf-8") as f:
        yaml.safe_dump(cfg, f, sort_keys=False, allow_unicode=True)
    return {"status": "saved"}


@router.get("/admin/pacemaker/config")
async def get_pacemaker_config():
    """Gibt aktuelle Pacemaker-Konfiguration zurück."""
    settings = get_settings()
    pm = settings.legacy.pacemaker
    return {
        "interval_seconds": pm.interval_seconds,
        "keep_alive_minutes": pm.keep_alive_minutes,
    }


@router.post("/admin/pacemaker/config")
@invalidates_settings  # Patch 156: Cache nach YAML-Write neu laden
async def post_pacemaker_config(request: Request):
    """Speichert keep_alive_minutes in config.yaml (wirkt nach Neustart)."""
    data = await request.json()
    minutes = int(data.get("keep_alive_minutes", 120))
    if minutes < 1 or minutes > 480:
        raise HTTPException(400, "keep_alive_minutes muss zwischen 1 und 480 liegen.")

    if not CONFIG_YAML_PATH.exists():
        raise HTTPException(500, "config.yaml nicht gefunden.")

    with open(CONFIG_YAML_PATH, "r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f)

    cfg.setdefault("legacy", {}).setdefault("pacemaker", {})
    cfg["legacy"]["pacemaker"]["keep_alive_minutes"] = minutes

    temp_fd, temp_path = tempfile.mkstemp(dir=CONFIG_YAML_PATH.parent, suffix=".tmp")
    try:
        with os.fdopen(temp_fd, "w", encoding="utf-8") as f:
            yaml.dump(cfg, f, allow_unicode=True, default_flow_style=False, sort_keys=False)
        os.replace(temp_path, CONFIG_YAML_PATH)
    except Exception as e:
        if os.path.exists(temp_path):
            os.unlink(temp_path)
        raise HTTPException(500, f"Fehler beim Schreiben: {e}")

    return {"status": "ok", "keep_alive_minutes": minutes}


# ============================================================
# Whisper Docker Restart (Patch 119)
# ============================================================

@router.post("/admin/whisper/restart")
async def restart_whisper_endpoint():
    """Manueller Whisper-Container-Restart aus dem Hel-Dashboard."""
    from zerberus.whisper_watchdog import restart_whisper_container, check_whisper_health

    loop = asyncio.get_event_loop()
    success = await loop.run_in_executor(None, restart_whisper_container)
    await asyncio.sleep(10)
    healthy = await check_whisper_health()

    return {
        "restart_success": success,
        "post_restart_healthy": healthy,
    }


# ============================================================
# Profil-Übersicht (Patch 61) – readonly, kein password_hash
# ============================================================

@router.get("/admin/profiles")
async def get_profiles():
    """
    Gibt alle konfigurierten Profile zurück – ohne password_hash.
    Patch 61: Für die readonly Profil-Übersicht im Hel-Dashboard.
    """
    if not CONFIG_YAML_PATH.exists():
        return []
    with open(CONFIG_YAML_PATH, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    profiles = data.get("profiles", {})
    result = []
    for key, val in profiles.items():
        if not isinstance(val, dict):
            continue
        result.append({
            "key": key,
            "display_name": val.get("display_name", key),
            "permission_level": val.get("permission_level", "guest"),
            "allowed_model": val.get("allowed_model"),
            "temperature": val.get("temperature"),
        })
    return result


# ============================================================
# User-Verwaltung: Hash-Test & Passwort-Reset (Patch 83)
# ============================================================

@router.post("/admin/auth/test-hash")
async def test_hash(request: Request):
    """Testet ob ein Passwort zum gespeicherten Hash eines Profils passt."""
    import bcrypt
    body = await request.json()
    profile_key = body.get("profile", "").strip()
    password = body.get("password", "")

    if not profile_key or not password:
        raise HTTPException(status_code=400, detail="profile und password erforderlich")

    if not CONFIG_YAML_PATH.exists():
        raise HTTPException(status_code=500, detail="config.yaml nicht gefunden")

    with open(CONFIG_YAML_PATH, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    profiles = data.get("profiles", {})
    profile = profiles.get(profile_key)
    if not profile:
        raise HTTPException(status_code=404, detail=f"Profil '{profile_key}' nicht gefunden")

    stored_hash = profile.get("password_hash", "")
    if not stored_hash:
        return {"match": False, "hash_exists": False}

    try:
        match = bcrypt.checkpw(password.encode("utf-8"), stored_hash.encode("utf-8"))
    except Exception as e:
        logger.error(f"[DEBUG-83] Hash-Test Fehler: {e}")
        return {"match": False, "hash_exists": True, "error": str(e)}

    return {"match": match, "hash_exists": True}


@router.post("/admin/auth/reset-password")
@invalidates_settings  # Patch 156: Cache nach YAML-Write neu laden (ersetzt manuellen reload_settings())
async def reset_password(request: Request):
    """Setzt das Passwort eines Profils neu (bcrypt, rounds=12)."""
    import bcrypt
    body = await request.json()
    profile_key = body.get("profile", "").strip()
    new_password = body.get("password", "")

    if not profile_key or len(new_password) < 4:
        raise HTTPException(status_code=400, detail="profile erforderlich, password mind. 4 Zeichen")

    if not CONFIG_YAML_PATH.exists():
        raise HTTPException(status_code=500, detail="config.yaml nicht gefunden")

    with open(CONFIG_YAML_PATH, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    profiles = data.get("profiles", {})
    if profile_key not in profiles:
        raise HTTPException(status_code=404, detail=f"Profil '{profile_key}' nicht gefunden")

    new_hash = bcrypt.hashpw(new_password.encode("utf-8"), bcrypt.gensalt(rounds=12)).decode("utf-8")
    data["profiles"][profile_key]["password_hash"] = new_hash

    temp_fd, temp_path = tempfile.mkstemp(dir=CONFIG_YAML_PATH.parent, suffix=".tmp")
    try:
        with os.fdopen(temp_fd, "w", encoding="utf-8") as tmp:
            yaml.dump(data, tmp, allow_unicode=True, default_flow_style=False, sort_keys=False)
        os.replace(temp_path, CONFIG_YAML_PATH)
    except Exception as e:
        if os.path.exists(temp_path):
            os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"Fehler beim Speichern: {e}")

    # Patch 156: reload via @invalidates_settings — kein manueller Call mehr noetig
    logger.info(f"[DEBUG-83] Passwort für Profil '{profile_key}' zurückgesetzt.")
    return {"success": True, "profile": profile_key}


# ============================================================
# Provider-Blacklist (Patch 63)
# ============================================================

@router.get("/admin/provider_blacklist")
async def get_provider_blacklist():
    """Gibt aktuelle OpenRouter Provider-Blacklist zurück."""
    if not CONFIG_YAML_PATH.exists():
        return {"blacklist": []}
    with open(CONFIG_YAML_PATH, "r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f) or {}
    blacklist = cfg.get("openrouter", {}).get("provider_blacklist", [])
    return {"blacklist": blacklist}


@router.post("/admin/provider_blacklist")
@invalidates_settings  # Patch 156: Cache nach YAML-Write neu laden (ersetzt manuellen reload_settings())
async def post_provider_blacklist(request: Request):
    """Speichert neue Provider-Blacklist in config.yaml und lädt Settings neu."""
    data = await request.json()
    blacklist = data.get("blacklist", [])
    if not isinstance(blacklist, list):
        raise HTTPException(400, "blacklist muss eine Liste sein.")

    if not CONFIG_YAML_PATH.exists():
        raise HTTPException(500, "config.yaml nicht gefunden.")

    with open(CONFIG_YAML_PATH, "r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f) or {}

    cfg.setdefault("openrouter", {})["provider_blacklist"] = [str(p).lower() for p in blacklist]

    temp_fd, temp_path = tempfile.mkstemp(dir=CONFIG_YAML_PATH.parent, suffix=".tmp")
    try:
        with os.fdopen(temp_fd, "w", encoding="utf-8") as f:
            yaml.dump(cfg, f, allow_unicode=True, default_flow_style=False, sort_keys=False)
        os.replace(temp_path, CONFIG_YAML_PATH)
    except Exception as e:
        if os.path.exists(temp_path):
            os.unlink(temp_path)
        raise HTTPException(500, f"Fehler beim Schreiben: {e}")

    # Patch 156: reload via @invalidates_settings — kein manueller Call mehr noetig
    return {"status": "ok"}


# ============================================================
# Metrics History mit BERT-Sentiment (Patch 57)
# ============================================================

@router.get("/metrics/profiles")
async def metrics_profiles():
    """
    Patch 95: Liefert die distinct profile_keys aus `interactions` für den
    Per-User-Filter im Hel-Metriken-Dashboard. Leere Liste falls die
    Patch-92-Spalte noch nicht existiert oder keine Daten vorhanden sind.
    """
    from zerberus.core.database import _async_session_maker
    from sqlalchemy import text

    async with _async_session_maker() as session:
        col_result = await session.execute(text("PRAGMA table_info(interactions)"))
        cols = {row[1] for row in col_result.fetchall()}
        if "profile_key" not in cols:
            return {"profiles": []}
        rows = await session.execute(text(
            "SELECT DISTINCT profile_key FROM interactions "
            "WHERE profile_key IS NOT NULL AND profile_key != '' "
            "ORDER BY profile_key"
        ))
        return {"profiles": [r[0] for r in rows.fetchall()]}


@router.get("/metrics/history")
async def metrics_history(
    limit: int = 200,
    from_date: Optional[str] = None,
    to_date: Optional[str] = None,
    profile_key: Optional[str] = None,
):
    """
    Gibt N ausgewertete Messages zurück inkl. aller Chart-Metriken.

    Patch 91: Optional Zeitraum-Filter (ISO-Dates) und `profile_key`-Filter.
    `profile_key` wird nur angewendet wenn die Spalte existiert (Patch 92 legt sie an);
    sonst wird der Parameter ignoriert.

    Response: {"meta": {...}, "results": [...]} — results ist das alte Array
    (abwärtskompatibel: Frontend kann body.results oder body direkt lesen).
    """
    from zerberus.core.database import _async_session_maker
    from sqlalchemy import text

    async with _async_session_maker() as session:
        # Spaltenliste dynamisch prüfen
        col_result = await session.execute(text("PRAGMA table_info(message_metrics)"))
        cols = {row[1] for row in col_result.fetchall()}
        has_bert = "bert_sentiment_label" in cols and "bert_sentiment_score" in cols

        # Patch 92-ready: profile_key-Filter nur wenn Spalte existiert
        int_cols_result = await session.execute(text("PRAGMA table_info(interactions)"))
        int_cols = {row[1] for row in int_cols_result.fetchall()}
        has_profile_key = "profile_key" in int_cols

        # Patch 64: BERT-Score als gerichteter Wert (0=negativ, 0.5=neutral, 1=positiv).
        bert_cols = (
            """, mm.bert_sentiment_label,
    CASE
        WHEN mm.bert_sentiment_label = 'positive' THEN mm.bert_sentiment_score
        WHEN mm.bert_sentiment_label = 'negative' THEN (1.0 - mm.bert_sentiment_score)
        WHEN mm.bert_sentiment_label = 'neutral'  THEN 0.5
        ELSE (i.sentiment + 1.0) / 2.0
    END as bert_sentiment_score"""
            if has_bert
            else ", NULL as bert_sentiment_label, (i.sentiment + 1.0) / 2.0 as bert_sentiment_score"
        )

        where_clauses = ["i.role = 'user'"]
        params: dict = {"limit": limit}

        if from_date:
            where_clauses.append("i.timestamp >= :from_date")
            params["from_date"] = from_date
        if to_date:
            # inklusive Tagesende
            where_clauses.append("i.timestamp <= :to_date")
            params["to_date"] = to_date + " 23:59:59" if len(to_date) == 10 else to_date
        if profile_key and has_profile_key:
            where_clauses.append("i.profile_key = :profile_key")
            params["profile_key"] = profile_key

        where_sql = " AND ".join(where_clauses)

        try:
            result = await session.execute(text(f"""
                SELECT i.id, i.timestamp, i.role, i.session_id, i.content,
                       mm.word_count, mm.ttr, mm.shannon_entropy, mm.vader_compound
                       {bert_cols}
                FROM interactions i
                JOIN message_metrics mm ON i.id = mm.message_id
                WHERE {where_sql}
                ORDER BY i.timestamp DESC
                LIMIT :limit
            """), params)
            rows = result.fetchall()
        except Exception as e:
            logger.warning(f"[metrics/history] Fehler: {e}")
            return {
                "meta": {"from": from_date, "to": to_date, "count": 0, "profile_key": profile_key, "error": str(e)},
                "results": [],
            }

        # Rolling-Window-TTR (Patch 64)
        import re as _re
        ROLLING_WINDOW = 50
        raw_rows = [dict(row._mapping) for row in rows]
        chron = list(reversed(raw_rows))          # älteste zuerst
        token_lists = [
            _re.findall(r'\b\w+\b', (r.get("content") or "").lower())
            for r in chron
        ]
        # Zusätzliche Metriken: hapax_ratio + avg_word_length (für Patch 91 Chart)
        for i, row in enumerate(chron):
            start = max(0, i - ROLLING_WINDOW + 1)
            window_tokens: list = []
            for tl in token_lists[start:i + 1]:
                window_tokens.extend(tl)
            if window_tokens:
                row["rolling_ttr"] = round(len(set(window_tokens)) / len(window_tokens), 3)
            else:
                row["rolling_ttr"] = row.get("ttr")

            # Per-Message hapax + avg_word_length
            msg_tokens = token_lists[i]
            if msg_tokens:
                from collections import Counter
                counts = Counter(msg_tokens)
                hapax = sum(1 for _, c in counts.items() if c == 1)
                row["hapax_ratio"] = round(hapax / len(msg_tokens), 3)
                row["avg_word_length"] = round(sum(len(t) for t in msg_tokens) / len(msg_tokens), 2)
            else:
                row["hapax_ratio"] = None
                row["avg_word_length"] = None

            # Alias bert_sentiment (ohne _score) für Frontend-Kompakt-Key
            row["bert_sentiment"] = row.get("bert_sentiment_score")
            # created_at-Alias (Frontend erwartet created_at)
            row["created_at"] = row.get("timestamp")
            row.pop("content", None)

        results = list(reversed(chron))
        return {
            "meta": {
                "from": from_date,
                "to": to_date,
                "count": len(results),
                "profile_key": profile_key,
                "profile_key_supported": has_profile_key,
            },
            "results": results,
        }


# ============================================================
# Patch 96 – Testreport-Viewer (H-F04)
# ============================================================

# zerberus/app/routers/hel.py.parents[2] == zerberus/  →  zerberus/tests/report
_REPORT_DIR = Path(__file__).resolve().parents[2] / "tests" / "report"


@router.get("/tests/report", response_class=HTMLResponse)
async def tests_report():
    """Letzten Playwright-HTML-Report ausliefern (full_report.html)."""
    p = _REPORT_DIR / "full_report.html"
    if not p.exists():
        return JSONResponse(
            {"error": "Kein Testreport vorhanden. Bitte pytest ausführen."},
            status_code=404,
        )
    try:
        return HTMLResponse(p.read_text(encoding="utf-8"))
    except Exception as e:
        logger.warning(f"[tests/report] Lesefehler: {e}")
        return JSONResponse({"error": f"Lesefehler: {e}"}, status_code=500)


# Patch 170 (B5): Einzelne Reports per Name ausliefern
_ALLOWED_REPORT_NAMES = {"full_report", "fenrir_report", "loki_report"}


@router.get("/tests/report/{name}", response_class=HTMLResponse)
async def tests_report_named(name: str):
    """Einzelnen HTML-Report per Name ausliefern (z.B. fenrir_report, loki_report).

    Whitelist verhindert Path-Traversal — nur bekannte Report-Namen sind erlaubt.
    """
    if name not in _ALLOWED_REPORT_NAMES:
        return JSONResponse(
            {"error": f"Unbekannter Report: {name}"},
            status_code=404,
        )
    p = _REPORT_DIR / f"{name}.html"
    if not p.exists():
        return JSONResponse(
            {"error": f"Kein {name}.html vorhanden. Bitte pytest ausführen."},
            status_code=404,
        )
    try:
        return HTMLResponse(p.read_text(encoding="utf-8"))
    except Exception as e:
        logger.warning(f"[tests/report/{name}] Lesefehler: {e}")
        return JSONResponse({"error": f"Lesefehler: {e}"}, status_code=500)


@router.get("/tests/reports")
async def tests_reports_list():
    """Alle .html-Reports im Report-Ordner mit mtime + size."""
    if not _REPORT_DIR.exists():
        return {"reports": []}
    files = sorted(
        _REPORT_DIR.glob("*.html"),
        key=lambda x: x.stat().st_mtime,
        reverse=True,
    )
    return {
        "reports": [
            {
                "name": f.name,
                "mtime": f.stat().st_mtime,
                "size": f.stat().st_size,
            }
            for f in files
        ]
    }


# ====================================================================
# Patch 191: Prosodie-Admin-Status
# ====================================================================
# Worker-Protection: NUR Aggregate (Counter, Modus, Timestamps).
# KEINE individuellen Mood/Valence/Arousal-Werte. Audio nicht in DB.
# Logging-Tags: [PROSODY-ADMIN-191]
# ====================================================================

@router.get("/admin/prosody/status")
async def admin_prosody_status():
    """Aggregierter Prosodie-Status für Hel-Admin (P191).

    Worker-Protection: gibt nur Modus, Counter und Timestamps zurück —
    NIEMALS individuelle Mood-/Valence-/Arousal-Werte. Audio-Bytes werden
    nicht in der DB gespeichert (siehe ProsodyManager.analyze).
    """
    settings = get_settings()
    from zerberus.modules.prosody.manager import get_prosody_manager
    mgr = get_prosody_manager(settings)
    status = mgr.admin_status()
    logger.info("[PROSODY-ADMIN-191] Status abgefragt")
    return status